from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/long_flood_duration_audit_france_2005_2024.docx")

SOURCES = [
    ("S1", "data.gouv.fr, Arrêtés de catastrophe naturelle par commune (GASPAR-derived; source updated September 2025)", "https://www.data.gouv.fr/datasets/arretes-de-catastrophe-naturelle-par-commune"),
    ("S2", "French Ministry for Ecological Transition, Key facts about floods in France", "https://www.ecologie.gouv.fr/en/public-policies/key-facts-about-floods-france"),
    ("S3", "BRGM, Groundwater, drought and flooding: frequently asked questions", "https://www.brgm.fr/fr/actualite/dossier-thematique/eau-souterraine-secheresse-inondations-faq-questions-frequentes"),
    ("S4", "Moselle Prefecture, accelerated CatNat procedure for the May 2024 floods", "https://www.moselle.gouv.fr/Actualites/Securite/Protection-publique-et-securite-civile/Inondations-mai-2024-Procedure-de-reconnaissance-de-l-Etat-de-catastrophe-naturelle-acceleree"),
    ("S5", "Balaruc-les-Bains municipality, request for CatNat recognition following the October 2023 inversac", "https://www.ville-balaruc-les-bains.com/actualites/inversac-demande-de-reconnaissance-de-catastrophe-naturelle"),
    ("S6", "Hauts-de-France DREAL, CatNat recognition after the heavy rains of November 2023", "https://www.hauts-de-france.developpement-durable.gouv.fr/IMG/pdf/cp_reconnaissance_de_l_etat_de_catastrophe_naturelle_pour_13_communes.pdf"),
    ("S7", "ARIA (French Ministry), erosion of a flood-protection dike at Hénouville on 13 February 2020", "https://www.aria.developpement-durable.gouv.fr/accident/55110/"),
    ("S8", "BRGM, hydrogeological expert report: groundwater flooding at Mazingarbe, February-April 2013", "https://infoterre.brgm.fr/rapports/RP-62372-FR.pdf"),
    ("S9", "Nord Prefecture, historical CatNat recognition table (includes Fenain 20 Nov 2013-6 Feb 2014)", "https://www.nord.gouv.fr/content/download/62567/393221/file/20190510_tableau_arretes_CATNAT.pdf"),
    ("S10", "BRGM, groundwater status on 1 January 2021", "https://www.brgm.fr/fr/actualite/communique-presse/nappes-eau-souterraine-au-1er-janvier-2021"),
    ("S11", "French Interior Ministry, report on the Seine and tributary floods of January-February 2018", "https://www.interieur.gouv.fr/documentation/rapports/crue-de-seine-et-de-ses-affluents-de-janvier-fevrier-2018.html"),
    ("S12", "Pas-de-Calais Department, historical floods from November 2023 and renewed record levels in January 2024", "https://www.pasdecalais.fr/inondations-2023-2024-retour-sur-les-aides-recues-par-le-departement-du-pas-de-calais-pour-la"),
    ("S13", "European Flood Awareness System, bulletin December 2023-January 2024", "https://european-flood.emergency.copernicus.eu/sites/default/files/bulletins-documents/2024/EFAS_Bimonthly_Bulletin_Dec2023_Jan2024_0.pdf"),
    ("S14", "Eaufrance, national hydrological bulletin for February 2023", "https://www.eaufrance.fr/publications/bsh/2023-02"),
    ("S15", "Météo-France, 2021: a year of strong contrasts", "https://meteofrance.com/actualites-et-dossiers/actualites/2021-une-annee-de-forts-contrastes-en-france"),
    ("S16", "Météo-France, summer 2024: locally violent thunderstorms and rainfall anomalies", "https://meteofrance.com/actualites-et-dossiers/actualites/ete-2024-chaud-orages-localement-violents-deux-vagues-chaleur"),
    ("S17", "Gironde Department, roads flooded and closed, 25 February 2024", "https://www.gironde.fr/espace-presse/point-de-situation-des-routes-inondees-et-fermees"),
    ("S18", "French Interior Ministry, 209 communes affected in January-February 2018 recognised as CatNat", "https://www.interieur.gouv.fr/archive/209-communes-touchees-par-les-inondations-en-janvier-et-fevrier-2018-reconnues-en-etat-de-catastrophe-naturelle"),
]

GASPAR = [
    (1, "31 Dec 2023-29 Jun 2024", 181, "Moselle: Buding, Guinkirchen, Hayange, Königsmacker, Rémering-lès-Puttelange, Schorbach", "Recognition record is valid; continuous surface flooding is not supported.", "Likely a groundwater/administrative window. The exceptional reported flood episode was 17-20 May 2024, far shorter than 181 days.", "S1, S3, S4"),
    (2, "16 Oct 2023-3 Apr 2024", 170, "Balaruc-les-Bains (Hérault)", "Valid CatNat period.", "Plausible as a prolonged groundwater/inversac phenomenon, but not as 170 days of uniformly flooded surface. The municipality reported a new inversac from 18 October.", "S1, S3, S5"),
    (3, "1 Feb-9 Jul 2020", 159, "Hénouville (Seine-Maritime)", "Valid CatNat period.", "A real local flooding/dike incident is documented on 13 February. No independent evidence found for uninterrupted flooding through 9 July; treat the full span as administrative.", "S1, S7"),
    (4, "2 Nov 2023-24 Mar 2024", 143, "Nampont, Vercourt and Vron (Somme)", "Valid CatNat periods for three communes.", "Heavy November rains and CatNat recognition are corroborated. The five-month span is better interpreted as a wet-season/groundwater damage period than one continuous flood.", "S1, S6"),
    (5, "11 Feb-20 Jun 2013", 129, "Angaïs (Pyrénées-Atlantiques)", "Valid CatNat period.", "Groundwater flooding can persist for months. The exact administrative dates are supported by the official dataset, but no independent event narrative was found; confidence in continuous inundation is low.", "S1, S3"),
    (6, "20 Dec 2013-16 Apr 2014", 117, "Ossun (Hautes-Pyrénées)", "Valid CatNat period.", "Hydrologically possible as groundwater flooding over a wet winter; evidence supports the recognition interval, not constant flood depth or extent.", "S1, S3"),
    (7, "1 Dec 2013-1 Mar 2014", 90, "Sermérieu (Isère)", "Valid CatNat period.", "Plausible only as a groundwater/impact window. No evidence found for a single continuous river flood lasting 90 days.", "S1, S3"),
    (8, "1 Dec 2020-28 Feb 2021", 89, "12 communes in Landes", "Valid multi-commune CatNat period.", "A long groundwater response is plausible; BRGM reported the lingering effect of abundant 2019-2020 recharge plus autumn/winter rainfall. Do not read it as 89 days of continuous surface inundation everywhere.", "S1, S3, S10"),
    (9, "1 Feb-25 Apr 2013", 83, "Mazingarbe (Pas-de-Calais)", "Valid CatNat period, independently examined by BRGM.", "The strongest GASPAR case for a genuinely prolonged phenomenon: BRGM explicitly investigated groundwater flooding from February to April 2013. Still, impacts likely varied within the interval.", "S1, S8"),
    (10, "4 Mar-25 May 2015", 82, "Angaïs (Pyrénées-Atlantiques)", "Valid CatNat period.", "Plausible as a groundwater episode; no independent exact-date narrative found. Administrative-period interpretation remains safest.", "S1, S3"),
    (10, "16 Jan-7 Apr 2024", 82, "Thaon (Calvados)", "Valid CatNat period; tied at rank 10.", "A multi-month groundwater impact is possible, but no independent exact-date narrative was found. Do not interpret as a continuously inundated commune.", "S1, S3"),
]

JRC = [
    (1, "1 Jan-1 Apr 2024", 91, "Pas-de-Calais; 7 communes", "Real winter flooding occurred, but the raster interval is an observation/merge envelope.", "November 2023 and January 2024 were distinct major waves; a single uninterrupted 91-day flood is not established.", "S12, S13"),
    (2, "10 Oct 2022-2 Jan 2023", 84, "North-western and northern France; 193 communes", "Spatially and temporally over-broad for one physical event.", "No authoritative evidence found for one continuous France flood over the full interval; likely merges separate detections.", "S2"),
    (3, "30 Jan-17 Apr 2023", 77, "Northern/western France; 212 communes", "Observation envelope, not a verified continuous flood.", "The February hydrological bulletin describes national conditions, but does not support a single 77-day flood across this footprint.", "S14"),
    (4, "1 Jan-5 Mar 2018", 63, "Nord and Pas-de-Calais; 8 communes", "Partly aligned with a genuine major flood season.", "The Seine and tributary floods are documented for January-February, but the JRC footprint and full 63-day span should not be equated with one continuous local flood.", "S11, S18"),
    (5, "28 Sep-30 Nov 2020", 63, "Atlantic/western France; 256 communes", "Likely aggregation of separate wet-weather detections.", "No single event source supports continuous flooding across the entire footprint and interval.", "S2"),
    (6, "14 Feb-18 Apr 2022", 63, "Brittany, Channel coast and north; 41 communes", "Low physical plausibility as one continuous flood.", "No corroborating authoritative event narrative found for the full interval; treat as a raster-clustering artifact until raw weekly detections are inspected.", "S2"),
    (7, "6 May-22 Jul 2024", 63, "Gironde, Landes and Pyrénées-Atlantiques; 36 communes", "Some storms/flood impacts occurred, but not a verified 63-day continuous event.", "Summer 2024 had locally violent storms and regional rainfall anomalies; this supports intermittent episodes, not continuous inundation.", "S16"),
    (8, "4 Jan-8 Mar 2021", 63, "Very broad France footprint; 852 communes", "Almost certainly an aggregation envelope.", "The national scale and commune count make a single physical flood implausible; 2021 was climatically contrasting, with multiple episodes.", "S15"),
    (9, "10 Feb-13 Apr 2020", 63, "Very broad France footprint; 1,557 communes", "Almost certainly an aggregation envelope.", "The exceptionally broad footprint is incompatible with one coherent continuous flood; inspect weekly components before analysis.", "S2"),
    (10, "30 Dec 2019-24 Feb 2020", 56, "Very broad France footprint; 1,675 communes", "Almost certainly an aggregation envelope.", "This is the broadest top-ten footprint and strongly indicates merged detections rather than a single flood.", "S2"),
]


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink"); hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "2E74B5"); rpr.append(color)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single"); rpr.append(underline)
    run.append(rpr); t = OxmlElement("w:t"); t.text = text; run.append(t); hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]; cell.width = Inches(w); shade(cell, "D9E2F3")
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.size = Pt(8.5)
    for row in rows:
        cells = table.add_row().cells
        for i, (value, w) in enumerate(zip(row, widths)):
            cells[i].width = Inches(w); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            if i in (0, 2): p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value)); r.font.size = Pt(8.2)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell in row.cells:
            set_cell_margins(cell)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    return table


def add_event_profiles(doc, heading, records):
    doc.add_heading(heading, level=1)
    for rank, dates, days, place, dataset, physical, refs in records:
        p = doc.add_paragraph(style="Heading 2")
        p.add_run(f"Rank {rank} — {days} days — {dates}")
        p = doc.add_paragraph(); p.add_run("Location/footprint: ").bold = True; p.add_run(place)
        p = doc.add_paragraph(); p.add_run("Dataset judgement: ").bold = True; p.add_run(dataset)
        p = doc.add_paragraph(); p.add_run("Physical-event judgement: ").bold = True; p.add_run(physical)
        p = doc.add_paragraph(); p.add_run("Evidence: ").bold = True; p.add_run(refs)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Title", 25, "17365D", 0, 8), ("Subtitle", 12, "5B6573", 0, 14),
        ("Heading 1", 16, "2E74B5", 16, 8), ("Heading 2", 12.5, "2E74B5", 10, 4),
        ("Heading 3", 11.5, "1F4D78", 8, 4)):
        s = styles[name]; s.font.name = "Calibri"; s.font.size = Pt(size); s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after)
    title = doc.add_paragraph(style="Title"); title.add_run("Audit of the Longest Flood Durations in France")
    sub = doc.add_paragraph(style="Subtitle"); sub.add_run("Top-ten GASPAR/CatNat and JRC records, 2005–2024 | Data validation and manual web verification")
    p = doc.add_paragraph(); p.add_run("Prepared: 31 July 2026\n").bold = True
    p.add_run("Scope: France; events intersecting the requested period. JRC source coverage available in the repository begins in 2015.")
    doc.add_paragraph("Bottom line", style="Heading 1")
    p = doc.add_paragraph(); p.add_run("The reported 316-day GASPAR and 288-day JRC floods are not reproduced. ").bold = True
    p.add_run("Using the current official GASPAR-derived CatNat file, the maximum unique recognition interval is 181 days. In the France JRC event table, the maximum is 91 days. The earlier number 288 is the count of unique JRC France events in the 2015–2024 comparison, not a duration.")
    p = doc.add_paragraph(); p.add_run("Interpretation rule: ").bold = True
    p.add_run("GASPAR dates are legal/administrative recognition periods attached to communes; JRC dates are raster observation and clustering envelopes. Neither variable proves that a place remained continuously inundated throughout the full interval.")
    doc.add_heading("Method and limitations", level=1)
    for text in (
        "GASPAR/CatNat: downloaded the current public commune-level file from data.gouv.fr (S1), filtered flood records whose start and end dates fall within 1 January 2005–31 December 2024, calculated end minus start, and grouped identical intervals into one ranked episode. Ties are retained.",
        "JRC: deduplicated the repository's France commune-event table by JRC event identifier, retained events intersecting 2005–2024, and ranked the supplied duration_days field. The local JRC archive starts in 2015, so the JRC ranking is effectively 2015–2024, not a complete 2005–2024 census.",
        "Manual verification: searched official government, BRGM, municipal, EFAS and Météo-France material for the dates and locations. 'No independent evidence found' means the exact long span was not corroborated in the sources reviewed; it does not prove that no damage occurred.",
        "Duration convention: reported days equal end date minus start date, matching the JRC filename convention. Calendar-day-inclusive counts would be one day higher.",
    ):
        doc.add_paragraph(text, style="List Bullet")
    doc.add_heading("Ranked overview", level=1)
    doc.add_heading("GASPAR/CatNat unique recognition periods", level=2)
    add_table(doc, ["Rank", "Dates", "Days", "Communes/area", "Assessment"],
              [(r, d, n, p, ("Plausible prolonged process" if r in (2, 8, 9) else "Administrative span; caution")) for r,d,n,p,*_ in GASPAR],
              [0.45, 1.2, 0.45, 2.55, 1.85])
    doc.add_paragraph("Source: current GASPAR-derived CatNat data (S1). Rank 10 is tied.", style="Caption")
    doc.add_heading("JRC France raster events", level=2)
    add_table(doc, ["Rank", "Dates", "Days", "France footprint", "Assessment"],
              [(r, d, n, p, "Observation/merge envelope") for r,d,n,p,*_ in JRC],
              [0.45, 1.2, 0.45, 2.55, 1.85])
    doc.add_paragraph("Source: repository JRC France canonical event table; manual corroboration sources listed below.", style="Caption")
    add_event_profiles(doc, "GASPAR/CatNat event-by-event verification", GASPAR)
    add_event_profiles(doc, "JRC event-by-event verification", JRC)
    doc.add_heading("Conclusions and recommended data treatment", level=1)
    for text in (
        "Do not label these values 'days flooded'. Use 'administrative recognition-period length' for GASPAR and 'JRC event-envelope duration' for JRC.",
        "Flag GASPAR durations above 60 days for groundwater/administrative review. Preserve the hazard subtype where available; the simplified current download labels all as 'Inondations'.",
        "For JRC durations above 35 days, inspect weekly raster components and split clusters when there are temporal gaps or disconnected footprints. The 852-, 1,557- and 1,675-commune events are especially clear aggregation warnings.",
        "Do not compare GASPAR and JRC duration values as if they measured the same phenomenon. Match using location, temporal overlap and event subtype, and retain separate confidence fields.",
        "If a 316- or 288-day row still appears in another workbook, trace its exact source file, event identifier and aggregation code; it is not present in the audited canonical France tables.",
    ):
        doc.add_paragraph(text, style="List Bullet")
    doc.add_heading("Local data provenance", level=1)
    doc.add_paragraph("GASPAR/CatNat ranking input: data/processed/catnat_latest_20260720.csv, downloaded from the S1 resource on 31 July 2026. The repository's older data/gaspar_floods.xlsx was last updated in 2022 and was not used for the final 2024-complete ranking.")
    doc.add_paragraph("JRC ranking input: data/processed/flood_outputs/jrc_gaspar_comparison_7d/jrc_france_commune_events_canonical.csv. The longest France event in that canonical table is 91 days.")
    doc.add_heading("References", level=1)
    for sid, label, url in SOURCES:
        p = doc.add_paragraph(); p.add_run(f"{sid}. {label}. ")
        add_hyperlink(p, urlparse(url).netloc + urlparse(url).path, url)
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("France flood-duration audit | GASPAR/CatNat and JRC | 2005–2024")
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
