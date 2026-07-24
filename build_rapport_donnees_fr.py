from pathlib import Path
import importlib.util, itertools, math, os, textwrap
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "outputs" / "rapport_alternance_donnees"
FIG = OUT / "figures"
OUT.mkdir(parents=True, exist_ok=True)
FIG.mkdir(parents=True, exist_ok=True)

spec = importlib.util.spec_from_file_location("shapefile", ROOT / ".venv/Lib/site-packages/shapefile.py")
shapefile = importlib.util.module_from_spec(spec)
spec.loader.exec_module(shapefile)

BLUE = "1F4E78"; MID = "2E75B6"; PALE = "DDEBF7"; TEAL = "2A7F7F"
GREEN = "548235"; AMBER = "C58A00"; RED = "C0504D"; GRAY = "666666"; LIGHT = "F3F6F8"

def font(size=22, bold=False):
    candidates = [
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / ("arialbd.ttf" if bold else "arial.ttf"),
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / ("calibrib.ttf" if bold else "calibri.ttf"),
    ]
    for p in candidates:
        if p.exists(): return ImageFont.truetype(str(p), size)
    return ImageFont.load_default()

def add_wrapped(draw, xy, text, width, fill, fnt, spacing=6, anchor=None):
    lines = textwrap.wrap(text, width=width)
    draw.multiline_text(xy, "\n".join(lines), fill=fill, font=fnt, spacing=spacing, anchor=anchor)

def flow_figure():
    im = Image.new("RGB", (1600, 720), "white"); d = ImageDraw.Draw(im)
    d.text((70, 45), "Complémentarité des jeux de données", fill="#1F4E78", font=font(42, True))
    boxes = [
        (80, 180, 390, 490, "#DDEBF7", "GASPAR", "Évènements administratifs\nCATNAT par commune\net par période"),
        (470, 180, 780, 490, "#E2F0D9", "TRI France / ISPRA Italie", "Emprises de danger ou\nde surface inondable\nselon des scénarios"),
        (860, 180, 1170, 490, "#FFF2CC", "Zones ripariennes", "Occupation du sol dans\nles corridors fluviaux\nCopernicus 2018"),
        (1250, 180, 1530, 490, "#EDEDED", "Décision", "Croisement spatial\n+ cohérence temporelle\n+ traçabilité"),
    ]
    for x1,y1,x2,y2,c,title,body in boxes:
        d.rounded_rectangle((x1,y1,x2,y2), 24, fill=c, outline="#A6A6A6", width=3)
        d.text(((x1+x2)//2,y1+58), title, fill="#1F1F1F", font=font(28,True), anchor="ma")
        d.multiline_text(((x1+x2)//2,y1+135), body, fill="#444444", font=font(24), spacing=10, anchor="ma", align="center")
    for x in (410, 800, 1190):
        d.line((x,335,x+40,335), fill="#2E75B6", width=8)
        d.polygon([(x+40,335),(x+22,322),(x+22,348)], fill="#2E75B6")
    d.text((800, 620), "Un évènement reconnu n’implique pas automatiquement que chaque point de la commune a été inondé.", fill="#9C5700", font=font(27,True), anchor="mm")
    p=FIG/"figure_01_complementarite.png"; im.save(p); return p

def project(pt, bounds, box):
    x,y=pt; xmin,ymin,xmax,ymax=bounds; l,t,r,b=box
    return (l+(x-xmin)/(xmax-xmin)*(r-l), b-(y-ymin)/(ymax-ymin)*(b-t))

def draw_shapes(draw, shp_path, bounds, box, outline, fill=None, max_shapes=9000, width=1):
    rd=shapefile.Reader(str(shp_path), encoding="latin1")
    step=max(1, math.ceil(len(rd)/max_shapes))
    for idx in range(0, len(rd), step):
        s = rd.shape(idx)
        pts=s.points
        starts=list(s.parts)+[len(pts)]
        for a,b in zip(starts,starts[1:]):
            segment = pts[a:b]
            point_step = max(1, math.ceil(len(segment) / 500))
            poly=[project(p,bounds,box) for p in segment[::point_step]]
            if len(poly)>2:
                draw.polygon(poly, fill=fill, outline=outline)

def riparian_figure():
    paths_fr = sorted((ROOT/"data/raw/France_Riparian").rglob("rpz_*.shp"))
    paths_it = sorted((ROOT/"data/raw/Italy_Riparian").rglob("rpz_*.shp"))
    panels=[("France — 6 unités de livraison",paths_fr,(3200000,1750000,4450000,3380000)),
            ("Italie — 3 unités de livraison",paths_it,(4000000,1500000,5070000,2670000))]
    im=Image.new("RGB",(1600,850),"white"); d=ImageDraw.Draw(im)
    d.text((70,35),"Zones ripariennes Copernicus 2018 — aperçu des couches du projet",fill="#1F4E78",font=font(38,True))
    for j,(title,paths,bounds) in enumerate(panels):
        box=(80+j*780,145,740+j*780,725)
        d.rectangle(box,fill="#F7FAFC",outline="#B4C6D7",width=3)
        for k,p in enumerate(paths):
            draw_shapes(d,p,bounds,box,outline="#2F6B4F",fill="#9FD5B3",max_shapes=2200)
        d.text(((box[0]+box[2])//2,105),title,fill="#2F6B4F",font=font(28,True),anchor="mm")
    d.text((800,790),"Polygones d’occupation/couverture du sol dans les corridors ripariens — projection ETRS89 / LAEA Europe (EPSG:3035).",fill="#555555",font=font(22),anchor="mm")
    p=FIG/"figure_02_riparian.png"; im.save(p); return p

def france_tri_figure():
    folder=ROOT/"data/raw/tri_2020_sig_di"
    layers=[("Débordement de cours d’eau","n_inondable_01_01for_s.shp","#1F77B4"),
            ("Ruissellement","n_inondable_02_01for_s.shp","#FF9F1C"),
            ("Submersion marine","n_inondable_03_01for_s.shp","#6A4C93")]
    bounds=(-5.8,41.0,10.2,51.5); box=(100,125,1120,805)
    im=Image.new("RGB",(1600,900),"white"); d=ImageDraw.Draw(im)
    d.text((70,35),"TRI France 2020 — scénario fréquent (« 01For »)",fill="#1F4E78",font=font(38,True))
    d.rectangle(box,fill="#F7FAFC",outline="#B4C6D7",width=3)
    for label,name,color in layers:
        draw_shapes(d,folder/name,bounds,box,outline=color,fill=None,max_shapes=11000,width=1)
    y=210
    for label,name,color in layers:
        d.rectangle((1190,y-13,1230,y+13),fill=color)
        d.text((1250,y),label,fill="#333333",font=font(23),anchor="lm"); y+=70
    d.text((1190,465),"Couche de territoire : n_tri_s\n131 entités au niveau national",fill="#555555",font=font(22),spacing=8)
    add_wrapped(d,(1190,590),"L’aperçu est limité à la France métropolitaine ; les fichiers nationaux comprennent également les territoires ultramarins.",28,"#777777",font(20))
    p=FIG/"figure_03_tri_france.png"; im.save(p); return p

def italy_hazard_figure():
    folder=ROOT/"data/raw/Mosaicatura_ISPRA_2020_aree_pericolosita_idraulica"
    layers=[("P1 — faible probabilité","LPH_Mosaicatura_ISPRA_2020_pericolosita_idraulica_bassa.shp","#8EC5FF"),
            ("P2 — probabilité moyenne","MPH_Mosaicatura_ISPRA_2020_pericolosita_idraulica_media.shp","#2E75B6"),
            ("P3 — probabilité élevée","HPH_Mosaicatura_ISPRA_2020_pericolosita_idraulica_elevata.shp","#C0504D")]
    bounds=(4050000,1515000,5060000,2660000); box=(100,125,1120,805)
    im=Image.new("RGB",(1600,900),"white"); d=ImageDraw.Draw(im)
    d.text((70,35),"Italie — mosaïque ISPRA 2020 de la dangerosité hydraulique",fill="#1F4E78",font=font(38,True))
    d.rectangle(box,fill="#F7FAFC",outline="#B4C6D7",width=3)
    for label,name,color in layers:
        draw_shapes(d,folder/name,bounds,box,outline=color,fill=None,max_shapes=1500,width=1)
    y=220
    for label,name,color in layers:
        d.rectangle((1190,y-13,1230,y+13),fill=color)
        d.text((1250,y),label,fill="#333333",font=font(23),anchor="lm"); y+=70
    add_wrapped(d,(1190,500),"Dans le workflow Italie, ces polygones jouent le rôle de filtre spatial complémentaire aux évènements historiques HANZE.",26,"#555555",font(22))
    p=FIG/"figure_04_tri_italie.png"; im.save(p); return p

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"),fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); mar=tcPr.first_child_found_in("w:tcMar")
    if mar is None: mar=OxmlElement("w:tcMar"); tcPr.append(mar)
    for m,v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node=mar.find(qn("w:"+m))
        if node is None: node=OxmlElement("w:"+m); mar.append(node)
        node.set(qn("w:w"),str(v)); node.set(qn("w:type"),"dxa")

def keep_with_next(p): p.paragraph_format.keep_with_next=True
def add_caption(doc,text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(10)
    r=p.add_run(text); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(89,89,89)
    return p
def add_figure(doc,path,caption,width=6.35):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True
    p.add_run().add_picture(str(path),width=Inches(width)); add_caption(doc,caption)
def add_callout(doc,label,text,fill="DDEBF7"):
    table=doc.add_table(rows=1,cols=1); table.alignment=WD_TABLE_ALIGNMENT.CENTER
    cell=table.cell(0,0); set_cell_shading(cell,fill); set_cell_margins(cell,120,180,120,180)
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(label+" "); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
def add_bullets(doc,items):
    for item in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(4)
        p.add_run(item)
def add_table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; set_cell_shading(c,"E8EEF5"); set_cell_margins(c)
        c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.bold=True; r.font.size=Pt(9)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=str(val); set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for r in cells[i].paragraphs[0].runs: r.font.size=Pt(9)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

def page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run=paragraph.add_run("Page ")
    fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE")
    run._r.addnext(fld)

def build():
    figs=[flow_figure(),riparian_figure(),france_tri_figure(),italy_hazard_figure()]
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Inches(.75); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.85); sec.right_margin=Inches(.85)
    styles=doc.styles
    normal=styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11)
    normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.10
    for name,size,color,bef,aft in [("Heading 1",16,MID,16,8),("Heading 2",13,MID,12,6),("Heading 3",12,BLUE,8,4)]:
        s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
        s.paragraph_format.space_before=Pt(bef); s.paragraph_format.space_after=Pt(aft); s.paragraph_format.keep_with_next=True
    footer=sec.footer.paragraphs[0]; footer.text="Rapport d’alternance — Présentation des données"; footer.runs[0].font.size=Pt(8); footer.runs[0].font.color.rgb=RGBColor(120,120,120); page_number(footer)

    # Couverture
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(30)
    r=p.add_run("RAPPORT D’ALTERNANCE"); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(TEAL)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(10)
    r=p.add_run("Présentation et utilisation des données d’inondation"); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=RGBColor.from_string(BLUE)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(24)
    r=p.add_run("Zones ripariennes Copernicus, base GASPAR et zonages TRI en France et en Italie"); r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(GRAY)
    add_callout(doc,"Objet du document —","Fournir un texte directement réutilisable dans la partie « Présentation des données » du rapport, en précisant la nature, le format, les attributs, la logique d’utilisation et les limites de chaque source.","E2F0D9")
    add_figure(doc,figs[0],"Figure 1 — Complémentarité des sources mobilisées dans l’analyse.",6.4)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Version du 24 juillet 2026").italic=True
    doc.add_page_break()

    doc.add_heading("1. Positionnement général des sources",level=1)
    doc.add_paragraph("L’étude combine des données de natures différentes. Leur complémentarité est essentielle : aucune source ne suffit, à elle seule, à démontrer qu’un actif précis a effectivement été inondé à une date donnée. GASPAR renseigne l’existence d’un évènement reconnu à l’échelle communale ; les couches TRI ou ISPRA décrivent des emprises exposées selon des scénarios ; les zones ripariennes caractérisent les corridors proches des cours d’eau ; enfin, les données raster JRC apportent une confirmation événementielle localisée lorsque l’évènement est observé.")
    add_table(doc,["Source","Nature","Unité spatiale","Information principale","Rôle dans le workflow"],[
        ["Riparian Zones","Vecteur — polygones","Corridors fluviaux européens","Occupation/couverture du sol en 2018","Filtre spatial complémentaire"],
        ["GASPAR / CATNAT","Table attributaire","Commune + période","Reconnaissance administrative d’un évènement","Sélection temporelle et communale"],
        ["TRI France 2020","Vecteur — polygones","Territoires et surfaces inondables","Scénarios fréquent, moyen et exceptionnel","Validation spatiale en France"],
        ["ISPRA Italie 2020","Vecteur — polygones","Mosaïque nationale","Danger hydraulique P1, P2 et P3","Validation spatiale en Italie"],
    ],[1.15,1.05,1.15,1.65,1.5])
    add_callout(doc,"Point de vocabulaire —","Dans la suite, l’expression « TRI Italie » est utilisée de manière fonctionnelle pour faciliter la comparaison. Le jeu italien réellement présent dans le projet est la mosaïque ISPRA de danger hydraulique ; sa nomenclature et sa construction ne sont pas identiques à celles du TRI français.","FFF2CC")

    doc.add_heading("2. Données Riparian Zones 2018",level=1)
    doc.add_heading("2.1 Définition et finalité",level=2)
    doc.add_paragraph("Le produit Riparian Zones du Copernicus Land Monitoring Service cartographie l’occupation et la couverture du sol dans les zones riveraines et les plaines alluviales européennes. Une zone riparienne correspond à l’interface entre un milieu terrestre et un système d’eau douce. Elle est fortement influencée par la proximité du cours d’eau, les débordements, l’humidité des sols et la dynamique de la végétation.")
    doc.add_paragraph("Le produit n’est donc pas une carte d’inondation historique. Il décrit un contexte territorial favorable aux interactions entre le cours d’eau et les surfaces voisines. Dans l’étude, il intervient comme indicateur spatial secondaire : lorsqu’un évènement GASPAR est compatible avec la commune et la période mais que le point est situé hors des polygones TRI retenus, l’intersection avec une zone riparienne peut soutenir la plausibilité d’une exposition fluviale.")
    add_figure(doc,figs[1],"Figure 2 — Aperçu des polygones Riparian Zones 2018 chargés pour la France et l’Italie.",6.4)
    doc.add_heading("2.2 Format et structure des fichiers",level=2)
    doc.add_paragraph("Les données sont fournies au format ESRI Shapefile. Une couche exploitable repose au minimum sur le fichier géométrique .shp, l’index .shx, la table attributaire .dbf et le fichier de projection .prj. Les couches du projet sont des polygones projetés en ETRS89 / LAEA Europe (EPSG:3035), un système métrique adapté aux analyses harmonisées à l’échelle européenne.")
    add_table(doc,["Champ observé","Signification"],[
        ["UID","Identifiant de l’objet géographique."],["DU_ID","Identifiant de l’unité de livraison Copernicus."],
        ["CODE_1_18 à CODE_4_18","Niveaux hiérarchiques de la nomenclature d’occupation/couverture du sol pour 2018."],
        ["UA_18","Lien éventuel avec une classe Urban Atlas."],["NODATA_18","Indicateur d’absence ou d’indisponibilité de l’information."],
        ["COMMENT_18","Commentaire de production ou exception cartographique."],["AREA_HA","Surface du polygone en hectares."],
    ],[1.7,4.8])
    doc.add_paragraph("Dans le dépôt, six unités de livraison sont utilisées sous France_Riparian, soit 1 480 130 polygones, et trois unités sous Italy_Riparian, soit 755 779 polygones. Ces chiffres décrivent les fichiers présents dans le projet ; ils ne doivent pas être interprétés comme un nombre de zones inondables indépendantes, car les objets représentent des classes d’occupation du sol découpées par unité de livraison.")
    doc.add_heading("2.3 Utilisation et limites",level=2)
    add_bullets(doc,[
        "Test effectué : intersection géométrique entre le point étudié et les polygones ripariens.",
        "Avantage : source européenne homogène, détaillée et directement exploitable dans un SIG.",
        "Limite : présence dans un corridor riparien ≠ preuve d’une inondation observée.",
        "Limite : le millésime 2018 décrit l’occupation du sol à une date de référence et peut ne pas refléter des transformations ultérieures.",
        "Précaution : conserver l’identifiant de l’unité de livraison et le système de coordonnées pour assurer la traçabilité.",
    ])

    doc.add_heading("3. Base GASPAR et arrêtés CATNAT",level=1)
    doc.add_heading("3.1 Nature administrative de la donnée",level=2)
    doc.add_paragraph("GASPAR, pour « Gestion assistée des procédures administratives relatives aux risques naturels », est le système national français de référence pour plusieurs informations de prévention et de réglementation : plans de prévention des risques, atlas des zones inondables, documents d’information préventive et procédures de reconnaissance de l’état de catastrophe naturelle. Dans ce travail, l’exploitation porte sur la table des arrêtés CATNAT.")
    doc.add_paragraph("Une ligne associe un identifiant d’arrêté, une commune, un type de risque et une période d’évènement, complétés par les dates administratives de publication. La donnée est tabulaire : elle ne contient pas la géométrie exacte de la surface inondée. Le rattachement spatial se fait donc par le code commune, puis par croisement avec des polygones TRI ou ripariens.")
    add_table(doc,["Champ","Contenu"],[
        ["cod_nat_catnat","Identifiant national de l’arrêté CATNAT."],["cod_commune","Code INSEE de la commune concernée."],
        ["lib_commune","Nom de la commune."],["num_risque_jo / lib_risque_jo","Code et libellé du phénomène reconnu."],
        ["dat_deb / dat_fin","Début et fin de l’évènement."],["dat_pub_arrete / dat_pub_jo","Dates de publication de l’arrêté et au Journal officiel."],
        ["dat_maj","Date de mise à jour de l’enregistrement."],
    ],[2.15,4.35])
    add_callout(doc,"Lecture correcte —","Un enregistrement GASPAR signifie qu’un phénomène a fait l’objet d’une reconnaissance administrative pour une commune et une période. Il ne signifie pas que toute la commune a été inondée, ni que la profondeur d’eau est connue.","FCE4D6")
    doc.add_paragraph("Le fichier brut présent dans le projet contient 260 799 enregistrements et couvre 34 704 codes communaux. Parmi eux, 148 245 lignes ont un libellé contenant « Inondations » ; la période observée s’étend de 1982 à 2025. Ces statistiques sont propres à la copie locale analysée et pourront évoluer lors d’une actualisation de la base.")
    doc.add_heading("3.2 Logique de traitement",level=2)
    add_bullets(doc,[
        "Filtrer les phénomènes liés aux inondations et aux coulées de boue.",
        "Normaliser les codes INSEE et conserver l’historique des codes communaux lorsque nécessaire.",
        "Vérifier le chevauchement entre la période de l’évènement et la fenêtre temporelle propre à chaque exposition.",
        "Associer le point à sa commune, puis appliquer le contrôle spatial TRI / zone riparienne.",
        "Conserver les dates et l’identifiant de l’arrêté pour rendre chaque résultat explicable.",
    ])
    doc.add_heading("3.3 Forces et limites",level=2)
    doc.add_paragraph("GASPAR offre une profondeur historique et une bonne traçabilité administrative. En revanche, sa granularité communale est trop large pour conclure directement sur un actif géolocalisé. Des changements de codes INSEE, des communes fusionnées, des libellés hétérogènes et un décalage entre date de l’évènement et date de publication doivent être pris en compte.")

    doc.add_heading("4. TRI France — rapportage 2020",level=1)
    doc.add_heading("4.1 Définition",level=2)
    doc.add_paragraph("Les Territoires à risque important d’inondation (TRI) sont identifiés dans le cadre de la directive européenne 2007/60/CE. Ils concentrent des enjeux humains, économiques ou patrimoniaux majeurs susceptibles d’être touchés. La livraison nationale utilisée ici correspond au rapportage 2020, version 2, distribué au format Shapefile selon le standard Directive Inondation.")
    doc.add_paragraph("La cartographie TRI distingue généralement trois niveaux d’évènement : fréquent, moyen et exceptionnel. Le workflow simplifié du projet retient les surfaces du scénario fréquent, repérées par le suffixe 01For, ainsi que la limite globale des territoires TRI.")
    add_figure(doc,figs[2],"Figure 3 — Surfaces inondables du scénario fréquent utilisées dans le workflow France (aperçu métropolitain).",6.4)
    doc.add_heading("4.2 Couches effectivement mobilisées",level=2)
    add_table(doc,["Fichier","Type d’inondation","Nombre d’entités","Rôle"],[
        ["n_inondable_01_01for_s","Débordement de cours d’eau","78 913","Polygones de scénario fréquent."],
        ["n_inondable_02_01for_s","Ruissellement","409","Polygones de scénario fréquent."],
        ["n_inondable_03_01for_s","Submersion marine","28 064","Polygones de scénario fréquent."],
        ["n_tri_s","Limite des TRI","131","Distinguer l’intérieur d’un TRI de l’extérieur."],
    ],[2.0,1.55,1.0,1.95])
    doc.add_paragraph("Les champs des surfaces inondables comprennent notamment l’identifiant de la surface, le type d’inondation, le scénario, le cours d’eau et l’identifiant du TRI. Les fichiers sont en coordonnées géographiques WGS 84. Les limites n_tri_s contiennent notamment le nom et l’identifiant du territoire ainsi que des informations de population.")
    doc.add_heading("4.3 Règle de décision utilisée en France",level=2)
    add_bullets(doc,[
        "Si le point intersecte une surface 01For, le signal GASPAR compatible est conservé.",
        "Si le point est à l’intérieur de la limite d’un TRI mais hors des surfaces 01For, le signal GASPAR n’est pas retenu par la règle simplifiée.",
        "Si le point est hors des limites TRI, une intersection avec une zone riparienne peut servir de solution de repli.",
    ])
    add_callout(doc,"Limite méthodologique —","Cette règle privilégie le scénario fréquent et n’exploite pas les scénarios moyen ou exceptionnel. Elle est adaptée à un filtrage prudent, mais elle peut écarter des points exposés uniquement lors d’évènements rares.","FFF2CC")

    doc.add_heading("5. Italie — mosaïque ISPRA de danger hydraulique",level=1)
    doc.add_heading("5.1 Équivalent fonctionnel du TRI français",level=2)
    doc.add_paragraph("Pour l’Italie, le dépôt contient la mosaïque nationale ISPRA 2020 des zones de danger hydraulique. Elle résulte de l’harmonisation des périmètres produits par les autorités de bassin de district, les régions et les provinces autonomes dans le cadre du décret législatif 49/2010, qui transpose la directive européenne sur les inondations.")
    doc.add_paragraph("Trois couches polygonales sont disponibles en ETRS89 / LAEA Europe (EPSG:3035). Leur table attributaire est volontairement simple et contient principalement le libellé du scénario.")
    add_figure(doc,figs[3],"Figure 4 — Mosaïque ISPRA 2020 des trois niveaux de danger hydraulique en Italie.",5.35)
    add_table(doc,["Couche","Interprétation","Ordre de grandeur officiel 2020"],[
        ["P3 / HPH","Danger élevé ; évènements fréquents, période de retour d’environ 20 à 50 ans.","16 224 km², soit 5,4 % du territoire."],
        ["P2 / MPH","Danger moyen ; évènements moins fréquents, période de retour d’environ 100 à 200 ans.","30 194 km², soit 10 % du territoire."],
        ["P1 / LPH","Danger faible ; faible probabilité ou scénario extrême.","42 376 km², soit 14 % du territoire."],
    ],[1.25,3.35,1.9])
    doc.add_heading("5.2 Utilisation dans le workflow Italie",level=2)
    doc.add_paragraph("Dans la branche italienne, les évènements historiques sont fournis par HANZE plutôt que par GASPAR. Après filtrage temporel et rapprochement administratif, le point est classé par intersection avec les polygones ISPRA. Le résultat permet de préciser si l’actif se trouve dans une zone de danger élevé, moyen ou faible. Contrairement au workflow France, la version actuelle du script Italie n’emploie pas les zones ripariennes comme solution de repli.")
    doc.add_heading("5.3 Limites de comparaison France–Italie",level=2)
    add_bullets(doc,[
        "Les deux pays appliquent la directive Inondations, mais les produits nationaux, nomenclatures et processus d’harmonisation diffèrent.",
        "Les scénarios français « fréquent / moyen / exceptionnel » sont proches conceptuellement de P3 / P2 / P1, sans constituer une équivalence parfaite objet par objet.",
        "Les polygones expriment une dangerosité ou une emprise potentielle, pas la survenue certaine d’un évènement à la date étudiée.",
        "La précision varie selon les données sources régionales et le réseau hydrographique effectivement modélisé.",
    ])

    doc.add_heading("6. Synthèse méthodologique et recommandations",level=1)
    doc.add_heading("6.1 Chaîne de décision recommandée",level=2)
    add_bullets(doc,[
        "1. Localiser le point et contrôler ses coordonnées.",
        "2. Identifier l’unité administrative correspondante.",
        "3. Sélectionner les évènements compatibles avec la période d’étude.",
        "4. Tester l’intersection avec les zonages nationaux de danger.",
        "5. En France, utiliser la zone riparienne uniquement selon la règle de repli documentée.",
        "6. Lorsque disponible, confirmer le signal par une donnée événementielle raster JRC.",
        "7. Conserver pour chaque résultat la source, le millésime, la couche, le scénario et la règle appliquée.",
    ])
    doc.add_heading("6.2 Formulation proposée pour le rapport",level=2)
    add_callout(doc,"Texte prêt à intégrer —","L’analyse repose sur une combinaison de données administratives, cartographiques et environnementales. La base GASPAR permet d’identifier les évènements reconnus à l’échelle communale et de contrôler leur cohérence temporelle. Les zonages TRI en France et la mosaïque ISPRA en Italie apportent une information spatiale sur les emprises potentiellement exposées selon différents scénarios de fréquence. Les zones ripariennes Copernicus décrivent, quant à elles, l’occupation du sol des corridors fluviaux et sont utilisées comme indicateur complémentaire, sans être assimilées à une preuve d’inondation. Cette approche croisée réduit le risque d’interpréter un évènement communal comme une exposition certaine de chaque actif et améliore la traçabilité de la décision.","E2F0D9")
    doc.add_heading("6.3 Points de vigilance",level=2)
    add_bullets(doc,[
        "Ne pas confondre reconnaissance CATNAT, danger potentiel et observation effective.",
        "Harmoniser les systèmes de coordonnées avant toute intersection spatiale.",
        "Éviter de comparer directement le nombre de polygones entre pays : le découpage cartographique est différent.",
        "Documenter les millésimes : Riparian 2018, TRI France 2020, ISPRA Italie 2020, copie locale GASPAR mise à jour jusqu’en 2025.",
        "Présenter le résultat comme un outil de screening et non comme une expertise hydraulique ou juridique.",
    ])

    doc.add_heading("7. Sources et références",level=1)
    refs=[
        "Copernicus Land Monitoring Service, « Riparian Zones » : https://land.copernicus.eu/en/products/riparian-zones",
        "Copernicus, Product User Manual — Riparian Zones LC/LU 2012 et 2018 : https://land.copernicus.eu/en/technical-library/product-user-manual-riparian-zones-land-cover-land-use-2012-and-2018-and-land-cover-land-use-change-2012-2018",
        "Géorisques, glossaire GASPAR : https://www.georisques.gouv.fr/glossaire/gaspar",
        "Géorisques, Zonages Inondation — Rapportage 2020 : https://www.georisques.gouv.fr/donnees/bases-de-donnees/zonages-inondation-rapportage-2020",
        "Géorisques, dossier expert sur les inondations : https://www.georisques.gouv.fr/consulter-les-dossiers-thematiques/dossier-expert-sur-les-inondations",
        "ISPRA, « Le alluvioni » : https://www.isprambiente.gov.it/it/attivita/suolo-e-territorio/dissesto-idrogeologico/le-alluvioni",
        "ISPRA, Rapport 356/2021 : https://www.isprambiente.gov.it/it/pubblicazioni/rapporti/dissesto-idrogeologico-in-italia-pericolosita-e-indicatori-di-rischio-edizione-2021",
        "Références internes au projet : docs/tri_2020_sig_di_reference.md, docs/gaspar_all_dates_workflow.md et docs/check_italy_points_against_jrc_hanze_guide.md.",
    ]
    add_bullets(doc,refs)
    doc.add_paragraph("Les statistiques de fichiers et les noms de champs indiqués dans ce document ont été vérifiés directement sur les données présentes dans le dépôt au 24 juillet 2026.")

    path=OUT/"Presentation_donnees_Riparian_GASPAR_TRI_FR_IT_final.docx"
    doc.save(path)
    return path

if __name__=="__main__":
    print(build())
