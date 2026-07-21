from __future__ import annotations

import argparse
import datetime as dt
import tempfile
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = Path(__file__).resolve().parent
DEFAULT_OUTPUT = REPORT_DIR / "rapport_alternance_data_collection.docx"
DEFAULT_FIGURES_DIR = REPORT_DIR / "generated_figures"

GREEN = RGBColor(0, 102, 68)
DARK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_FILL = "EEF3F8"
LIGHT_GREEN = "E8F3EE"
LIGHT_NOTE = "F7F7F7"

MATCH_RATE_ROWS = [
    ("Commune 7 jours", 22.9, 14.2),
    ("Commune 30 jours", 27.8, 16.7),
    ("Département 7 jours", 43.4, 42.5),
    ("Département 30 jours", 59.4, 53.8),
]

SOURCE_ROWS = [
    (
        "JRC / Copernicus",
        "Raster 20 m, événement",
        "Identifier les pixels inondés, tabulariser les événements et calculer des profondeurs.",
        "Gestion du CRS projeté, volumes de données élevés et filtrage strict des fichiers officiels.",
    ),
    (
        "GASPAR / CatNat",
        "Commune, période administrative",
        "Comparer les reconnaissances administratives aux événements JRC et enrichir le screening France.",
        "Historique des codes communes, logique CatNat différente du footprint satellitaire.",
    ),
    (
        "HANZE",
        "Événement historique, NUTS3",
        "Compléter la profondeur historique lorsque l'événement n'est pas directement confirmé par JRC.",
        "Normalisation des dates, éclatement des régions multi-NUTS3 et cohérence temporelle.",
    ),
    (
        "TRI / couches ripariennes",
        "Polygones de risque",
        "Appliquer une règle spatiale de filtrage pour la branche France GASPAR/HANZE.",
        "Sélection du sous-ensemble utile et documentation de la règle métier.",
    ),
    (
        "Eurostat LAU / NUTS",
        "Polygones administratifs Europe",
        "Construire la table canonique des événements par unité locale puis les agrégations NUTS.",
        "Cohérence pan-européenne et respect des projections géographiques.",
    ),
    (
        "AdminExpress / INSEE",
        "Commune courante + historique",
        "Réconcilier les codes français actuels et historiques pour les jointures JRC/GASPAR.",
        "Mouvements de communes, homonymes, cas particuliers comme la Corse.",
    ),
    (
        "ISPRA / Italie",
        "Polygones d'aléa hydraulique",
        "Étendre la logique de screening à l'Italie et documenter un second écosystème national.",
        "Harmonisation de standards nationaux différents avec le socle JRC/HANZE.",
    ),
]

MISSION_ROWS = [
    (
        "Construire un socle géodata",
        "Passer de sources brutes hétérogènes à des tables propres, historisées et reliables.",
        "Pipelines JRC Europe, harmonisation France LAU vers INSEE, extensions Italie.",
    ),
    (
        "Industrialiser le screening",
        "Rendre reproductible la vérification point par point pour plusieurs workflows métier.",
        "Scripts France T20, France collaterals, Italie T20, Italie collaterals, build FLOOD_LGD.",
    ),
    (
        "Produire des outils d'analyse",
        "Faciliter l'exploration, le contrôle visuel et la discussion métier autour des écarts observés.",
        "Applications Streamlit, audits JRC vs GASPAR, notes méthodologiques et supports de restitution.",
    ),
    (
        "Préparer les analyses LGD",
        "Livrer une base exploitable pour des analyses ultérieures de défaut et de perte.",
        "Exports consolidés, dictionnaires de colonnes, guides de lecture et points d'extension.",
    ),
]

WORKFLOW_ROWS = [
    (
        "France T20",
        "check_points_against_jrc_floods.py",
        "build_flood_lgd_exports.py",
        "JRC + GASPAR + HANZE",
    ),
    (
        "France collaterals",
        "check_points_against_jrc_floods_collaterals.py",
        "build_flood_lgd_exports_collaterals.py",
        "JRC + GASPAR + HANZE",
    ),
    (
        "Italie T20",
        "check_italy_points_against_jrc_hanze.py",
        "build_flood_lgd_exports_italy.py",
        "JRC + HANZE",
    ),
    (
        "Italie collaterals",
        "check_italy_points_against_jrc_hanze_collaterals.py",
        "build_flood_lgd_exports_collaterals_italy.py",
        "JRC + HANZE",
    ),
]

TEST_ROWS = [
    (
        "tests/test_check_points_against_jrc_floods.py",
        "Contrôle la logique de candidats JRC, les fenêtres temporelles et les décisions de hit raster.",
    ),
    (
        "tests/test_check_points_against_gaspar_floods.py",
        "Sécurise la branche GASPAR et les règles de filtrage propres à la France.",
    ),
    (
        "tests/test_build_flood_lgd_exports.py",
        "Valide la consolidation finale, les defaults des wrappers et plusieurs cas collaterals / Italie.",
    ),
    (
        "tests/test_france_commune_activity.py",
        "Vérifie les fonctions de réconciliation communale et la logique partagée avec l'application Streamlit.",
    ),
    (
        "tests/test_flood_preview.py",
        "Stabilise les comportements du moteur de preview raster utilisé dans le notebook et l'application.",
    ),
]

SCRIPT_MAP_ROWS = [
    (
        "src/granular_tabularization.py",
        "Tabularisation JRC Europe",
        "Produit la table canonique par LAU et les agrégations NUTS.",
    ),
    (
        "src/france_lau_to_insee.py",
        "Harmonisation France",
        "Rapproche LAU, communes INSEE courantes et historique des mouvements communaux.",
    ),
    (
        "src/check_points_against_jrc_floods.py",
        "Screening France T20",
        "Confirme les événements JRC au voisinage d'un point et prépare les branches GASPAR/HANZE.",
    ),
    (
        "src/check_points_against_jrc_floods_collaterals.py",
        "Screening France collaterals",
        "Adapte la logique point par point au cas des sûretés avec identifiants répétés.",
    ),
    (
        "src/check_italy_points_against_jrc_hanze.py",
        "Screening Italie T20",
        "Combine la confirmation JRC et la logique HANZE / TRI côté Italie.",
    ),
    (
        "src/build_flood_lgd_exports.py",
        "Consolidation finale",
        "Regroupe les sources, fusionne les épisodes proches et écrit les colonnes FLOOD_LGD.",
    ),
    (
        "src/gaspar_jrc_france_map_app.py",
        "Application France",
        "Visualise l'activité communale GASPAR vs JRC sur une période donnée.",
    ),
    (
        "src/app.py + src/flood_preview.py",
        "Exploration raster",
        "Permettent la prévisualisation rapide de TIFF lourds avec plusieurs stratégies de rendu.",
    ),
]

COMPLETION_ROWS = [
    (
        "Résultats chiffrés portefeuille / LGD",
        "Les sorties confidentielles et les analyses métiers finales sont stockées sur un autre poste.",
        "generated_figures/lgd_portfolio_results.png",
    ),
    (
        "Mix des sources dans FLOOD_LGD",
        "Le rapport doit encore montrer la répartition réelle JRC / GASPAR / HANZE sur vos exports finaux.",
        "generated_figures/flood_lgd_source_mix_france_t20.png",
    ),
    (
        "Captures d'écran des applications",
        "Les captures définitives doivent refléter l'environnement de démonstration utilisé pour la soutenance.",
        "generated_figures/france_commune_app.png et generated_figures/raster_dashboard.png",
    ),
    (
        "Tableaux d'annexe orientés métier",
        "Les annexes finales pourront reprendre des extractions propriétaires de workbook ou de portefeuille.",
        "generated_figures/lgd_workbook_extract.png",
    ),
]

FIGURE_HINT_ROWS = [
    (
        "france_commune_app.png",
        "Capture d'écran finale de l'application France GASPAR vs JRC.",
        "Optionnel mais recommandé pour remplacer la figure de fallback.",
    ),
    (
        "raster_dashboard.png",
        "Capture d'écran finale du dashboard raster JRC.",
        "Optionnel mais désormais reconnu dans la section 4.1.",
    ),
    (
        "jrc_gaspar_comparison_snapshot.png",
        "Synthèse auto-générée depuis les coverage_overview du poste métier.",
        "Remplace le graphique intégré par défaut dans la section 4.3.",
    ),
    (
        "flood_lgd_source_mix_france_t20.png",
        "Répartition des sources retenues dans l'export France T20.",
        "Générée automatiquement si le fichier FLOOD_LGD est fourni.",
    ),
    (
        "lgd_portfolio_results.png",
        "Figure libre pour vos résultats finaux de portefeuille ou votre première lecture LGD.",
        "À produire manuellement ou à partir d'un notebook métier.",
    ),
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Génère un rapport d'alternance Word à partir du dépôt DataCollection, "
            "avec figures de fallback et emplacements réservés pour les sorties confidentielles."
        )
    )
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--figures-dir", type=Path, default=DEFAULT_FIGURES_DIR)
    parser.add_argument("--student-name", default="Juan David Alonso")
    parser.add_argument("--company", default="BNP Paribas")
    parser.add_argument(
        "--program",
        default="Master 2 MOSEF, Paris 1 Panthéon-Sorbonne",
    )
    parser.add_argument(
        "--period",
        default="[à renseigner]",
        help="Période d'alternance affichée sur la page de garde.",
    )
    parser.add_argument(
        "--mission-title",
        default=(
            "Construction d'un pipeline de données d'inondation pour le screening "
            "du risque physique et la préparation d'analyses LGD"
        ),
    )
    parser.add_argument("--tutor", default="[à renseigner]")
    parser.add_argument("--manager", default="[à renseigner]")
    return parser.parse_args()


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def get_font(size: int, bold: bool = False):
    candidates = [
        Path(r"C:\Windows\Fonts\segoeuib.ttf") if bold else Path(r"C:\Windows\Fonts\segoeui.ttf"),
        Path(r"C:\Windows\Fonts\arialbd.ttf") if bold else Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Mettez à jour le champ dans Word si le sommaire n'apparaît pas."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18

    title_style = document.styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(26)
    title_style.font.bold = True
    title_style.font.color.rgb = DARK
    title_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.12

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header_p.add_run("Rapport d'alternance | DataCollection")
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = MUTED

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run("Page ")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MUTED
    add_page_field(footer_p)


def style_table(table, widths: list[float], *, header_fill: str = LIGHT_FILL, font_size: int = 9) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)

    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_fill)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    set_repeat_table_header(table.rows[0])


def add_table(document: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float], *, header_fill: str = LIGHT_FILL, font_size: int = 9) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            row.cells[idx].text = value
    style_table(table, widths, header_fill=header_fill, font_size=font_size)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.08
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.08
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_paragraph(document: Document, text: str, *, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.italic = italic


def add_note_box(document: Document, title: str, text: str, *, fill: str = LIGHT_GREEN) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    cell.text = ""
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)

    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run(title)
    run_title.bold = True
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(10.5)

    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_after = Pt(0)
    p_body.paragraph_format.line_spacing = 1.06
    run_body = p_body.add_run(text)
    run_body.font.name = "Calibri"
    run_body.font.size = Pt(10)


def add_code_block(document: Document, code: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_NOTE)
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(code.splitlines()):
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        if idx < len(code.splitlines()) - 1:
            run.add_break()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if draw.textlength(candidate, font=font) <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def create_placeholder_image(path: Path, title: str, body: str) -> None:
    width, height = 1400, 820
    image = Image.new("RGB", (width, height), (250, 251, 252))
    draw = ImageDraw.Draw(image)

    draw.rounded_rectangle((50, 50, width - 50, height - 50), radius=24, outline=(182, 191, 200), width=3, fill=(255, 255, 255))
    draw.rectangle((50, 50, width - 50, 140), fill=(232, 243, 238))

    title_font = get_font(34, bold=True)
    body_font = get_font(24)
    hint_font = get_font(22, bold=True)

    draw.text((90, 82), title, fill=(15, 23, 42), font=title_font)

    body_lines = wrap_text(draw, body, body_font, width - 180)
    y = 220
    for line in body_lines:
        draw.text((90, y), line, fill=(71, 85, 105), font=body_font)
        y += 42

    hint = "Le builder remplacera automatiquement ce visuel dès qu'un fichier du même nom sera disponible."
    hint_lines = wrap_text(draw, hint, hint_font, width - 180)
    y += 36
    for line in hint_lines:
        draw.text((90, y), line, fill=(0, 102, 68), font=hint_font)
        y += 38

    image.save(path)


def create_pipeline_figure(path: Path) -> None:
    width, height = 1500, 620
    image = Image.new("RGB", (width, height), (248, 250, 252))
    draw = ImageDraw.Draw(image)
    title_font = get_font(34, bold=True)
    box_title = get_font(22, bold=True)
    box_body = get_font(18)

    draw.text((70, 30), "Chaîne de valeur du dépôt DataCollection", fill=(15, 23, 42), font=title_font)
    draw.text((70, 72), "Du raster brut à l'export FLOOD_LGD puis aux outils d'audit et de visualisation", fill=(71, 85, 105), font=box_body)

    boxes = [
        (
            (80, 160, 330, 380),
            "Sources brutes",
            "JRC, GASPAR,\nHANZE, TRI,\nLAU/NUTS,\nAdminExpress, ISPRA",
            (230, 244, 234),
        ),
        (
            (390, 160, 640, 380),
            "Normalisation",
            "Parsing, filtres\nofficiels, CRS,\nhistorique des codes,\nlookup France",
            (232, 239, 248),
        ),
        (
            (700, 160, 950, 380),
            "Screening",
            "Workflows France\net Italie,\nT20 et collaterals,\npoint vs area",
            (243, 238, 252),
        ),
        (
            (1010, 160, 1260, 380),
            "Consolidation",
            "Fusion des épisodes,\npriorités de source,\ncolonnes FLOOD_LGD,\nexports CSV / Excel",
            (255, 244, 229),
        ),
        (
            (1320, 160, 1470, 380),
            "Restitution",
            "Apps Streamlit,\naudits,\ndéploiement,\ndocumentation",
            (250, 233, 235),
        ),
    ]

    for x0, y0, x1, y1, title, body, fill in [
        (box[0][0], box[0][1], box[0][2], box[0][3], box[1], box[2], box[3]) for box in boxes
    ]:
        draw.rounded_rectangle((x0, y0, x1, y1), radius=22, fill=fill, outline=(174, 184, 194), width=3)
        draw.text((x0 + 18, y0 + 20), title, fill=(15, 23, 42), font=box_title)
        body_lines = body.splitlines()
        y = y0 + 72
        for line in body_lines:
            draw.text((x0 + 18, y), line, fill=(71, 85, 105), font=box_body)
            y += 34

    arrow_y = 270
    for start_x in [330, 640, 950, 1260]:
        draw.line((start_x + 15, arrow_y, start_x + 55, arrow_y), fill=(71, 85, 105), width=6)
        draw.polygon(
            [(start_x + 55, arrow_y), (start_x + 35, arrow_y - 12), (start_x + 35, arrow_y + 12)],
            fill=(71, 85, 105),
        )

    footer = "Cette structuration est ce qui rend possible, ensuite, l'analyse métier et la future lecture LGD."
    draw.text((70, 500), footer, fill=(0, 102, 68), font=box_title)
    image.save(path)


def create_match_rate_chart(path: Path) -> None:
    width, height = 1500, 840
    image = Image.new("RGB", (width, height), (248, 250, 252))
    draw = ImageDraw.Draw(image)

    title_font = get_font(34, bold=True)
    label_font = get_font(22)
    tick_font = get_font(18)
    legend_font = get_font(20)

    draw.text((90, 28), "Taux de match JRC vs GASPAR déjà documentés dans le dépôt", fill=(15, 23, 42), font=title_font)
    draw.text((90, 68), "Lecture synthétique des audits France : commune vs département, fenêtre 7 jours vs 30 jours", fill=(71, 85, 105), font=label_font)

    chart_x0, chart_y0, chart_x1, chart_y1 = 130, 150, 1420, 690
    chart_w = chart_x1 - chart_x0
    chart_h = chart_y1 - chart_y0
    max_pct = 70.0

    for pct in range(0, 71, 10):
        y = chart_y1 - (pct / max_pct) * chart_h
        draw.line((chart_x0, y, chart_x1, y), fill=(220, 227, 234), width=1)
        draw.text((70, y - 10), f"{pct}%", fill=(100, 116, 139), font=tick_font)

    group_width = 230
    bar_width = 62
    gap_in_group = 34
    group_gap = 75
    start_x = chart_x0 + 40
    colors = {"JRC": (37, 99, 235), "GASPAR": (245, 118, 39)}

    for idx, (label, jrc_rate, gaspar_rate) in enumerate(MATCH_RATE_ROWS):
        x_base = start_x + idx * (group_width + group_gap)
        for offset, name, rate in [(0, "JRC", jrc_rate), (bar_width + gap_in_group, "GASPAR", gaspar_rate)]:
            x0 = x_base + offset
            x1 = x0 + bar_width
            y0 = chart_y1 - (rate / max_pct) * chart_h
            draw.rounded_rectangle((x0, y0, x1, chart_y1), radius=10, fill=colors[name])
            value = f"{rate:.1f}%"
            tw = draw.textlength(value, font=tick_font)
            draw.text((x0 + (bar_width - tw) / 2, y0 - 28), value, fill=(15, 23, 42), font=tick_font)

        label_lines = textwrap.wrap(label, width=16)
        y = chart_y1 + 16
        for line in label_lines:
            tw = draw.textlength(line, font=tick_font)
            draw.text((x_base + 20 + (bar_width + gap_in_group + bar_width - tw) / 2, y), line, fill=(15, 23, 42), font=tick_font)
            y += 24

    legend_x = 1030
    legend_y = 110
    draw.rounded_rectangle((legend_x, legend_y, legend_x + 320, legend_y + 90), radius=16, fill=(255, 255, 255), outline=(203, 213, 225))
    draw.rectangle((legend_x + 20, legend_y + 24, legend_x + 48, legend_y + 52), fill=colors["JRC"])
    draw.text((legend_x + 62, legend_y + 22), "JRC", fill=(15, 23, 42), font=legend_font)
    draw.rectangle((legend_x + 160, legend_y + 24, legend_x + 188, legend_y + 52), fill=colors["GASPAR"])
    draw.text((legend_x + 202, legend_y + 22), "GASPAR", fill=(15, 23, 42), font=legend_font)

    footer = "Source : docs/gaspar_jrc_match_audit_fr.md et docs/gaspar_jrc_horizon_audit_fr.md"
    draw.text((90, 770), footer, fill=(71, 85, 105), font=tick_font)
    image.save(path)


def insert_picture_with_caption(document: Document, image_path: Path, caption: str, *, width: float = 6.1) -> None:
    document.add_picture(str(image_path), width=Inches(width))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(3)
    caption_paragraph.paragraph_format.space_after = Pt(8)
    run = caption_paragraph.add_run(caption)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = MUTED


def resolve_or_placeholder(name: str, figures_dir: Path, temp_dir: Path, *, fallback: Path | None, title: str, missing_note: str) -> Path:
    preferred = figures_dir / name
    if preferred.exists():
        return preferred
    if fallback is not None and fallback.exists():
        return fallback
    placeholder = temp_dir / name
    create_placeholder_image(placeholder, title, missing_note)
    return placeholder


def add_cover_page(document: Document, args: argparse.Namespace) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Rapport d'alternance")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = DARK

    document.add_paragraph("")

    company = document.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run(args.company.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(29)
    run.bold = True
    run.font.color.rgb = GREEN

    mission = document.add_paragraph()
    mission.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mission.paragraph_format.space_before = Pt(12)
    run = mission.add_run(args.mission_title)
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.font.color.rgb = DARK

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Version structurée à partir du dépôt DataCollection, prête à intégrer les résultats et figures du poste métier."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = MUTED

    document.add_paragraph("")
    document.add_paragraph("")

    meta_table = document.add_table(rows=6, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_rows = [
        ("Auteur", args.student_name),
        ("Formation", args.program),
        ("Entreprise", args.company),
        ("Période", args.period),
        ("Tuteur / tutrice", args.tutor),
        ("Manager", args.manager),
    ]
    for row_cells, (label, value) in zip(meta_table.rows, meta_rows, strict=True):
        row_cells.cells[0].text = label
        row_cells.cells[1].text = value
    style_table(meta_table, [1.6, 4.9], header_fill=LIGHT_FILL, font_size=10)

    document.add_paragraph("")
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(10)
    run = note.add_run(
        f"Document généré le {dt.date.today().strftime('%d/%m/%Y')}."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    document.add_page_break()


def add_toc_page(document: Document) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run("Table des matières")

    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(4)
    run = note.add_run(
        "Le sommaire ci-dessous peut être mis à jour automatiquement dans Word si nécessaire."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    toc_paragraph = document.add_paragraph()
    add_toc_field(toc_paragraph)

    document.add_paragraph("")
    plan = document.add_paragraph()
    run = plan.add_run("Plan du rapport")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    add_bullets(
        document,
        [
            "Introduction",
            "1. Cadre de l'alternance et enjeux métier",
            "2. Construction du socle de données d'inondation",
            "3. Industrialisation des workflows de screening et de sortie FLOOD_LGD",
            "4. Outils d'analyse, d'audit et de restitution",
            "5. Apports pour les futures analyses LGD",
            "6. Bilan personnel et conclusion",
            "Annexes",
        ],
    )

    document.add_page_break()


def add_introduction(document: Document) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run("Introduction")

    add_paragraph(
        document,
        "L'alternance s'inscrit dans un contexte de montée en puissance du risque physique climatique dans les dispositifs d'analyse du risque de crédit. L'un des besoins les plus concrets consiste à transformer des sources externes très hétérogènes en objets suffisamment propres, documentés et reproductibles pour être mobilisés dans des workflows d'analyse puis, à terme, dans des lectures LGD.",
    )
    add_paragraph(
        document,
        "Le dépôt DataCollection représente le cœur technique de cette mission. Il regroupe la collecte de données, la normalisation géospatiale, le screening point par point, la consolidation vers des tables FLOOD_LGD, plusieurs applications Streamlit de visualisation ainsi qu'un ensemble important de guides, de tests et de notes d'audit.",
    )
    add_paragraph(
        document,
        "Le présent rapport est donc volontairement centré sur ce patrimoine concret. Les résultats confidentiels et certaines figures finales étant stockés sur un autre poste, une version complète et soutenable a été rédigée dès maintenant, mais structurée pour réintégrer facilement les derniers graphiques, captures d'écran et tableaux métier au moment de la finalisation.",
    )
    add_note_box(
        document,
        "Note méthodologique",
        "Tout ce qui est décrit dans ce document est adossé à des scripts, guides ou visuels présents dans DataCollection. Lorsque des résultats quantitatifs finaux ne sont pas disponibles sur ce poste, le rapport insère un emplacement réservé clair plutôt qu'une estimation arbitraire.",
    )


def add_section_one(document: Document) -> None:
    h1 = document.add_paragraph(style="Heading 1")
    h1.add_run("1. Cadre de l'alternance et enjeux métier")

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("1.1 Enjeux métier et positionnement de la mission")
    add_paragraph(
        document,
        "Le fil rouge de l'alternance a été de rapprocher deux mondes qui dialoguent rarement de façon fluide : d'un côté des sources géospatiales et historiques complexes, de l'autre des besoins métiers de screening du risque physique devant rester lisibles, auditables et réutilisables dans des analyses de portefeuille.",
    )
    add_paragraph(
        document,
        "Dans cette perspective, la mission n'a pas consisté uniquement à 'récupérer des données'. Elle a consisté à construire une chaîne robuste allant de la source brute jusqu'à un objet exploitable par des équipes risques : tables harmonisées, règles de rapprochement explicites, exports finaux, documentation de lecture et outils de visualisation pour vérifier ou discuter les écarts.",
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("1.2 Missions principales réellement couvertes dans DataCollection")
    add_table(
        document,
        ["Axe", "Objectif", "Livrables visibles dans le dépôt"],
        MISSION_ROWS,
        [1.45, 2.15, 2.9],
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("1.3 Valeur ajoutée de la démarche")
    add_bullets(
        document,
        [
            "Passer d'une logique exploratoire à une logique industrialisable, avec des scripts nommés, des options CLI et des comportements documentés.",
            "Séparer nettement les briques : collecte et normalisation, screening, consolidation finale, visualisation, audit.",
            "Documenter non seulement 'comment exécuter', mais aussi 'pourquoi la logique est défendable' sur le plan géospatial et métier.",
            "Préparer le terrain pour les analyses LGD futures sans bloquer la production actuelle sur l'absence temporaire des résultats confidentiels.",
        ],
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("1.4 Bilan personnel intermédiaire")
    add_paragraph(
        document,
        "Cette mission a permis de progresser simultanément sur trois dimensions. Sur le plan technique, elle a conduit à consolider des compétences en géodata Python, en structuration de pipelines et en restitution outillée. Sur le plan métier, elle a aussi permis de traduire des besoins de screening du risque physique en règles concrètes de données. Enfin, sur le plan de l'autonomie, elle a renforcé un raisonnement en termes de produit interne : un bon script n'est utile que s'il est compréhensible, vérifiable et transmissible.",
    )


def add_section_two(document: Document, pipeline_figure: Path, italy_figure: Path) -> None:
    h1 = document.add_paragraph(style="Heading 1")
    h1.add_run("2. Construction du socle de données d'inondation")

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("2.1 Inventaire des sources et complémentarité des jeux de données")
    add_paragraph(
        document,
        "Le dépôt n'est pas bâti autour d'une seule source. Il articule au contraire plusieurs familles de données, chacune couvrant une partie différente du problème : footprint satellitaire, reconnaissance administrative, historique d'événements, polygones de risque et référentiels administratifs. C'est précisément cette complémentarité qui rend le projet utile, mais aussi exigeant à maintenir.",
    )
    add_table(
        document,
        ["Source", "Maille / format", "Usage principal", "Difficulté principale"],
        SOURCE_ROWS,
        [1.25, 1.35, 2.15, 1.75],
        font_size=8.6,
    )
    insert_picture_with_caption(
        document,
        pipeline_figure,
        "Figure 1. Vue d'ensemble de la chaîne de valeur DataCollection, depuis les sources brutes jusqu'aux exports et outils de restitution.",
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("2.2 Tabularisation JRC à l'échelle européenne")
    add_paragraph(
        document,
        "La brique structurante du dépôt est la tabularisation des rasters JRC. Le script src/granular_tabularization.py parcourt les fichiers officiels, applique un filtrage strict des noms de fichiers, respecte le système de projection projeté des TIFF et produit une table canonique par LAU. Cette table devient ensuite la base des agrégations NUTS0 à NUTS3.",
    )
    add_bullets(
        document,
        [
            "Un fichier TIFF officiel est traité comme un événement en soi, ce qui évite de regrouper artificiellement des épisodes distincts.",
            "Le workflow rejette automatiquement les dérivés d'affichage non officiels pour éviter les doubles comptes.",
            "La sortie canonique par LAU rend ensuite les agrégations supérieures déterministes et comparables entre pays.",
            "Le respect du CRS natif des rasters est traité comme une contrainte scientifique et non comme un simple détail d'affichage.",
        ],
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("2.3 Harmonisation France : LAU, INSEE courant et historique communal")
    add_paragraph(
        document,
        "Pour la France, la valeur du projet tient aussi à la réconciliation entre nomenclatures administratives. Le script src/france_lau_to_insee.py transforme les sorties JRC Europe en tables France exploitables par commune actuelle, tout en conservant la possibilité de remonter des codes anciens vers les codes courants. Cette étape est indispensable pour rendre les rapprochements avec GASPAR défendables.",
    )
    add_paragraph(
        document,
        "Le même enjeu est au centre du moteur partagé src/france_commune_activity.py : résolution d'un code actuel direct quand c'est possible, sinon usage de l'historique INSEE, puis fallbacks prudents sur les noms uniques. Cette logique de réconciliation n'est pas un confort ; elle conditionne la qualité de la comparaison source par source et la lisibilité de l'application cartographique.",
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("2.4 Ouverture Italie et adaptation à un second écosystème national")
    add_paragraph(
        document,
        "Le dépôt ne se limite pas à la France. Les scripts check_italy_points_against_jrc_hanze.py et build_flood_lgd_exports_italy.py montrent une généralisation de la logique vers l'Italie, avec un couplage JRC + HANZE et l'intégration d'éléments ISPRA. Cette extension est importante car elle prouve que la démarche n'est pas monolithique : elle peut absorber un contexte national différent tout en gardant une structure commune.",
    )
    add_paragraph(
        document,
        "Du point de vue de l'alternance, cette extension italienne joue un double rôle. Elle montre d'abord la robustesse de l'architecture de code. Elle montre ensuite une montée en abstraction : au lieu d'avoir un pipeline unique, le dépôt porte désormais un cadre réutilisable que l'on peut spécialiser selon la source, le pays et le cas d'usage.",
    )
    insert_picture_with_caption(
        document,
        italy_figure,
        "Figure 2. Exemple d'extension du socle de données au cas italien avec des polygones d'aléa hydraulique ISPRA.",
    )


def add_section_three(document: Document, flood_mix_figure: Path) -> None:
    h1 = document.add_paragraph(style="Heading 1")
    h1.add_run("3. Industrialisation des workflows de screening et de sortie FLOOD_LGD")

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("3.1 Séparation explicite entre check et build")
    add_paragraph(
        document,
        "L'une des décisions les plus structurantes du dépôt est la séparation entre la phase de check et la phase de build. La phase de check conserve les lignes par source et par événement candidat ; la phase de build reconstruit ensuite une lecture consolidée au grain 'point x épisode d'inondation'. Cette séparation améliore à la fois la traçabilité, l'auditabilité et la capacité à corriger un seul maillon sans refaire toute la chaîne.",
    )
    add_table(
        document,
        ["Workflow", "Script de check", "Script de build", "Sources finales"],
        WORKFLOW_ROWS,
        [1.1, 2.0, 2.0, 1.4],
        font_size=8.8,
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("3.2 Screening France et Italie : logique point, zone et variantes collaterals")
    add_paragraph(
        document,
        "Pour la France T20, le script de check combine plusieurs logiques complémentaires : géocodage du point vers sa LAU, sélection des candidats JRC, confirmation raster dans un buffer proche et un buffer large, puis préparation des branches GASPAR et HANZE. Les workflows collaterals ajoutent une contrainte concrète de qualité des identifiants : certains labels métiers peuvent être répétés ou vides, d'où la création d'un point_id séquentiel propre au pipeline.",
    )
    add_paragraph(
        document,
        "Le cas italien reprend cette philosophie tout en l'adaptant aux sources disponibles. La continuité entre T20 et collaterals, France et Italie, montre qu'il ne s'agit pas d'un unique script ad hoc, mais d'un ensemble cohérent de wrappers spécialisés autour d'un cœur logique partageable.",
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("3.3 Consolidation finale : rôle de build_flood_lgd_exports.py")
    add_paragraph(
        document,
        "La consolidation finale ne fait pas de calcul spatial supplémentaire ; elle réouvre le workbook source, récupère les métadonnées métier nécessaires et fusionne les preuves JRC, GASPAR et HANZE dans une table commune. La règle d'agrégation par défaut est celle d'un regroupement des événements proches dans une fenêtre de 30 jours, avec une priorité de source JRC > GASPAR > HANZE pour les champs de synthèse.",
    )
    add_note_box(
        document,
        "Point clé de lecture",
        "Le grain final n'est pas 'une ligne par point'. Il est 'une ligne par point et par épisode consolidé'. Un même point peut donc réapparaître plusieurs fois si plusieurs épisodes distincts sont détectés.",
    )
    add_note_box(
        document,
        "Autre point clé",
        "JRC est la seule source qui peut alimenter directement les champs de profondeur. GASPAR et HANZE servent surtout à qualifier l'existence d'un épisode et à enrichir la robustesse de la détection.",
        fill=LIGHT_FILL,
    )
    insert_picture_with_caption(
        document,
        flood_mix_figure,
        "Figure 3. Emplacement réservé pour la répartition réelle des sources retenues dans l'export FLOOD_LGD France T20.",
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("3.4 Documentation, qualité et tests")
    add_paragraph(
        document,
        "La transmissibilité du dépôt a également été travaillée. Les guides markdown présents dans docs/ détaillent non seulement les commandes d'exécution, mais aussi le raisonnement derrière les choix de géotraitement, la logique des colonnes et les conventions de sortie. Cette documentation évite que le projet repose sur une connaissance implicite détenue par une seule personne.",
    )
    add_table(
        document,
        ["Module de test", "Rôle principal"],
        TEST_ROWS,
        [2.7, 3.8],
        font_size=8.9,
    )


def add_section_four(document: Document, raster_figure: Path, app_figure: Path, comparison_figure: Path) -> None:
    h1 = document.add_paragraph(style="Heading 1")
    h1.add_run("4. Outils d'analyse, d'audit et de restitution")

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("4.1 Prévisualisation raster et pédagogie géospatiale")
    add_paragraph(
        document,
        "Le dépôt contient un second apport important : des outils pour expliquer et vérifier ce que font réellement les données. Le duo src/app.py et src/flood_preview.py permet d'explorer des rasters lourds sans charger systématiquement l'intégralité des fichiers à pleine résolution. Cette brique rend la donnée plus accessible à des utilisateurs qui ne sont pas nécessairement spécialistes SIG.",
    )
    add_bullets(
        document,
        [
            "Sélection des rasters officiels par année et par filtre de nom.",
            "Découpage intelligent de la zone utile avant rendu détaillé.",
            "Plusieurs stratégies d'affichage selon la densité et la taille de l'événement.",
            "Clarification pédagogique du problème de projection entre CRS scientifique et web map.",
        ],
    )
    insert_picture_with_caption(
        document,
        raster_figure,
        "Figure 4. Emplacement prévu pour une capture du dashboard raster JRC ou d'un visuel équivalent d'exploration géospatiale.",
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("4.2 Application France GASPAR vs JRC")
    add_paragraph(
        document,
        "L'application France constitue l'un des livrables les plus visibles de l'alternance. Son intérêt n'est pas de comparer deux rasters mais de montrer, à l'échelle communale et sur une période choisie, quelles communes sont actives dans GASPAR, dans JRC, ou dans les deux. Elle sert à la fois d'outil d'analyse, de démonstration et de contrôle qualité.",
    )
    add_paragraph(
        document,
        "Le moteur partagé résout les codes communaux, applique les filtres temporels, agrège l'activité par commune courante et alimente la cartographie. Le fait d'avoir également préparé des options de déploiement (Oracle Always Free, Render et Streamlit Community Cloud) montre que l'objectif n'était pas seulement exploratoire : la diffusion et la réutilisation de l'outil ont elles aussi été travaillées.",
    )
    insert_picture_with_caption(
        document,
        app_figure,
        "Figure 5. Exemple de lecture communale des écarts GASPAR vs JRC sur Grand Est, T3 2021.",
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("4.3 Audits quantitatifs et vérifications manuelles")
    add_paragraph(
        document,
        "Les notes docs/gaspar_jrc_match_audit_fr.md et docs/gaspar_jrc_horizon_audit_fr.md documentent déjà des résultats chiffrés et des cas d'usage analysés manuellement. Ces livrables sont précieux car ils évitent de réduire l'écart entre sources à un simple bug technique : ils montrent au contraire qu'une partie du mismatch reflète de vraies différences de granularité, de temporalité et de support spatial.",
    )
    add_table(
        document,
        [
            "Fenêtre / niveau",
            "Taux JRC",
            "Taux GASPAR",
            "Lecture rapide",
        ],
        [
            (
                "Commune 7 jours",
                "22,9 %",
                "14,2 %",
                "Match très faible au niveau communal avec fenêtre stricte.",
            ),
            (
                "Commune 30 jours",
                "27,8 %",
                "16,7 %",
                "Gain modéré, mais le désaccord reste important.",
            ),
            (
                "Département 7 jours",
                "43,4 %",
                "42,5 %",
                "Le recouvrement remonte fortement quand on relâche la granularité spatiale.",
            ),
            (
                "Département 30 jours",
                "59,4 %",
                "53,8 %",
                "La lecture départementale montre qu'une partie du mismatch vient de la fragmentation communale.",
            ),
        ],
        [1.4, 1.0, 1.0, 3.1],
        font_size=8.8,
    )
    insert_picture_with_caption(
        document,
        comparison_figure,
        "Figure 6. Synthèse des taux de match JRC vs GASPAR ; le visuel peut être remplacé automatiquement par une version régénérée depuis l'autre PC.",
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("4.4 Documentation de déploiement et logique produit")
    add_paragraph(
        document,
        "Enfin, une véritable couche de 'produit interne' a aussi été produite : guides de déploiement, notes d'usage, dictionnaires de colonnes, supports de présentation et scripts de génération de rapports. Cette dimension est importante dans un contexte d'alternance, car elle transforme un travail de code en patrimoine opérationnel partageable par l'équipe après la mission.",
    )


def add_section_five(document: Document, proprietary_figure: Path) -> None:
    h1 = document.add_paragraph(style="Heading 1")
    h1.add_run("5. Apports pour les futures analyses LGD")

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("5.1 Ce qui est déjà exploitable dès maintenant")
    add_bullets(
        document,
        [
            "Une chaîne de données d'inondation reproductible, depuis les rasters et bases externes jusqu'à des tables harmonisées.",
            "Des workflows de screening point par point adaptés à plusieurs cas métier et plusieurs pays.",
            "Une logique explicite de consolidation vers des exports FLOOD_LGD lisibles et documentés.",
            "Des outils d'audit et de visualisation qui facilitent la validation manuelle et la communication des écarts.",
            "Un ensemble de guides et de tests qui réduit le risque de dépendance à une seule personne.",
        ],
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("5.2 Ce qui reste à compléter depuis l'autre PC")
    add_paragraph(
        document,
        "La version actuelle du rapport est volontairement complète sur la méthode et sur les livrables techniques. En revanche, certains résultats confidentiels doivent encore être insérés pour la version finale de soutenance ou de remise. Le tableau ci-dessous recense explicitement ces éléments au lieu de les laisser implicites.",
    )
    add_table(
        document,
        ["Élément à compléter", "Pourquoi il n'apparaît pas ici", "Nom de fichier attendu"],
        COMPLETION_ROWS,
        [1.7, 3.0, 1.8],
        font_size=8.7,
    )
    insert_picture_with_caption(
        document,
        proprietary_figure,
        "Figure 6. Emplacement réservé pour les résultats portefeuille / LGD à intégrer depuis le poste métier.",
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("5.3 Prochaines étapes logiques pour la suite de l'analyse")
    add_numbered(
        document,
        [
            "Régénérer les figures quantitatives à partir des exports finaux FLOOD_LGD et des sorties de comparaison stockées sur le poste métier.",
            "Croiser les variables de flood avec les données de portefeuille et les dates de défaut pour produire une première lecture LGD ou pré-LGD.",
            "Tester plusieurs fenêtres temporelles et plusieurs mailles d'agrégation afin de mesurer la sensibilité des conclusions aux choix méthodologiques.",
            "Sélectionner quelques cas de défaut ou d'exposition emblématiques pour illustrer qualitativement la chaîne de bout en bout dans le rapport final.",
            "Capitaliser sur les applications Streamlit pour préparer une démonstration orale claire lors de la soutenance.",
        ],
    )


def add_section_six(document: Document) -> None:
    h1 = document.add_paragraph(style="Heading 1")
    h1.add_run("6. Bilan personnel et conclusion")

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("6.1 Enseignements tirés de l'alternance")
    add_paragraph(
        document,
        "Au-delà des scripts eux-mêmes, cette alternance a montré l'importance de construire un objet technique crédible pour des utilisateurs non techniques. Elle a imposé un raisonnement simultané en géodata, en logique d'audit, en qualité de sortie et en lisibilité métier. Elle a aussi mis en évidence qu'un bon résultat en science des données ne tient pas seulement à un calcul ; il tient à la capacité à retracer la donnée, expliquer la méthode et assumer les limites du périmètre.",
    )
    add_paragraph(
        document,
        "L'alternance a également conduit à faire progresser la manière de structurer le travail : découper un problème en briques, nommer proprement les workflows, documenter les arbitrages et ajouter des tests lorsque la logique devient fragile. Cette discipline est particulièrement importante sur un sujet comme le risque physique, où l'on manipule des sources hétérogènes et des interprétations qui doivent rester auditables.",
    )

    h2 = document.add_paragraph(style="Heading 2")
    h2.add_run("6.2 Conclusion générale")
    add_paragraph(
        document,
        "À ce stade, l'alternance a déjà permis de transformer DataCollection en un socle beaucoup plus complet qu'un simple dossier de scripts. Le dépôt porte désormais une chaîne de traitement cohérente, multi-sources et multi-usages, accompagnée d'outils de visualisation, d'audit et de déploiement. Même si l'analyse LGD complète reste à finaliser avec les résultats confidentiels du poste métier, l'infrastructure nécessaire à cette suite est déjà largement en place.",
    )
    add_paragraph(
        document,
        "En ce sens, la valeur de la mission tient autant dans les résultats immédiats que dans la capacité offerte à l'équipe de prolonger l'analyse. Le présent rapport reflète précisément cette idée : il documente un travail déjà substantiel, tout en préparant proprement la dernière étape d'enrichissement chiffré.",
    )


def add_annexes(document: Document) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)

    h1 = document.add_paragraph(style="Heading 1")
    h1.add_run("Annexe A. Cartographie synthétique des scripts et livrables")
    add_table(
        document,
        ["Fichier", "Rôle", "Contribution au rapport"],
        SCRIPT_MAP_ROWS,
        [2.2, 1.5, 2.8],
        font_size=8.7,
    )

    h1 = document.add_paragraph(style="Heading 1")
    h1.add_run("Annexe B. Commandes à exécuter sur l'autre PC")
    add_paragraph(
        document,
        "Les commandes ci-dessous permettent de générer les figures finales puis de reconstruire automatiquement le rapport avec ces visuels. Elles utilisent les scripts ajoutés dans docs/alternance_report/.",
    )
    add_code_block(
        document,
        "\n".join(
            [
                "python docs\\alternance_report\\generate_figures_for_other_pc.py \\",
                '  --comparison-7d-dir "data\\processed\\jrc_gaspar_comparison_flexible_7d" \\',
                '  --comparison-30d-dir "data\\processed\\jrc_gaspar_comparison_flexible_30d" \\',
                '  --flood-lgd "france_t20=outputs\\flood_lgd_export\\T20_Anonymised_FLOOD_LGD.csv" \\',
                '  --flood-lgd "france_collateral=outputs\\flood_lgd_export\\my_collaterals_points_FLOOD_LGD.csv" \\',
                '  --flood-lgd "italy_t20=outputs\\flood_lgd_export\\T20_Anonymised_italy_FLOOD_LGD.csv" \\',
                '  --flood-lgd "italy_collateral=outputs\\flood_lgd_export\\my_italy_collaterals_points_FLOOD_LGD.csv" \\',
                '  --copy-figure "france_commune_app=D:\\captures\\france_commune_app.png" \\',
                '  --copy-figure "raster_dashboard=D:\\captures\\raster_dashboard.png" \\',
                '  --copy-figure "lgd_portfolio_results=D:\\captures\\lgd_portfolio_results.png" \\',
                '  --out-dir "docs\\alternance_report\\generated_figures"',
            ]
        ),
    )
    add_code_block(
        document,
        "\n".join(
            [
                "python docs\\alternance_report\\build_rapport_alternance.py \\",
                '  --figures-dir "docs\\alternance_report\\generated_figures" \\',
                '  --period "à confirmer" \\',
                '  --tutor "à confirmer" \\',
                '  --manager "à confirmer"',
            ]
        ),
    )

    h1 = document.add_paragraph(style="Heading 1")
    h1.add_run("Annexe C. Fichiers de figures reconnus automatiquement par le builder")
    add_table(
        document,
        ["Nom de fichier", "Usage dans le rapport", "Statut"],
        FIGURE_HINT_ROWS,
        [1.8, 3.6, 1.1],
        font_size=8.8,
    )


def build_document(args: argparse.Namespace) -> Path:
    ensure_dir(args.output.parent)
    ensure_dir(args.figures_dir)

    with tempfile.TemporaryDirectory() as tmp_dir_str:
        temp_dir = Path(tmp_dir_str)
        pipeline_figure = temp_dir / "pipeline_architecture.png"
        match_chart = temp_dir / "match_rates.png"
        create_pipeline_figure(pipeline_figure)
        create_match_rate_chart(match_chart)

        italy_fallback = PROJECT_ROOT / "docs" / "Screenshot 2026-06-15 133559LMH.png"
        france_app_fallback = PROJECT_ROOT / "docs" / "assets" / "gaspar_jrc_match_audit" / "grand_est_2021_q3.png"

        raster_figure = resolve_or_placeholder(
            "raster_dashboard.png",
            args.figures_dir,
            temp_dir,
            fallback=None,
            title="Capture du dashboard raster",
            missing_note=(
                "Ajoutez une capture d'écran finale du dashboard raster JRC dans "
                "generated_figures/raster_dashboard.png pour illustrer la logique de prévisualisation."
            ),
        )
        app_figure = resolve_or_placeholder(
            "france_commune_app.png",
            args.figures_dir,
            temp_dir,
            fallback=france_app_fallback,
            title="Capture de l'application France",
            missing_note=(
                "Ajoutez une capture d'écran finale de l'application France GASPAR vs JRC "
                "dans generated_figures/france_commune_app.png pour remplacer ce visuel."
            ),
        )
        italy_figure = resolve_or_placeholder(
            "italy_hazard_map.png",
            args.figures_dir,
            temp_dir,
            fallback=italy_fallback,
            title="Illustration Italie",
            missing_note=(
                "Ajoutez un visuel italien final si vous souhaitez montrer le cas ISPRA "
                "ou une carte équivalente issue de vos travaux."
            ),
        )
        flood_mix_figure = resolve_or_placeholder(
            "flood_lgd_source_mix_france_t20.png",
            args.figures_dir,
            temp_dir,
            fallback=None,
            title="Mix des sources FLOOD_LGD",
            missing_note=(
                "Générez cette figure sur l'autre PC via generate_figures_for_other_pc.py "
                "à partir de l'export France T20 final."
            ),
        )
        proprietary_figure = resolve_or_placeholder(
            "lgd_portfolio_results.png",
            args.figures_dir,
            temp_dir,
            fallback=None,
            title="Résultats portefeuille / LGD",
            missing_note=(
                "Insérez ici votre graphique final de portefeuille, votre première lecture LGD "
                "ou toute figure confidentielle produite sur le poste métier."
            ),
        )
        comparison_figure = resolve_or_placeholder(
            "jrc_gaspar_comparison_snapshot.png",
            args.figures_dir,
            temp_dir,
            fallback=match_chart,
            title="Synthèse comparaison JRC vs GASPAR",
            missing_note=(
                "Ce visuel peut être remplacé automatiquement par la figure générée sur l'autre PC "
                "à partir des dossiers comparison 7d / 30d."
            ),
        )

        document = Document()
        configure_document(document)
        add_cover_page(document, args)
        add_toc_page(document)
        add_introduction(document)
        add_section_one(document)
        add_section_two(document, pipeline_figure, italy_figure)
        add_section_three(document, flood_mix_figure)
        add_section_four(document, raster_figure, app_figure, comparison_figure)
        add_section_five(document, proprietary_figure)
        add_section_six(document)
        add_annexes(document)
        document.save(args.output)

    return args.output


def main() -> None:
    args = parse_args()
    output_path = build_document(args)
    print(f"Rapport généré : {output_path}")


if __name__ == "__main__":
    main()
