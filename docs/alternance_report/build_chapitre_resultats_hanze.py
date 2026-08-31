from __future__ import annotations

import math
from pathlib import Path

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "alternance_report" / "chapitre_resultats_hanze"
DOCX_OUT = OUT_DIR / "resultats_conclusion_jrc_gaspar_hanze.docx"
RESULTS_7D = ROOT / "data" / "processed" / "jrc_gaspar_hanze_comparison_7d"
RESULTS_30D = ROOT / "data" / "processed" / "jrc_gaspar_hanze_comparison_30d"
NUTS3_MAP = (
    ROOT
    / "docs"
    / "alternance_report"
    / "annexe_nuts_lau"
    / "carte_nuts3_france_italie.png"
)

NAVY = "17365D"
BLUE = "2E74B5"
INK = "1F2937"
GRAY = "667085"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FB"
GOLD = "A66A00"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def image_font(size: int, bold: bool = False):
    paths = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for path in paths:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def set_run_font(run, size=None, bold=None, italic=None, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = rgb(color)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, headers, rows, widths, font_size=9.2):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shade_cell(cell, LIGHT_GRAY)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=font_size, bold=True, color=NAVY)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            cells[index].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT
            )
            if row_index % 2 == 1:
                shade_cell(cells[index], PALE_BLUE)
            for run in cells[index].paragraphs[0].runs:
                set_run_font(run, size=font_size)
    set_table_geometry(table, widths)
    return table


def add_caption(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, italic=True, color=GRAY)


def add_body(document, text, keep=False):
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.keep_together = keep
    return paragraph


def add_callout(document, label, text):
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_BLUE)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label} — ")
    set_run_font(label_run, size=10.2, bold=True, color=NAVY)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.2, color=INK)
    set_table_geometry(table, [9360])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def make_rate_chart(summary_7d: pd.DataFrame, summary_30d: pd.DataFrame, output: Path):
    width, height = 1700, 920
    image = Image.new("RGB", (width, height), (248, 250, 252))
    draw = ImageDraw.Draw(image)
    draw.text((75, 45), "Taux d’événements appariés selon la fenêtre temporelle", font=image_font(38, True), fill=(23, 54, 93))
    draw.text((76, 96), "Comparaison au niveau NUTS 3 — France, 2015–2024", font=image_font(23), fill=(102, 112, 133))

    categories = []
    for comparison in ["JRC vs GASPAR", "JRC vs HANZE", "GASPAR vs HANZE"]:
        row7 = summary_7d.loc[summary_7d["comparison"].eq(comparison)].iloc[0]
        row30 = summary_30d.loc[summary_30d["comparison"].eq(comparison)].iloc[0]
        categories.extend(
            [
                (f"{row7['left_source']} dans\n{row7['right_source']}", row7["left_match_share"], row30["left_match_share"]),
                (f"{row7['right_source']} dans\n{row7['left_source']}", row7["right_match_share"], row30["right_match_share"]),
            ]
        )

    chart_left, chart_top, chart_right, chart_bottom = 100, 180, 1610, 735
    draw.line((chart_left, chart_bottom, chart_right, chart_bottom), fill=(151, 161, 176), width=2)
    for tick in range(0, 101, 20):
        y = chart_bottom - (tick / 100) * (chart_bottom - chart_top)
        draw.line((chart_left, y, chart_right, y), fill=(222, 226, 232), width=1)
        draw.text((44, y), f"{tick} %", font=image_font(17), anchor="lm", fill=(102, 112, 133))

    group_w = (chart_right - chart_left) / len(categories)
    bar_w = 58
    colors = [(46, 116, 181), (166, 106, 0)]
    for index, (label, value7, value30) in enumerate(categories):
        center = chart_left + group_w * (index + 0.5)
        for offset, value, color in ((-bar_w / 2, value7, colors[0]), (bar_w / 2, value30, colors[1])):
            x0 = center + offset - bar_w / 2
            x1 = center + offset + bar_w / 2
            y0 = chart_bottom - value * (chart_bottom - chart_top)
            draw.rounded_rectangle((x0, y0, x1, chart_bottom), radius=6, fill=color)
            draw.text((center + offset, y0 - 12), f"{value * 100:.1f} %", font=image_font(17, True), anchor="ms", fill=(31, 41, 55))
        draw.multiline_text((center, 768), label, font=image_font(17), anchor="ma", align="center", fill=(31, 41, 55), spacing=4)

    draw.rectangle((590, 860, 616, 886), fill=colors[0])
    draw.text((628, 873), "Fenêtre de 7 jours", font=image_font(18), anchor="lm", fill=(31, 41, 55))
    draw.rectangle((880, 860, 906, 886), fill=colors[1])
    draw.text((918, 873), "Fenêtre de 30 jours", font=image_font(18), anchor="lm", fill=(31, 41, 55))
    image.save(output, dpi=(220, 220))


def make_source_diagram(output: Path):
    image = Image.new("RGB", (1700, 880), (248, 250, 252))
    draw = ImageDraw.Draw(image, "RGBA")
    draw.text((75, 45), "Trois sources, trois dimensions complémentaires", font=image_font(38, True), fill=(23, 54, 93))
    draw.text((76, 96), "La comparaison rapproche des objets différents sans confondre leur signification", font=image_font(23), fill=(102, 112, 133))
    centers = [(400, 390), (850, 390), (1300, 390)]
    colors = [(46, 116, 181, 220), (166, 106, 0, 220), (63, 124, 97, 220)]
    titles = ["JRC", "GASPAR", "HANZE"]
    subtitles = [
        "Observation physique\nsatellitaire",
        "Reconnaissance\nadministrative CatNat",
        "Inventaire historique\ndocumenté",
    ]
    for center, color, title, subtitle in zip(centers, colors, titles, subtitles):
        cx, cy = center
        draw.ellipse((cx - 175, cy - 175, cx + 175, cy + 175), fill=color, outline=(255, 255, 255), width=5)
        draw.text((cx, cy - 35), title, font=image_font(36, True), anchor="mm", fill=(255, 255, 255))
        draw.multiline_text((cx, cy + 45), subtitle, font=image_font(21), anchor="mm", align="center", fill=(255, 255, 255), spacing=6)
    for left, right in ((0, 1), (1, 2)):
        x0 = centers[left][0] + 180
        x1 = centers[right][0] - 180
        draw.line((x0, 390, x1, 390), fill=(102, 112, 133), width=5)
        draw.polygon([(x1, 390), (x1 - 18, 380), (x1 - 18, 400)], fill=(102, 112, 133))
    draw.rounded_rectangle((280, 660, 1420, 800), radius=22, fill=(232, 238, 245), outline=(186, 199, 214), width=2)
    draw.text((850, 700), "Règle commune de comparaison", font=image_font(25, True), anchor="mm", fill=(23, 54, 93))
    draw.text((850, 752), "Même territoire NUTS 3 + intervalles temporels compatibles", font=image_font(23), anchor="mm", fill=(31, 41, 55))
    image.save(output, dpi=(220, 220))


def result_rows(summary: pd.DataFrame):
    rows = []
    for _, row in summary.iterrows():
        rows.append(
            (
                row["comparison"],
                f"{row['left_matched_events']}/{row['left_total_events']} ({row['left_match_share'] * 100:.1f} %)",
                f"{row['right_matched_events']}/{row['right_total_events']} ({row['right_match_share'] * 100:.1f} %)",
            )
        )
    return rows


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    summary_7d = pd.read_csv(RESULTS_7D / "pairwise_summary.csv")
    summary_30d = pd.read_csv(RESULTS_30D / "pairwise_summary.csv")
    triple_7d = pd.read_csv(RESULTS_7D / "three_source_summary.csv").iloc[0]
    triple_30d = pd.read_csv(RESULTS_30D / "three_source_summary.csv").iloc[0]
    rate_chart = OUT_DIR / "comparaison_taux_7j_30j.png"
    source_diagram = OUT_DIR / "sources_jrc_gaspar_hanze.png"
    make_rate_chart(summary_7d, summary_30d, rate_chart)
    make_source_diagram(source_diagram)

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "Rapport d’alternance | DataCollection"
    set_run_font(header.runs[0], size=8.5, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Comparaison JRC–GASPAR–HANZE")
    set_run_font(footer_run, size=8.5, color=GRAY)

    # Editorial cover
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(90)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("RÉSULTATS ET CONCLUSION")
    set_run_font(run, size=11, bold=True, color=GOLD)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Comparaison des sources d’inondation\nJRC, GASPAR et HANZE")
    set_run_font(run, size=28, bold=True, color=NAVY)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(32)
    run = subtitle.add_run("France, 2015–2024 | Analyse au niveau NUTS 3")
    set_run_font(run, size=14, color=GRAY)
    lead = document.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.left_indent = Inches(0.6)
    lead.paragraph_format.right_indent = Inches(0.6)
    run = lead.add_run(
        "Document prêt à intégrer dans le rapport d’alternance : méthodologie, tableaux chiffrés, figures, interprétation et conclusion."
    )
    set_run_font(run, size=11.5, italic=True, color=INK)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("1. Construction de la comparaison", level=1)
    add_body(
        document,
        "Les trois sources décrivent des dimensions différentes du risque d’inondation. JRC fournit principalement une observation physique issue de produits satellitaires, GASPAR enregistre les reconnaissances administratives CatNat en France et HANZE rassemble des événements historiques documentés ainsi que certains de leurs impacts. La comparaison vise donc à mesurer leur compatibilité spatio-temporelle, sans supposer qu’elles représentent exactement le même objet.",
    )
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(source_diagram), width=Inches(6.45))
    add_caption(document, "Figure 1 — Positionnement complémentaire des trois sources.")

    add_table(
        document,
        ["Source", "Nature de l’information", "Unité événementielle utilisée"],
        [
            ("JRC", "Emprise et profondeur issues de produits satellitaires", "Identifiant d’événement JRC"),
            ("GASPAR", "Reconnaissance administrative CatNat par commune", "Proxy : arrêté + dates de début et de fin"),
            ("HANZE", "Inventaire historique et impacts documentés", "Identifiant d’événement HANZE"),
        ],
        [1500, 4100, 3760],
        font_size=9,
    )
    add_caption(document, "Tableau 1 — Nature et unité de comparaison des sources.")

    document.add_heading("1.1 Identifiant d’événement GASPAR", level=2)
    add_body(
        document,
        "Dans GASPAR, la variable cod_nat_catnat identifie un arrêté CatNat, c’est-à-dire un document administratif, et non une inondation précise. Un même arrêté peut regrouper plusieurs communes, dates ou épisodes différents ; inversement, une seule inondation peut apparaître dans plusieurs arrêtés. Pour approcher la notion d’événement, un identifiant a donc été construit en combinant le code de l’arrêté, la date de début et la date de fin.",
    )
    add_callout(
        document,
        "Limite méthodologique",
        "Cet identifiant est un proxy administratif. Les 3 505 groupes GASPAR ne doivent pas être interprétés comme 3 505 inondations physiques indépendantes.",
    )

    document.add_heading("1.2 Règle de rapprochement", level=2)
    add_body(
        document,
        "La comparaison porte sur la France entre le 1er janvier 2015 et le 31 décembre 2024. NUTS 3 est retenu comme niveau géographique commun, car HANZE ne fournit généralement pas une localisation communale. Deux événements sont considérés comme compatibles lorsqu’ils concernent le même territoire NUTS 3 et que leurs intervalles de dates se chevauchent après application d’une tolérance de sept ou trente jours.",
    )
    if NUTS3_MAP.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(NUTS3_MAP), width=Inches(6.4))
        add_caption(document, "Figure 2 — Exemple du découpage NUTS 3 utilisé pour l’harmonisation géographique.")

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("2. Résultats", level=1)
    add_body(
        document,
        "L’univers étudié comprend 288 événements JRC, 3 505 groupes d’événements GASPAR et 43 événements HANZE. Les tableaux suivants présentent, pour chaque paire de sources, le nombre d’événements trouvant au moins une correspondance et leur part dans la source considérée.",
    )
    document.add_heading("2.1 Fenêtre temporelle stricte de sept jours", level=2)
    add_table(
        document,
        ["Comparaison", "Première source appariée", "Seconde source appariée"],
        result_rows(summary_7d),
        [2400, 3480, 3480],
        font_size=9.2,
    )
    add_caption(document, "Tableau 2 — Couverture des appariements avec une fenêtre de sept jours.")
    add_body(
        document,
        f"La concordance simultanée des trois sources couvre {int(triple_7d['jrc_events_in_triple_matches'])} événements JRC, {int(triple_7d['gaspar_events_in_triple_matches'])} groupes GASPAR et {int(triple_7d['hanze_events_in_triple_matches'])} événements HANZE. Cette variante privilégie les rapprochements temporels les plus stricts.",
    )

    document.add_heading("2.2 Fenêtre temporelle élargie de trente jours", level=2)
    add_table(
        document,
        ["Comparaison", "Première source appariée", "Seconde source appariée"],
        result_rows(summary_30d),
        [2400, 3480, 3480],
        font_size=9.2,
    )
    add_caption(document, "Tableau 3 — Couverture des appariements avec une fenêtre de trente jours.")
    add_body(
        document,
        f"Avec trente jours, la concordance simultanée couvre {int(triple_30d['jrc_events_in_triple_matches'])} événements JRC, {int(triple_30d['gaspar_events_in_triple_matches'])} groupes GASPAR et {int(triple_30d['hanze_events_in_triple_matches'])} événements HANZE. L’augmentation du recouvrement montre l’importance des décalages de dates et de la fragmentation administrative.",
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("2.3 Comparaison visuelle des taux", level=2)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(rate_chart), width=Inches(6.45))
    add_caption(document, "Figure 3 — Taux d’événements appariés selon la source et la fenêtre temporelle.")
    add_callout(
        document,
        "Lecture",
        "Les pourcentages ne sont pas symétriques : 27 événements HANZE sur 43 peuvent correspondre à seulement 48 événements JRC sur 288. Chaque taux utilise le nombre total d’événements de sa propre source comme dénominateur.",
    )

    document.add_heading("2.4 Interprétation des résultats", level=2)
    add_body(
        document,
        "Le rapprochement JRC–GASPAR est le plus large : avec trente jours, 59,4 % des événements JRC et 53,8 % des groupes GASPAR trouvent une correspondance. Le passage de sept à trente jours augmente nettement le recouvrement, ce qui confirme qu’une partie des écarts provient des dates retenues dans la procédure administrative et de la fragmentation d’un épisode entre plusieurs groupes GASPAR.",
    )
    add_body(
        document,
        "Le rapprochement JRC–HANZE est plus faible du point de vue de JRC : 13,2 % avec sept jours et 16,7 % avec trente jours. En revanche, 53,5 % puis 62,8 % des événements HANZE trouvent une correspondance JRC. HANZE est donc plus sélective et ne constitue pas un inventaire exhaustif des événements détectés par satellite.",
    )
    add_body(
        document,
        "Tous les événements HANZE trouvent une correspondance GASPAR sur la période étudiée, mais ce résultat ne signifie pas que les deux bases sont équivalentes. Un événement HANZE régional peut correspondre à plusieurs groupes administratifs GASPAR. Le volume de correspondances reflète donc aussi la granularité et la construction de chaque source.",
    )

    document.add_heading("2.5 Limites et précautions", level=2)
    add_table(
        document,
        ["Source", "Principale limite pour l’interprétation"],
        [
            ("JRC", "Couverture liée à la disponibilité des produits satellitaires et à leur capacité de détection."),
            ("GASPAR", "Source administrative dépendant des demandes communales, des critères CatNat et du proxy d’événement construit."),
            ("HANZE", "Inventaire historique sélectif, avec une localisation souvent limitée au niveau NUTS 3."),
            ("Appariement", "Une compatibilité NUTS 3 et temporelle ne démontre pas une emprise physique identique."),
        ],
        [1700, 7660],
        font_size=9.3,
    )
    add_caption(document, "Tableau 4 — Principales limites des résultats.")

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("3. Conclusion", level=1)
    add_body(
        document,
        "La comparaison entre JRC, GASPAR et HANZE montre qu’aucune source ne fournit, à elle seule, une représentation complète des inondations. JRC apporte la preuve physique la plus proche du phénomène observé ; GASPAR décrit sa reconnaissance administrative et assurantielle ; HANZE apporte une profondeur historique et des informations documentaires sur les événements et leurs impacts.",
    )
    add_body(
        document,
        "L’intégration de HANZE confirme la complémentarité des sources. Une majorité des événements HANZE récents trouve une correspondance dans JRC et tous sont reliés à au moins un groupe GASPAR. Toutefois, le recouvrement reste faible du point de vue de JRC, ce qui montre que HANZE est une source sélective et non un catalogue exhaustif des emprises satellitaires.",
    )
    add_body(
        document,
        "Les résultats soulignent également l’importance des choix méthodologiques. La fenêtre de sept jours fournit une comparaison stricte, tandis que la fenêtre de trente jours absorbe davantage les décalages temporels au prix d’un risque plus élevé d’associer des épisodes voisins. Les appariements doivent donc être interprétés comme des compatibilités spatio-temporelles et non comme l’identification certaine d’un même événement physique.",
    )
    add_body(
        document,
        "La stratégie retenue pour DataCollection consiste ainsi à hiérarchiser les preuves. Une confirmation JRC au voisinage d’un point apporte l’information physique la plus précise. Une correspondance simultanée entre JRC, GASPAR et HANZE renforce la confiance dans l’existence de l’événement. GASPAR ou HANZE seules constituent plutôt des signaux administratifs ou historiques, à compléter avec les couches TRI, ISPRA ou d’autres observations spatiales.",
    )
    add_body(
        document,
        "En définitive, le principal apport de la mission réside dans la construction d’une méthode capable de combiner des sources hétérogènes sans confondre leur signification. Les exports FLOOD_LGD issus de cette chaîne fournissent désormais une base documentée et auditable pour étudier l’incidence future des inondations sur le risque de crédit et la perte en cas de défaut.",
    )
    add_callout(
        document,
        "Résultat central",
        "La valeur du dispositif ne vient pas d’une source unique, mais de la convergence contrôlée entre preuve physique, reconnaissance administrative et documentation historique.",
    )

    for current in document.sections:
        current.page_width = Inches(8.5)
        current.page_height = Inches(11)
        current.top_margin = Inches(1)
        current.bottom_margin = Inches(1)
        current.left_margin = Inches(1)
        current.right_margin = Inches(1)
        current.header_distance = Inches(0.492)
        current.footer_distance = Inches(0.492)
        current.header.is_linked_to_previous = True
        current.footer.is_linked_to_previous = True

    document.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    build_document()
