from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


DEFAULT_OUTPUT = Path("docs/alternance_report/note_synthese_alternance_data_collection.docx")
BODY_FONT = "Arial"
TITLE_COLOR = RGBColor(21, 56, 87)
ACCENT_COLOR = RGBColor(0, 102, 68)
MUTED_COLOR = RGBColor(96, 96, 96)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Génère une note de synthèse d'alternance en français au format demandé "
            "par les consignes MoSEF."
        )
    )
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--student-name", default="Juan David Alonso")
    parser.add_argument("--student-number", default="[à renseigner]")
    parser.add_argument(
        "--program",
        default="MASTER MODELISATIONS STATISTIQUES ECONOMIQUES ET FINANCIERES (MOSEF DATA SCIENCES)",
    )
    parser.add_argument("--promotion", default="2025-2026")
    parser.add_argument("--company", default="BNP Paribas")
    parser.add_argument(
        "--mission-title",
        default=(
            "Construction d'un pipeline de données d'inondation pour le screening "
            "du risque physique et la préparation d'analyses LGD"
        ),
    )
    parser.add_argument("--manager", default="[à renseigner]")
    parser.add_argument("--unit", default="Risk AIR / Data / Paris")
    parser.add_argument("--tutor", default="[à renseigner]")
    return parser.parse_args()


def set_run_font(run, *, size: int, bold: bool = False, italic: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)


def apply_a4_layout(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)


def restart_page_numbering(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def configure_document(document: Document) -> None:
    for section in document.sections:
        apply_a4_layout(section)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after, color in [
        ("Heading 1", 15, 14, 6, TITLE_COLOR),
        ("Heading 2", 13, 10, 5, ACCENT_COLOR),
    ]:
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.08


def add_body_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    set_run_font(run, size=10, color=MUTED_COLOR)
    add_page_field(paragraph)


def add_cover_page(document: Document, args: argparse.Namespace) -> None:
    def centered(text: str, *, size: int, bold: bool = False, color: RGBColor = TITLE_COLOR, after: int = 0) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(after)
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)

    centered(args.student_name.upper(), size=14, bold=True, after=4)
    centered(f"Numéro étudiant : {args.student_number}", size=11, color=MUTED_COLOR, after=8)
    centered(args.program, size=12, bold=True, after=16)
    centered("NOTE DE SYNTHÈSE", size=20, bold=True, color=ACCENT_COLOR, after=6)
    centered(f"Promotion « {args.promotion} »", size=11, color=MUTED_COLOR, after=20)
    centered(args.company, size=16, bold=True, after=12)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(20)
    run = title.add_run(args.mission_title)
    set_run_font(run, size=14, bold=True, color=TITLE_COLOR)

    meta_lines = [
        ("Maître d'alternance", args.manager),
        ("Unité d'accueil – lieu de stage", args.unit),
        ("Tuteur pédagogique", args.tutor),
    ]
    for label, value in meta_lines:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        label_run = paragraph.add_run(f"{label} : ")
        set_run_font(label_run, size=11, bold=True, color=TITLE_COLOR)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=11)


def add_section_heading(document: Document, text: str) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run(text)


def add_paragraph(document: Document, text: str, *, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=12, italic=italic)


def add_note_body(document: Document) -> None:
    add_section_heading(document, "1. Cadre et enjeux de la mission")
    add_paragraph(
        document,
        "La présente note de synthèse a pour objet de présenter, de manière directement lisible pour un lecteur académique ou professionnel, le cadre de l'alternance menée au sein de BNP Paribas autour du dépôt DataCollection. Elle ne constitue pas un résumé exhaustif du rapport de stage : elle vise avant tout à expliciter le contexte de la mission, les objectifs poursuivis, les résultats les plus structurants et les suites opérationnelles qui paraissent les plus utiles.",
    )
    add_paragraph(
        document,
        "La mission s'inscrit dans un besoin métier clair : disposer d'une chaîne de traitement reproductible permettant de rapprocher des sources d'inondation externes, hétérogènes et parfois difficiles à auditer, avec des usages risques plus directement exploitables par des équipes de portefeuille et d'analyse crédit. L'enjeu ne portait donc pas uniquement sur la collecte de données, mais sur la transformation d'objets géospatiaux complexes en livrables interprétables, documentés et réutilisables.",
    )

    add_section_heading(document, "2. Objectifs opérationnels poursuivis")
    add_paragraph(
        document,
        "Le premier objectif a consisté à consolider un socle géodata robuste pour la France et l'Italie, en articulant plusieurs familles de sources : rasters JRC, historique GASPAR, événements HANZE, TRI, référentiels LAU/NUTS et rapprochements France INSEE. Le second objectif a été d'industrialiser des workflows de screening point par point, capables de traiter à la fois des cas T20 et des jeux de collatéraux. Le troisième objectif a porté sur la restitution : produire des exports FLOOD_LGD, des outils d'audit, des applications de visualisation et une documentation suffisamment claire pour réduire la dépendance à une seule personne.",
    )

    add_section_heading(document, "3. Principaux résultats obtenus")
    add_paragraph(
        document,
        "Le travail réalisé a permis d'aboutir à plusieurs résultats concrets. D'abord, le dépôt DataCollection porte désormais une chaîne de traitement cohérente allant de la source brute jusqu'aux tables finales FLOOD_LGD. Ensuite, les règles de rapprochement entre sources ont été explicitées et testées, notamment pour les cas France T20, France collaterals, Italie T20 et Italie collaterals. Par ailleurs, l'alternance a conduit à enrichir l'environnement de travail avec des applications Streamlit, des scripts d'audit, des guides détaillés, des dictionnaires de colonnes et des supports de restitution, ce qui renforce la lisibilité de la méthode et sa capacité de transmission dans le temps.",
    )
    add_paragraph(
        document,
        "Au-delà de la production technique, un apport important tient à la structuration de la démarche. Les sorties finales ne se limitent plus à des scripts isolés : elles forment un patrimoine de travail articulé, commenté et vérifiable. Cette structuration facilite la reprise par d'autres personnes de l'équipe, améliore les échanges avec des interlocuteurs non spécialistes du SIG et prépare plus proprement la phase suivante d'analyse portefeuille / LGD sur le poste métier.",
    )

    add_section_heading(document, "4. Préconisations et suites proposées")
    add_paragraph(
        document,
        "La priorité la plus immédiate consiste à réinjecter, depuis le poste métier, les résultats confidentiels qui ne figurent pas encore dans cette version du rapport : figures quantitatives finales, cas d'usage portefeuille et premières lectures LGD. Une seconde suite logique est de croiser de manière systématique les variables de flood avec les données d'exposition et les dates de défaut afin de tester plusieurs fenêtres temporelles et plusieurs mailles d'agrégation. Enfin, il paraît utile de poursuivre l'effort de capitalisation en stabilisant les derniers artefacts de démonstration, notamment pour la soutenance, afin de disposer d'un discours homogène entre la donnée brute, la méthode, les contrôles et les usages décisionnels.",
    )
    add_paragraph(
        document,
        "En synthèse, l'alternance a surtout permis de transformer une problématique de collecte géodata en un dispositif plus mature, orienté vers l'usage et l'auditabilité. Le dépôt produit aujourd'hui un cadre de travail crédible pour prolonger l'analyse, et la suite du travail doit désormais porter sur l'exploitation métier des résultats autant que sur l'approfondissement technique.",
    )


def build_document(args: argparse.Namespace) -> Path:
    document = Document()
    configure_document(document)
    add_cover_page(document, args)

    body_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    apply_a4_layout(body_section)
    restart_page_numbering(body_section, start=1)
    add_body_footer(body_section)

    add_note_body(document)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    return args.output


def main() -> None:
    args = parse_args()
    output = build_document(args)
    print(f"Note de synthèse générée : {output}")


if __name__ == "__main__":
    main()
