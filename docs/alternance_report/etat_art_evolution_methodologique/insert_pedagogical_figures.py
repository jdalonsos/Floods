from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
INPUT = ROOT / "etat_art_et_evolution_methodologique_inondations.docx"
OUTPUT = ROOT / "etat_art_et_evolution_methodologique_inondations_illustre.docx"
IMAGES = ROOT / "web_images"


def insert_after(paragraph, new_paragraph):
    paragraph._p.addnext(new_paragraph._p)
    return new_paragraph


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    element = p_pr.find(qn("w:keepNext"))
    if value and element is None:
        p_pr.append(OxmlElement("w:keepNext"))
    elif not value and element is not None:
        p_pr.remove(element)


def set_keep_together(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepLines")) is None:
        p_pr.append(OxmlElement("w:keepLines"))


def add_alt_text(inline_shape, text):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", text)
    doc_pr.set("title", "Schéma pédagogique sur les inondations")


def add_figure_after(doc, anchor, image_name, width, number, caption, reading, alt_text):
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(7)
    image_paragraph.paragraph_format.space_after = Pt(2)
    shape = image_paragraph.add_run().add_picture(str(IMAGES / image_name), width=Inches(width))
    add_alt_text(shape, alt_text)
    set_keep_with_next(image_paragraph)
    insert_after(anchor, image_paragraph)

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(0)
    caption_paragraph.paragraph_format.space_after = Pt(3)
    caption_paragraph.paragraph_format.left_indent = Inches(0.35)
    caption_paragraph.paragraph_format.right_indent = Inches(0.35)
    lead = caption_paragraph.add_run(f"Figure {number} — ")
    lead.bold = True
    lead.font.size = Pt(8.5)
    lead.font.color.rgb = RGBColor(55, 65, 81)
    body = caption_paragraph.add_run(caption)
    body.italic = True
    body.font.size = Pt(8.5)
    body.font.color.rgb = RGBColor(55, 65, 81)
    set_keep_with_next(caption_paragraph)
    insert_after(image_paragraph, caption_paragraph)

    reading_paragraph = doc.add_paragraph()
    reading_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    reading_paragraph.paragraph_format.left_indent = Inches(0.45)
    reading_paragraph.paragraph_format.right_indent = Inches(0.45)
    reading_paragraph.paragraph_format.space_before = Pt(2)
    reading_paragraph.paragraph_format.space_after = Pt(8)
    label = reading_paragraph.add_run("Lecture : ")
    label.bold = True
    label.font.size = Pt(9)
    label.font.color.rgb = RGBColor(31, 78, 121)
    explanation = reading_paragraph.add_run(reading)
    explanation.font.size = Pt(9)
    explanation.font.color.rgb = RGBColor(55, 65, 81)
    set_keep_together(reading_paragraph)
    insert_after(caption_paragraph, reading_paragraph)


def find_paragraph(doc, prefix):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraphe introuvable : {prefix}")


def add_reference_after(doc, anchor, text):
    paragraph = doc.add_paragraph()
    paragraph.style = anchor.style
    paragraph.paragraph_format.space_after = anchor.paragraph_format.space_after
    run = paragraph.add_run(text)
    run.font.size = Pt(8.5)
    insert_after(anchor, paragraph)
    return paragraph


def main():
    doc = Document(INPUT)
    if doc.inline_shapes:
        add_alt_text(
            doc.inline_shapes[0],
            "Schéma comparant les travaux antérieurs et l’approche méthodologique développée cette année pour mesurer le risque d’inondation.",
        )

    anchor_21 = find_paragraph(doc, "L’humidité préalable des sols")
    add_figure_after(
        doc,
        anchor_21,
        "schema_fr_saturation_sol.png",
        5.85,
        2,
        "Influence de l’humidité antérieure sur l’infiltration et le ruissellement. Schéma pédagogique adapté du BRGM / SIGES Rhin-Meuse (2025).",
        "au-dessus de la nappe, les pores du sol contiennent encore de l’air et peuvent stocker une partie de la pluie. Dans la zone saturée, les pores sont déjà remplis d’eau ; la capacité de stockage supplémentaire est donc faible, ce qui favorise le ruissellement et la montée des niveaux d’eau.",
        "Comparaison en français d’un sol sec, qui peut encore infiltrer et stocker de l’eau, et d’un sol saturé, qui favorise le ruissellement.",
    )

    anchor_22 = find_paragraph(doc, "La topographie, la pente")
    add_figure_after(
        doc,
        anchor_22,
        "schema_fr_impermeabilisation.png",
        5.85,
        3,
        "Effet de l’imperméabilisation sur la répartition de la pluie. Schéma pédagogique adapté du guide Vers la ville perméable (Agence de l’eau Rhône Méditerranée Corse, 2017).",
        "lorsque la part de surfaces imperméables augmente, l’infiltration diminue fortement tandis que le ruissellement augmente. L’eau rejoint alors les réseaux et les cours d’eau plus rapidement, ce qui accentue les pics de débit et le risque d’inondation urbaine.",
        "Comparaison en français d’un bassin végétalisé et d’un bassin très urbanisé, avec les parts indicatives d’infiltration et de ruissellement.",
    )

    anchor_23 = find_paragraph(doc, "Toutes les inondations ne relèvent pas")
    add_figure_after(
        doc,
        anchor_23,
        "schema_fr_types_inondation.png",
        5.85,
        4,
        "Principaux mécanismes d’inondation. Schéma pédagogique adapté du Guide métropolitain de l’aménagement résilient (Cerema, 2023).",
        "une inondation peut provenir de la mer, d’un cours d’eau, du ruissellement de surface, d’un réseau d’assainissement dépassé ou d’une remontée de nappe. Plusieurs mécanismes peuvent se produire simultanément et toucher les mêmes bâtiments ou infrastructures.",
        "Schéma entièrement en français présentant la submersion marine, le débordement de cours d’eau, le ruissellement pluvial et la remontée de nappe.",
    )

    # Remove the original empty hard-page-break paragraph. With the added figures
    # it can create a nearly empty page before section 4.
    heading_4 = find_paragraph(doc, "4. Recentrage")
    previous = heading_4._p.getprevious()
    if previous is not None:
        previous_text = "".join(previous.itertext()).strip()
        page_breaks = [
            br for br in previous.findall(".//" + qn("w:br"))
            if br.get(qn("w:type")) == "page"
        ]
        if not previous_text and page_breaks:
            previous.getparent().remove(previous)
    heading_4.paragraph_format.page_break_before = False

    last_reference = find_paragraph(doc, "Joint Research Centre (2024)")
    ref = add_reference_after(
        doc,
        last_reference,
        "BRGM / SIGES Rhin-Meuse (2025). L’eau souterraine : de l’eau contenue dans les roches. https://sigesrm.brgm.fr/L-eau-souterraine-de-l-eau-contenue-dans-les-roches",
    )
    ref = add_reference_after(
        doc,
        ref,
        "Agence de l’eau Rhône Méditerranée Corse (2017). Vers la ville perméable : comment désimperméabiliser les sols ? Guide technique du SDAGE.",
    )
    add_reference_after(
        doc,
        ref,
        "Cerema (2023). Guide métropolitain de l’aménagement résilient face au risque d’inondation.",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
