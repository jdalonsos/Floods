from __future__ import annotations

import importlib.util
from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
BASE_SCRIPT = ROOT / "docs" / "alternance_report" / "build_chapitre_resultats_hanze.py"
OUT_DIR = ROOT / "docs" / "alternance_report" / "etat_art_evolution_methodologique"
DOCX_OUT = OUT_DIR / "etat_art_et_evolution_methodologique_inondations.docx"
SCHEMA_OUT = OUT_DIR / "evolution_methodologique.png"

spec = importlib.util.spec_from_file_location("report_helpers", BASE_SCRIPT)
helpers = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(helpers)


def add_body(document: Document, text: str):
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.25
    return paragraph


def add_source(document: Document, text: str):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        helpers.set_run_font(run, size=8.2, color=helpers.INK)


def configure_document(document: Document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = helpers.rgb(helpers.INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, helpers.BLUE, 18, 10),
        ("Heading 2", 13, helpers.BLUE, 12, 6),
        ("Heading 3", 12, helpers.NAVY, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = helpers.rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_centered(document: Document, text: str, size: float, color: str, bold=False, italic=False, after=8):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    helpers.set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def make_schema(path: Path):
    width, height = 1500, 760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title = helpers.image_font(42, True)
    subtitle = helpers.image_font(22)
    hfont = helpers.image_font(30, True)
    body = helpers.image_font(23)
    small = helpers.image_font(20)

    draw.text((70, 48), "Évolution de la mesure du risque d’inondation", font=title, fill="#17365D")
    draw.text(
        (70, 105),
        "Du signal météorologique à la preuve d’un événement observé ou documenté",
        font=subtitle,
        fill="#667085",
    )

    boxes = [
        (65, 190, 670, 630, "#F7EEDC", "#A66A00"),
        (830, 190, 1435, 630, "#E8EEF5", "#2E74B5"),
    ]
    for x1, y1, x2, y2, fill, outline in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=28, fill=fill, outline=outline, width=4)

    draw.text((110, 230), "Travaux antérieurs", font=hfont, fill="#7A4B00")
    previous = [
        "Échelle principalement française",
        "TRI + GASPAR + Copernicus",
        "Précipitations extrêmes comme proxy",
        "Indicateurs binaires par établissement",
    ]
    y = 305
    for item in previous:
        draw.ellipse((112, y + 7, 126, y + 21), fill="#A66A00")
        draw.text((145, y), item, font=body, fill="#1F2937")
        y += 66

    draw.text((880, 230), "Approche développée cette année", font=hfont, fill="#17365D")
    current = [
        "Extension et harmonisation européennes",
        "JRC + GASPAR + HANZE",
        "Inondations observées ou documentées",
        "Comparaison spatio-temporelle à NUTS 3",
    ]
    y = 305
    for item in current:
        draw.ellipse((882, y + 7, 896, y + 21), fill="#2E74B5")
        draw.text((915, y), item, font=body, fill="#1F2937")
        y += 66

    draw.line((690, 410, 810, 410), fill="#667085", width=8)
    draw.polygon([(810, 410), (780, 392), (780, 428)], fill="#667085")
    draw.text((643, 456), "Recentrage", font=small, fill="#667085")
    draw.text((570, 490), "sur le phénomène physique", font=small, fill="#667085")
    image.save(path, quality=95)


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_schema(SCHEMA_OUT)
    document = Document()
    configure_document(document)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(118)
    add_centered(document, "REVUE DE LITTÉRATURE ET APPORTS DE LA MISSION", 10.5, helpers.GOLD, bold=True, after=18)
    add_centered(document, "État de l’art et évolution méthodologique", 28, helpers.NAVY, bold=True, after=8)
    add_centered(
        document,
        "Du dispositif développé l’année précédente à une mesure harmonisée des inondations à l’échelle européenne",
        14,
        helpers.GRAY,
        after=24,
    )
    add_centered(document, "Section prête à intégrer au rapport d’alternance", 10.5, helpers.GOLD, italic=True, after=72)
    helpers.add_callout(
        document,
        "Question méthodologique",
        "Pourquoi remplacer les indicateurs fondés sur les précipitations extrêmes par une identification directe, multi-source et européenne des inondations ?",
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("1. Positionnement par rapport aux travaux antérieurs", level=1)
    add_body(
        document,
        "Les travaux conduits lors de l’alternance précédente constituaient une première étape de mesure du risque physique d’inondation au niveau des établissements. Ils combinaient trois familles d’informations : les territoires à risque important d’inondation (TRI), les reconnaissances communales issues de GASPAR et des indicateurs de précipitations extrêmes produits à partir de données Copernicus. Cette approche permettait de construire des variables binaires simples, telles que la présence d’un établissement dans une zone TRI, l’existence d’une reconnaissance d’inondation dans sa commune ou le dépassement d’un percentile de précipitations.",
    )
    add_body(
        document,
        "Cette première architecture avait l’avantage d’être lisible et opérationnelle. Elle rapprochait une exposition structurelle, une reconnaissance administrative et un signal météorologique. Elle présentait toutefois deux limites majeures. D’une part, le périmètre demeurait principalement national, ce qui réduisait la comparabilité entre établissements situés dans différents pays européens. D’autre part, les précipitations extrêmes étaient utilisées comme approximation de l’inondation, alors qu’elles ne décrivent qu’un facteur déclencheur potentiel et non l’occurrence certaine du phénomène.",
    )
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(SCHEMA_OUT), width=Inches(6.45))
    helpers.add_caption(
        document,
        "Figure 1 — Synthèse de l’évolution méthodologique entre les deux années d’alternance.",
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("2. Pourquoi une forte précipitation n’implique pas nécessairement une inondation", level=1)
    add_body(
        document,
        "La littérature hydrologique montre que la relation entre pluie et inondation est non linéaire. Une précipitation intense constitue une sollicitation météorologique, mais sa transformation en ruissellement, en hausse du débit puis en débordement dépend de l’état du bassin versant. Deux épisodes présentant des cumuls similaires peuvent donc produire des conséquences très différentes. À l’inverse, une pluie d’intensité modérée peut provoquer une crue importante lorsqu’elle survient sur un sol déjà saturé ou au sein d’un bassin où les niveaux d’eau sont élevés.",
    )
    document.add_heading("2.1 L’état hydrique antérieur du bassin", level=2)
    add_body(
        document,
        "L’humidité préalable des sols contrôle la part de l’eau qui peut être infiltrée ou qui devient rapidement du ruissellement. Sur un sol sec et perméable, une partie importante de la pluie peut être absorbée ou stockée. Sur un sol saturé, la capacité d’infiltration et de stockage est réduite ; une fraction plus élevée de la pluie rejoint alors les cours d’eau. Une analyse menée sur 1 353 bassins européens montre que les changements de processus générateurs de crues, notamment le passage de pluies sur sols secs à des pluies sur sols humides, expliquent davantage certaines périodes riches en crues que la seule évolution des pluies extrêmes.",
    )
    document.add_heading("2.2 Les caractéristiques physiques et l’échelle spatiale", level=2)
    add_body(
        document,
        "La topographie, la pente, la géologie, la perméabilité des sols, la densité du réseau hydrographique, la taille du bassin et l’urbanisation modifient la vitesse et le volume du ruissellement. Une pluie très localisée peut dépasser un percentile dans une maille climatique sans toucher les zones contributives d’un cours d’eau. À l’inverse, des précipitations moins intenses mais étendues sur l’ensemble d’un bassin peuvent produire une onde de crue importante. L’imperméabilisation urbaine accélère également les transferts vers les réseaux et favorise les inondations pluviales, tandis que les ouvrages hydrauliques, barrages et zones d’expansion peuvent atténuer le pic.",
    )
    document.add_heading("2.3 La diversité des mécanismes de crue", level=2)
    add_body(
        document,
        "Toutes les inondations ne relèvent pas du même mécanisme. Les crues fluviales résultent du dépassement de la capacité d’un cours d’eau ; les inondations pluviales peuvent être liées à la saturation des réseaux ou au ruissellement urbain ; les remontées de nappe dépendent du stockage souterrain ; les submersions côtières sont principalement commandées par le niveau marin, les vagues et les surcotes. En Europe centrale et septentrionale, la fonte nivale, seule ou combinée à la pluie, constitue également un moteur important. Un indicateur de précipitation unique ne peut donc pas représenter correctement cette diversité.",
    )
    helpers.add_callout(
        document,
        "Conséquence",
        "Le dépassement du 90e, du 95e ou du 99e percentile mesure une anomalie pluviométrique. Il ne constitue ni une preuve d’inondation, ni une mesure directe de son emprise, de sa profondeur ou de ses impacts.",
    )

    document.add_heading("3. Passage à l’échelle européenne", level=1)
    add_body(
        document,
        "L’extension à l’échelle européenne répond à l’évolution du portefeuille étudié et à la nature transfrontalière du risque. Les grands bassins, tels que le Rhin, le Danube ou l’Elbe, traversent plusieurs États ; un même épisode peut ainsi affecter simultanément plusieurs régions et plusieurs cadres administratifs. Une approche strictement française ne permettait ni de couvrir l’ensemble des établissements ni de comparer les niveaux d’exposition selon une méthode commune.",
    )
    add_body(
        document,
        "Cette extension nécessite une harmonisation géographique. Les données nationales ne disposent pas toujours des mêmes unités, définitions ou niveaux de détail. Le recours à la nomenclature NUTS, notamment NUTS 3, fournit un référentiel territorial commun pour apparier les événements et rattacher les établissements. Ce choix facilite les comparaisons européennes et l’intégration de HANZE, mais implique une perte de précision par rapport à une localisation communale ou à une emprise satellitaire. Le niveau NUTS 3 est donc utilisé comme unité de rapprochement, non comme preuve que toute la région a été inondée.",
    )
    document.add_page_break()
    document.add_heading("4. Recentrage sur les inondations observées ou documentées", level=1)
    add_body(
        document,
        "La méthode développée cette année privilégie des sources décrivant le phénomène d’inondation lui-même. Les produits du JRC apportent une observation satellitaire de l’emprise en eau ; GASPAR renseigne la reconnaissance administrative des catastrophes naturelles en France ; HANZE recense des événements historiques dommageables et leurs impacts à l’échelle européenne. Leur combinaison permet de distinguer trois dimensions complémentaires : la preuve physique, la reconnaissance institutionnelle et la documentation historique.",
    )
    helpers.add_table(
        document,
        ["Dimension", "Travaux antérieurs", "Amélioration apportée cette année"],
        [
            ("Périmètre", "Approche principalement française", "Extension européenne et référentiel NUTS"),
            ("Phénomène", "Précipitations extrêmes utilisées comme proxy", "Inondations observées ou documentées"),
            ("Sources", "TRI, GASPAR et Copernicus", "JRC, GASPAR et HANZE, complétés par les zonages"),
            ("Temporalité", "Indicateurs annuels binaires", "Événements datés et fenêtres d’appariement"),
            ("Interprétation", "Signal potentiel de risque", "Hiérarchie explicite du niveau de preuve"),
        ],
        [1700, 3200, 4460],
        font_size=8.7,
    )
    helpers.add_caption(document, "Tableau 1 — Principales évolutions de la méthode.")
    add_body(
        document,
        "L’abandon des variables de précipitation ne signifie pas qu’elles sont inutiles. Elles restent pertinentes pour expliquer les mécanismes, anticiper les événements ou construire des scénarios climatiques. En revanche, elles ne sont plus retenues comme variable principale d’occurrence dans la chaîne de DataCollection, car le besoin est d’identifier des inondations effectivement observées ou documentées au voisinage des établissements. Ce changement réduit les faux positifs associés aux pluies extrêmes sans inondation, mais il introduit d’autres limites : lacunes satellitaires, hétérogénéité des déclarations administratives et sélection des événements dommageables dans les bases historiques.",
    )

    document.add_heading("5. Apports et limites de la nouvelle approche", level=1)
    add_body(
        document,
        "La principale amélioration réside dans la validité conceptuelle de l’indicateur. La variable finale se rapproche davantage de l’objet que l’on souhaite mesurer : l’occurrence d’une inondation et non la seule présence d’un facteur météorologique favorable. L’approche devient également reproductible à l’échelle européenne grâce à un référentiel spatial commun, et plus robuste grâce à la confrontation de sources indépendantes.",
    )
    add_body(
        document,
        "Cette robustesse ne doit toutefois pas être assimilée à une exhaustivité parfaite. Le JRC dépend de la disponibilité et de la qualité des observations satellitaires ; GASPAR reflète un processus de reconnaissance administrative et non un inventaire scientifique homogène ; HANZE privilégie les événements documentés et dommageables. De plus, l’appariement par NUTS 3 et par fenêtre temporelle fournit une compatibilité, mais ne démontre pas à lui seul l’identité physique de deux enregistrements. Les résultats doivent donc être accompagnés d’un niveau de confiance fondé sur la convergence des sources.",
    )
    helpers.add_callout(
        document,
        "Amélioration centrale",
        "La méthode passe d’un indicateur de conditions favorables à une inondation à un indicateur fondé sur des preuves complémentaires de l’événement, harmonisées à l’échelle européenne.",
    )

    document.add_heading("6. Synthèse", level=1)
    add_body(
        document,
        "Les travaux de l’année précédente ont fourni une base utile en rapprochant exposition territoriale, reconnaissance CatNat et précipitations extrêmes. La mission actuelle prolonge ce travail tout en corrigeant deux limites : le périmètre national et l’assimilation implicite entre pluie extrême et inondation. Le passage à l’échelle européenne, l’intégration de JRC et de HANZE ainsi que la comparaison avec GASPAR permettent désormais de construire une mesure plus cohérente avec le phénomène étudié.",
    )
    add_body(
        document,
        "Le choix de ne plus utiliser les précipitations comme indicateur principal est donc méthodologique plutôt que pratique. Il repose sur le constat qu’une inondation résulte d’une combinaison de facteurs : précipitation, humidité des sols, stockage souterrain, neige, propriétés du bassin, urbanisation, état des cours d’eau et dispositifs de protection. La pluie demeure un déterminant essentiel, mais elle ne suffit pas à conclure à l’existence d’une inondation. Pour une analyse du risque de crédit, il est préférable de privilégier l’occurrence observée ou documentée, puis d’utiliser les variables climatiques comme facteurs explicatifs ou prospectifs.",
    )

    document.add_heading("Références bibliographiques", level=1)
    sources = [
        "Berghuijs, W. R. et al. (2019). Growing Spatial Scales of Synchronous River Flooding in Europe. Geophysical Research Letters, 46. https://doi.org/10.1029/2018GL081883",
        "Bertola, M. et al. (2021). Do small and large floods have the same drivers of change? A regional attribution analysis in Europe. Hydrology and Earth System Sciences, 25, 1347–1364. https://doi.org/10.5194/hess-25-1347-2021",
        "Paprotny, D. et al. (2024). HANZE v2.1: an improved database of flood impacts in Europe from 1870 to 2020. Earth System Science Data, 16, 5145–5170. https://doi.org/10.5194/essd-16-5145-2024",
        "Tellman, B. et al. (2021). Satellite imaging reveals increased proportion of population exposed to floods. Nature, 596, 80–86. https://doi.org/10.1038/s41586-021-03695-w",
        "Thober, S. et al. (2023). Shifts in flood generation processes exacerbate regional flood anomalies in Europe. Communications Earth & Environment, 4, 49. https://doi.org/10.1038/s43247-023-00714-8",
        "Joint Research Centre (2024). Global Flood Monitoring – Annual Product and Service Quality Assessment Report 2023. Publications Office of the European Union, JRC138200.",
    ]
    for source in sources:
        add_source(document, source)

    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True

    document.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    build_document()
