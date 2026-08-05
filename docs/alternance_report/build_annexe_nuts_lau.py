from __future__ import annotations

import hashlib
import io
import math
import sqlite3
import struct
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "alternance_report" / "annexe_nuts_lau"
NUTS_GPKG = ROOT / "data" / "raw" / "NUTS_RG_03M_2024_4326.gpkg"
LAU_GPKG = ROOT / "data" / "raw" / "LAU_RG_01M_2024_4326.gpkg"
DOCX_OUT = OUT_DIR / "annexe_nuts_lau_france_italie.docx"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
INK = "1F2937"


def gpkg_wkb(blob: bytes) -> bytes:
    if blob[:2] != b"GP":
        return blob
    flags = blob[3]
    envelope_code = (flags >> 1) & 0b111
    envelope_doubles = {0: 0, 1: 4, 2: 6, 3: 6, 4: 8}.get(envelope_code, 0)
    return blob[8 + envelope_doubles * 8 :]


def read_wkb(data: bytes):
    stream = io.BytesIO(data)

    def u32(endian):
        return struct.unpack(endian + "I", stream.read(4))[0]

    def f64(endian):
        return struct.unpack(endian + "d", stream.read(8))[0]

    def geometry():
        byte_order = stream.read(1)
        if not byte_order:
            return []
        endian = "<" if byte_order == b"\x01" else ">"
        raw_type = u32(endian)
        has_z = bool(raw_type & 0x80000000) or 1000 <= raw_type % 10000 < 2000
        has_m = bool(raw_type & 0x40000000) or 2000 <= raw_type % 10000 < 4000
        geom_type = raw_type & 0x0FFFFFFF
        geom_type %= 1000
        dims = 2 + int(has_z) + int(has_m)

        def point():
            values = [f64(endian) for _ in range(dims)]
            return values[0], values[1]

        if geom_type == 3:  # Polygon
            rings = []
            for _ in range(u32(endian)):
                rings.append([point() for _ in range(u32(endian))])
            return [rings]
        if geom_type == 6:  # MultiPolygon
            polygons = []
            for _ in range(u32(endian)):
                polygons.extend(geometry())
            return polygons
        raise ValueError(f"Unsupported WKB geometry type: {geom_type}")

    return geometry()


def load_features(path: Path, country: str, level: int | None):
    connection = sqlite3.connect(path)
    table = connection.execute("select table_name from gpkg_contents").fetchone()[0]
    if level is None:
        sql = f'SELECT GISCO_ID, LAU_NAME, geom FROM "{table}" WHERE CNTR_CODE = ?'
        rows = connection.execute(sql, (country,)).fetchall()
    else:
        sql = f'SELECT NUTS_ID, NAME_LATN, geom FROM "{table}" WHERE CNTR_CODE = ? AND LEVL_CODE = ?'
        rows = connection.execute(sql, (country, level)).fetchall()
    connection.close()
    features = []
    for code, name, blob in rows:
        try:
            polygons = read_wkb(gpkg_wkb(blob))
        except (ValueError, struct.error):
            continue
        features.append({"code": code, "name": name, "polygons": polygons})
    return features


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def polygon_bounds(polygons):
    points = [p for polygon in polygons for ring in polygon for p in ring]
    if not points:
        return None
    xs, ys = zip(*points)
    return min(xs), min(ys), max(xs), max(ys)


def stable_fill(code: str, level: int | None):
    palette = [
        (204, 224, 244),
        (220, 233, 246),
        (189, 215, 238),
        (230, 240, 249),
        (211, 228, 243),
        (197, 221, 239),
    ]
    if level is None:
        return (225, 235, 244)
    digest = hashlib.sha1(code.encode("utf-8")).digest()[0]
    return palette[digest % len(palette)]


def draw_panel(draw, box, features, bbox, title, level, show_labels):
    x0, y0, x1, y1 = box
    draw.rounded_rectangle(box, radius=18, fill=(250, 252, 254), outline=(205, 213, 222), width=2)
    draw.text((x0 + 24, y0 + 18), title, font=font(28, True), fill=(23, 54, 93))
    map_top = y0 + 64
    pad = 26
    lon0, lon1, lat0, lat1 = bbox
    width, height = x1 - x0 - 2 * pad, y1 - map_top - pad
    scale = min(width / (lon1 - lon0), height / (lat1 - lat0))
    used_w, used_h = (lon1 - lon0) * scale, (lat1 - lat0) * scale
    ox = x0 + pad + (width - used_w) / 2
    oy = map_top + (height - used_h) / 2

    def project(point):
        lon, lat = point
        return ox + (lon - lon0) * scale, oy + (lat1 - lat) * scale

    visible = []
    for feature in features:
        bounds = polygon_bounds(feature["polygons"])
        if not bounds:
            continue
        bx0, by0, bx1, by1 = bounds
        if bx1 < lon0 or bx0 > lon1 or by1 < lat0 or by0 > lat1:
            continue
        visible.append((feature, bounds))
        for polygon in feature["polygons"]:
            for ring_index, ring in enumerate(polygon):
                pts = [project(p) for p in ring if lon0 - 2 <= p[0] <= lon1 + 2 and lat0 - 2 <= p[1] <= lat1 + 2]
                if len(pts) < 3:
                    continue
                if ring_index == 0:
                    outline = (74, 103, 134) if level is not None else (120, 142, 164)
                    line_width = 2 if level in (1, 2) else 1
                    draw.polygon(pts, fill=stable_fill(feature["code"], level), outline=outline, width=line_width)
                else:
                    draw.polygon(pts, fill=(250, 252, 254))

    if show_labels:
        for feature, bounds in visible:
            bx0, by0, bx1, by1 = bounds
            cx, cy = project(((bx0 + bx1) / 2, (by0 + by1) / 2))
            label = feature["code"]
            label_font = font(15 if level == 1 else 11, True)
            tb = draw.textbbox((0, 0), label, font=label_font)
            draw.rounded_rectangle(
                (cx - (tb[2] - tb[0]) / 2 - 3, cy - 8, cx + (tb[2] - tb[0]) / 2 + 3, cy + 9),
                radius=3,
                fill=(255, 255, 255, 220),
            )
            draw.text((cx, cy), label, font=label_font, anchor="mm", fill=(23, 54, 93))
    return len(visible)


def make_map(level: int | None, out_path: Path):
    image = Image.new("RGB", (1800, 1080), (244, 247, 250))
    draw = ImageDraw.Draw(image, "RGBA")
    level_name = "LAU" if level is None else f"NUTS {level}"
    draw.text((70, 42), f"Découpage territorial {level_name} — France et Italie", font=font(38, True), fill=(23, 54, 93))
    subtitle = "Référentiel Eurostat/GISCO 2024 — projection EPSG:4326"
    draw.text((72, 92), subtitle, font=font(22), fill=(102, 112, 133))

    source = LAU_GPKG if level is None else NUTS_GPKG
    fr = load_features(source, "FR", level)
    it = load_features(source, "IT", level)
    fr_visible = draw_panel(
        draw,
        (60, 145, 880, 945),
        fr,
        (-5.8, 10.0, 41.0, 51.6),
        "France métropolitaine",
        level,
        level in (1, 2),
    )
    it_visible = draw_panel(
        draw,
        (920, 145, 1740, 945),
        it,
        (6.2, 19.0, 35.3, 47.4),
        "Italie",
        level,
        level in (1, 2),
    )
    line1 = f"Unités représentées dans le cadre : France {fr_visible} | Italie {it_visible}"
    line2 = "Lecture : chaque contour correspond à une unité du niveau sélectionné ; les territoires ultramarins français sont hors cadre."
    draw.text((72, 978), line1, font=font(19, True), fill=(31, 41, 55))
    draw.text((72, 1013), line2, font=font(17), fill=(102, 112, 133))
    image.save(out_path, dpi=(220, 220))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_caption(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)


def add_body(document, text):
    p = document.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.keep_together = True
    return p


def build_document(map_paths):
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.left_indent = Inches(0)
        style.paragraph_format.right_indent = Inches(0)
        style.paragraph_format.first_line_indent = Inches(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "Rapport d’alternance — Annexe territoriale"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor.from_string(MID_GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Annexe NUTS–LAU | France–Italie")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("ANNEXE GÉOGRAPHIQUE")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("Niveaux NUTS 1, NUTS 2, NUTS 3 et LAU en France et en Italie")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(23)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(NAVY)

    add_body(
        document,
        "Cette annexe présente la hiérarchie territoriale utilisée pour harmoniser les données géographiques du projet DataCollection. La nomenclature NUTS, définie par Eurostat, organise les territoires européens en trois niveaux statistiques emboîtés. Les LAU constituent l’échelon administratif local situé sous le niveau NUTS 3. Les cartes suivantes montrent concrètement comment la précision géographique augmente du niveau NUTS 1 jusqu’au niveau LAU.",
    )
    add_body(
        document,
        "La correspondance avec les divisions administratives nationales doit être comprise comme une approximation fonctionnelle. Les niveaux NUTS sont des unités statistiques européennes : ils coïncident souvent avec des collectivités administratives, mais cette équivalence n’est pas identique dans tous les pays ni à toutes les dates. Dans le projet, ces identifiants servent avant tout de langage géographique commun pour rapprocher JRC, GASPAR, HANZE, TRI et ISPRA.",
    )

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Niveau", "Échelle", "France", "Italie"]
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = value
        set_cell_shading(cell, LIGHT_GRAY)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
    rows = [
        ("NUTS 1", "Macro-régionale", "Grands ensembles statistiques", "Groupes de régions"),
        ("NUTS 2", "Régionale", "Régions", "Régions et provinces autonomes"),
        ("NUTS 3", "Intermédiaire", "Départements", "Provinces et entités équivalentes"),
        ("LAU", "Locale", "Communes", "Comuni"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_widths(table, [1200, 1800, 3100, 3260])
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9)
    add_caption(document, "Tableau A.1 — Lecture simplifiée des niveaux territoriaux utilisés dans DataCollection.")

    sections = [
        (
            1,
            "A.1 — NUTS 1 : les grands ensembles territoriaux",
            "Le niveau NUTS 1 constitue le niveau infranational le plus agrégé de la nomenclature. Il est adapté aux comparaisons macro-régionales et à la présentation de tendances générales. En France, il regroupe les régions au sein de grands ensembles statistiques ; en Italie, il rassemble également plusieurs régions. Ce niveau est utile pour synthétiser les résultats, mais il est trop large pour localiser précisément une exposition ou un sinistre.",
        ),
        (
            2,
            "A.2 — NUTS 2 : l’échelle régionale",
            "Le niveau NUTS 2 correspond à l’échelle régionale privilégiée par de nombreuses statistiques européennes. En France, il est proche du découpage des régions ; en Italie, il correspond aux régions et aux provinces autonomes. Dans DataCollection, cette maille facilite les comparaisons régionales de fréquence, de couverture et de concentration des événements d’inondation.",
        ),
        (
            3,
            "A.3 — NUTS 3 : l’échelle départementale ou provinciale",
            "Le niveau NUTS 3 fournit une lecture territoriale plus fine. En France, il correspond principalement aux départements ; en Italie, il recouvre les provinces, villes métropolitaines et unités équivalentes. Ce niveau joue un rôle important dans l’intégration de HANZE et dans les comparaisons JRC–GASPAR, car il réduit la fragmentation observée lorsqu’une correspondance stricte à l’échelle communale est insuffisante.",
        ),
        (
            None,
            "A.4 — LAU : l’unité administrative locale",
            "La LAU est le niveau administratif local situé sous NUTS 3. Elle correspond généralement à la commune en France et au comune en Italie. Dans DataCollection, il s’agit de l’échelle privilégiée pour rattacher un point d’exposition, identifier les communes intersectées par un raster JRC et effectuer les rapprochements les plus détaillés. La densité des contours sur la carte illustre la forte granularité de ce référentiel.",
        ),
    ]
    for index, (level, heading, paragraph) in enumerate(sections):
        document.add_section(WD_SECTION.NEW_PAGE)
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(2)
        spacer.add_run(" ").font.size = Pt(1)
        document.add_heading(heading, level=1)
        add_body(document, paragraph)
        image_path = map_paths[level]
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(image_path), width=Inches(6.7))
        label = "LAU" if level is None else f"NUTS {level}"
        add_caption(document, f"Figure A.{index + 1} — Découpage {label} de la France métropolitaine et de l’Italie (Eurostat/GISCO, millésime 2024).")
        note = document.add_paragraph()
        note.paragraph_format.space_before = Pt(4)
        note.paragraph_format.space_after = Pt(0)
        note_run = note.add_run(
            "Note de lecture — Les territoires ultramarins français sont inclus dans le référentiel source, mais ne sont pas représentés dans le cadrage cartographique afin de conserver une échelle lisible pour la France métropolitaine et l’Italie."
        )
        note_run.font.size = Pt(8.5)
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor.from_string(MID_GRAY)

    document.add_section(WD_SECTION.NEW_PAGE)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)
    spacer.add_run(" ").font.size = Pt(1)
    document.add_heading("A.5 — Utilisation dans la chaîne DataCollection", level=1)
    add_body(
        document,
        "La hiérarchie NUTS–LAU permet de conserver plusieurs niveaux de lecture au sein d’une même chaîne de données. Un point de portefeuille est d’abord localisé à l’échelle LAU lorsque ses coordonnées le permettent. Les résultats peuvent ensuite être agrégés vers NUTS 3, NUTS 2 et NUTS 1 sans perdre la traçabilité de l’unité locale d’origine. Cette organisation rend possibles des analyses allant de la commune jusqu’aux grands ensembles européens.",
    )
    add_body(
        document,
        "Cette structure facilite également la combinaison de sources de granularités différentes. Les rasters JRC peuvent être tabularisés à l’échelle LAU, GASPAR est rapproché par code INSEE en France, tandis que HANZE est fréquemment mobilisé au niveau NUTS 3. Les couches TRI françaises et ISPRA italiennes apportent enfin une information spatiale nationale plus détaillée. L’utilisation d’identifiants territoriaux harmonisés permet de réunir ces informations dans les exports FLOOD_LGD tout en conservant leur source et leur niveau de précision.",
    )
    add_body(
        document,
        "Il convient néanmoins de ne pas interpréter une agrégation territoriale comme une preuve directe d’exposition. Une correspondance NUTS 3 indique qu’un événement concerne la même unité territoriale qu’un point étudié, mais elle ne démontre pas que ce point a été inondé. Lorsque les données le permettent, la confirmation finale repose donc sur une intersection spatiale plus précise, notamment avec les pixels des rasters JRC ou avec les zones réglementaires disponibles.",
    )
    source = document.add_paragraph()
    source.paragraph_format.space_before = Pt(12)
    source.paragraph_format.space_after = Pt(0)
    run = source.add_run("Source cartographique : Eurostat/GISCO, référentiels NUTS 2024 et LAU 2024, projection WGS 84 (EPSG:4326).")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)

    # Reapply identical page geometry to every section. Word otherwise inherits
    # inconsistent section properties on alternating pages after NEW_PAGE breaks.
    for current in document.sections:
        current.page_width = Inches(8.5)
        current.page_height = Inches(11)
        current.top_margin = Inches(0.8)
        current.bottom_margin = Inches(0.75)
        current.left_margin = Inches(0.85)
        current.right_margin = Inches(0.85)
        current.header_distance = Inches(0.492)
        current.footer_distance = Inches(0.492)
        current.different_first_page_header_footer = False
        current.header.is_linked_to_previous = True
        current.footer.is_linked_to_previous = True

    document.save(DOCX_OUT)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    map_paths = {}
    for level in (1, 2, 3, None):
        label = "lau" if level is None else f"nuts{level}"
        path = OUT_DIR / f"carte_{label}_france_italie.png"
        make_map(level, path)
        map_paths[level] = path
    build_document(map_paths)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
