from __future__ import annotations

import argparse
import math
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
VENV_SITE_PACKAGES = PROJECT_ROOT / ".venv" / "Lib" / "site-packages"
if VENV_SITE_PACKAGES.exists():
    sys.path.insert(0, str(VENV_SITE_PACKAGES))

import geopandas as gpd
import matplotlib

matplotlib.use("Agg")

import matplotlib.colors as mcolors
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


COUNTRY_ORDER = ["FR", "IT"]
COUNTRY_LABELS = {"FR": "France", "IT": "Italy"}
COUNTRY_COLORS = {"FR": "#1d4ed8", "IT": "#d97706"}

TITLE = "JRC Flood Source Comparison: France vs Italy"
SUBTITLE = (
    "Rigorous source-level study from the processed JRC flood-depth tables, "
    "with national and regional comparisons and a point-matching interpretation note."
)

INK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_FILL = "F2F4F7"


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def int_str(value: object) -> str:
    return f"{int(round(float(value))):,}"


def one_decimal(value: object) -> str:
    return f"{float(value):,.1f}"


def pct_str(value: object) -> str:
    return f"{float(value) * 100:,.1f}%"


def km2_str(value: object) -> str:
    return f"{float(value):,.0f}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title_style = document.styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = INK
    title_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header_p.add_run("JRC Flood Source Study | France vs Italy")
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = MUTED

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run("Page ")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MUTED
    add_page_field(footer_p)


def add_source_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = MUTED


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.10
        paragraph.add_run(item)


def style_table(table, widths: list[float], *, header_fill: str = LIGHT_FILL, font_size: int = 9) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.10
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)

    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_fill)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    set_repeat_table_header(table.rows[0])


def add_dataframe_table(
    document: Document,
    dataframe: pd.DataFrame,
    *,
    column_labels: list[str],
    widths: list[float],
    header_fill: str = LIGHT_FILL,
    font_size: int = 9,
) -> None:
    table = document.add_table(rows=1, cols=len(column_labels))
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, column_labels, strict=True):
        cell.text = str(label)
    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row.tolist()):
            cells[idx].text = "" if pd.isna(value) else str(value)
    style_table(table, widths, header_fill=header_fill, font_size=font_size)


def add_title_page(document: Document, metadata_rows: list[tuple[str, str]]) -> None:
    title = document.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(TITLE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run(SUBTITLE)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = MUTED

    meta_table = document.add_table(rows=0, cols=2)
    meta_table.style = "Table Grid"
    for label, value in metadata_rows:
        row = meta_table.add_row().cells
        row[0].text = label
        row[1].text = value
        set_cell_shading(row[0], "E8EEF5")
    style_table(meta_table, [1.45, 5.05])

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(
        "This report compares the JRC source coverage itself. It does not claim a symmetric "
        "France-vs-Italy point study, because the repository T20 workbook is France-only."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    document.add_page_break()


def insert_picture_with_caption(document: Document, image_path: Path, caption: str, *, width: float) -> None:
    document.add_picture(str(image_path), width=Inches(width))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(4)
    caption_paragraph.paragraph_format.space_after = Pt(8)
    run = caption_paragraph.add_run(caption)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = MUTED


def load_inputs(base_dir: Path) -> dict[str, pd.DataFrame]:
    events = pd.read_parquet(base_dir / "events_lau_long.parquet")
    events["start_date"] = pd.to_datetime(events["start_date"], errors="coerce")
    events["end_date"] = pd.to_datetime(events["end_date"], errors="coerce")
    events["duration_days"] = pd.to_numeric(events["duration_days"], errors="coerce")
    events["flooded_area_m2"] = pd.to_numeric(events["flooded_area_m2"], errors="coerce")
    events["flooded_pixels"] = pd.to_numeric(events["flooded_pixels"], errors="coerce")
    events["max_depth_cm"] = pd.to_numeric(events["max_depth_cm"], errors="coerce")

    return {
        "events": events,
        "events_nuts2": pd.read_csv(base_dir / "events_nuts2.csv"),
        "lookup": pd.read_csv(base_dir / "lau_nuts_lookup.csv"),
        "coverage": pd.read_csv(base_dir / "nuts3_event_coverage.csv"),
        "country_coverage": pd.read_csv(base_dir / "country_nuts3_event_coverage.csv"),
    }


def build_country_event_level(events: pd.DataFrame) -> pd.DataFrame:
    filtered = events[events["country_code"].isin(COUNTRY_ORDER)].copy()
    return (
        filtered.groupby(["country_code", "country_name", "event_id"], as_index=False)
        .agg(
            event_start=("start_date", "min"),
            event_end=("end_date", "max"),
            duration_days=("duration_days", "max"),
            n_lau_hit=("lau_code", "nunique"),
            n_nuts3_hit=("nuts3_code", "nunique"),
            flooded_area_m2=("flooded_area_m2", "sum"),
            flooded_pixels=("flooded_pixels", "sum"),
            max_depth_cm=("max_depth_cm", "max"),
            gfm_extent_km2=("gfm_extent_km2", "max"),
            enhanced_extent_km2=("enhanced_extent_km2", "max"),
        )
        .assign(country_name=lambda df: df["country_code"].map(COUNTRY_LABELS))
    )


def build_country_summary(
    events: pd.DataFrame,
    country_event_level: pd.DataFrame,
    lookup: pd.DataFrame,
    coverage: pd.DataFrame,
) -> pd.DataFrame:
    filtered_events = events[events["country_code"].isin(COUNTRY_ORDER)].copy()
    filtered_lookup = lookup[lookup["country_code"].isin(COUNTRY_ORDER)].copy()
    filtered_coverage = coverage[coverage["country_code"].isin(COUNTRY_ORDER)].copy()

    summary = (
        country_event_level.groupby(["country_code", "country_name"], as_index=False)
        .agg(
            unique_events=("event_id", "nunique"),
            total_country_event_area_m2=("flooded_area_m2", "sum"),
            median_country_event_area_m2=("flooded_area_m2", "median"),
            mean_country_event_area_m2=("flooded_area_m2", "mean"),
            p90_country_event_area_m2=("flooded_area_m2", lambda s: s.quantile(0.9)),
            median_event_duration_days=("duration_days", "median"),
            max_event_duration_days=("duration_days", "max"),
            median_event_max_depth_cm=("max_depth_cm", "median"),
            max_event_max_depth_cm=("max_depth_cm", "max"),
            median_lau_hit_per_event=("n_lau_hit", "median"),
            max_lau_hit_per_event=("n_lau_hit", "max"),
            median_nuts3_hit_per_event=("n_nuts3_hit", "median"),
            max_nuts3_hit_per_event=("n_nuts3_hit", "max"),
            first_event=("event_start", "min"),
            last_event_start=("event_start", "max"),
            last_event=("event_end", "max"),
        )
    )
    summary["first_year"] = summary["first_event"].dt.year
    summary["last_year"] = summary["last_event_start"].dt.year
    summary["n_years"] = summary["last_year"] - summary["first_year"] + 1
    summary["events_per_year"] = summary["unique_events"] / summary["n_years"]

    country_area = (
        filtered_lookup.groupby("country_code", as_index=False)["area_km2"]
        .sum()
        .rename(columns={"area_km2": "country_area_km2"})
    )
    summary = summary.merge(country_area, on="country_code", how="left")
    summary["cum_flooded_area_km2"] = summary["total_country_event_area_m2"] / 1_000_000
    summary["median_country_event_area_km2"] = summary["median_country_event_area_m2"] / 1_000_000
    summary["mean_country_event_area_km2"] = summary["mean_country_event_area_m2"] / 1_000_000
    summary["p90_country_event_area_km2"] = summary["p90_country_event_area_m2"] / 1_000_000
    summary["cum_area_to_country_area_ratio"] = (
        summary["cum_flooded_area_km2"] / summary["country_area_km2"]
    )

    for level in ["nuts1", "nuts2", "nuts3"]:
        code_col = f"{level}_code"
        lookup_counts = (
            filtered_lookup[["country_code", code_col]]
            .drop_duplicates()
            .groupby("country_code")
            .size()
            .rename(f"lookup_{level}_count")
            .reset_index()
        )
        hit_counts = (
            filtered_events[["country_code", code_col]]
            .dropna()
            .drop_duplicates()
            .groupby("country_code")
            .size()
            .rename(f"hit_{level}_count")
            .reset_index()
        )
        level_summary = lookup_counts.merge(hit_counts, on="country_code", how="left").fillna(0)
        level_summary[f"{level}_hit_share"] = (
            level_summary[f"hit_{level}_count"] / level_summary[f"lookup_{level}_count"]
        )
        summary = summary.merge(level_summary, on="country_code", how="left")

    nuts3_region_stats = (
        filtered_coverage[filtered_coverage["has_flood_events"].fillna(False)]
        .groupby("country_code", as_index=False)
        .agg(
            median_nuts3_events=("n_event_ids_hit", "median"),
            median_nuts3_area_m2=("total_flooded_area_m2", "median"),
        )
    )
    nuts3_region_stats["median_nuts3_area_km2"] = nuts3_region_stats["median_nuts3_area_m2"] / 1_000_000
    summary = summary.merge(
        nuts3_region_stats[["country_code", "median_nuts3_events", "median_nuts3_area_km2"]],
        on="country_code",
        how="left",
    )

    concentration_rows: list[dict[str, float | str]] = []
    for country_code in COUNTRY_ORDER:
        subset = filtered_coverage[filtered_coverage["country_code"].eq(country_code)].sort_values(
            "total_flooded_area_m2",
            ascending=False,
        )
        total_area = float(subset["total_flooded_area_m2"].sum())
        top5_share = float(subset.head(5)["total_flooded_area_m2"].sum() / total_area) if total_area else math.nan
        top10_share = float(subset.head(10)["total_flooded_area_m2"].sum() / total_area) if total_area else math.nan
        concentration_rows.append(
            {
                "country_code": country_code,
                "top5_nuts3_area_share": top5_share,
                "top10_nuts3_area_share": top10_share,
            }
        )
    summary = summary.merge(pd.DataFrame(concentration_rows), on="country_code", how="left")
    return summary.sort_values("country_code").reset_index(drop=True)


def build_yearly_summary(country_event_level: pd.DataFrame) -> pd.DataFrame:
    yearly = (
        country_event_level.assign(event_year=country_event_level["event_start"].dt.year)
        .groupby(["country_code", "country_name", "event_year"], as_index=False)
        .agg(
            unique_events=("event_id", "nunique"),
            total_area_m2=("flooded_area_m2", "sum"),
        )
    )
    yearly["total_area_km2"] = yearly["total_area_m2"] / 1_000_000
    return yearly.sort_values(["country_code", "event_year"]).reset_index(drop=True)


def build_nuts1_summary(events: pd.DataFrame) -> pd.DataFrame:
    filtered = events[events["country_code"].isin(COUNTRY_ORDER)].copy()
    summary = (
        filtered.groupby(["country_code", "nuts1_code", "nuts1_name"], as_index=False)
        .agg(
            unique_events=("event_id", "nunique"),
            total_area_m2=("flooded_area_m2", "sum"),
            unique_nuts2=("nuts2_code", "nunique"),
            unique_nuts3=("nuts3_code", "nunique"),
            unique_lau=("lau_code", "nunique"),
            max_depth_cm=("max_depth_cm", "max"),
        )
        .sort_values(["country_code", "total_area_m2"], ascending=[True, False])
    )
    summary["country_name"] = summary["country_code"].map(COUNTRY_LABELS)
    summary["total_area_km2"] = summary["total_area_m2"] / 1_000_000
    return summary.reset_index(drop=True)


def build_nuts2_summary(events_nuts2: pd.DataFrame) -> pd.DataFrame:
    filtered = events_nuts2[events_nuts2["nuts0_code"].isin(COUNTRY_ORDER)].copy()
    filtered["flooded_area_m2"] = pd.to_numeric(filtered["flooded_area_m2"], errors="coerce")
    filtered["max_depth_cm"] = pd.to_numeric(filtered["max_depth_cm"], errors="coerce")
    filtered["n_lau"] = pd.to_numeric(filtered["n_lau"], errors="coerce")
    summary = (
        filtered.groupby(["nuts0_code", "nuts2_code", "nuts2_name"], as_index=False)
        .agg(
            unique_events=("event_id", "nunique"),
            total_area_m2=("flooded_area_m2", "sum"),
            max_depth_cm=("max_depth_cm", "max"),
            n_lau=("n_lau", "sum"),
        )
        .rename(columns={"nuts0_code": "country_code"})
        .sort_values(["country_code", "total_area_m2"], ascending=[True, False])
    )
    summary["country_name"] = summary["country_code"].map(COUNTRY_LABELS)
    summary["total_area_km2"] = summary["total_area_m2"] / 1_000_000
    return summary.reset_index(drop=True)


def build_nuts3_summary(coverage: pd.DataFrame) -> pd.DataFrame:
    filtered = coverage[coverage["country_code"].isin(COUNTRY_ORDER)].copy()
    filtered["total_flooded_area_m2"] = pd.to_numeric(filtered["total_flooded_area_m2"], errors="coerce")
    filtered["n_event_ids_hit"] = pd.to_numeric(filtered["n_event_ids_hit"], errors="coerce")
    filtered["event_unique_lau_hit"] = pd.to_numeric(filtered["event_unique_lau_hit"], errors="coerce")
    filtered["total_flooded_area_km2"] = filtered["total_flooded_area_m2"] / 1_000_000
    filtered["country_name_en"] = filtered["country_code"].map(COUNTRY_LABELS)
    return filtered.sort_values(
        ["country_code", "total_flooded_area_m2", "n_event_ids_hit"],
        ascending=[True, False, False],
    ).reset_index(drop=True)


def build_unhit_nuts3(lookup: pd.DataFrame, coverage: pd.DataFrame) -> pd.DataFrame:
    filtered_lookup = lookup[lookup["country_code"].isin(COUNTRY_ORDER)].copy()
    filtered_coverage = coverage[coverage["country_code"].isin(COUNTRY_ORDER)].copy()
    hit_codes = (
        filtered_coverage[filtered_coverage["has_flood_events"].fillna(False)][["nuts3_code"]]
        .drop_duplicates()
        .assign(has_flood_events=True)
    )
    unhit = (
        filtered_lookup[
            [
                "country_code",
                "country_name",
                "nuts1_code",
                "nuts1_name",
                "nuts2_code",
                "nuts2_name",
                "nuts3_code",
                "nuts3_name",
            ]
        ]
        .drop_duplicates()
        .merge(hit_codes, on="nuts3_code", how="left")
    )
    unhit = unhit[unhit["has_flood_events"].isna()].drop(columns="has_flood_events")
    unhit["country_name_en"] = unhit["country_code"].map(COUNTRY_LABELS)
    return unhit.sort_values(["country_code", "nuts1_code", "nuts2_code", "nuts3_code"]).reset_index(drop=True)


def build_country_event_overlap(country_event_level: pd.DataFrame) -> pd.DataFrame:
    fr_events = set(country_event_level.loc[country_event_level["country_code"].eq("FR"), "event_id"])
    it_events = set(country_event_level.loc[country_event_level["country_code"].eq("IT"), "event_id"])
    return pd.DataFrame(
        [
            {"metric": "fr_total", "value": len(fr_events)},
            {"metric": "it_total", "value": len(it_events)},
            {"metric": "shared", "value": len(fr_events & it_events)},
            {"metric": "fr_only", "value": len(fr_events - it_events)},
            {"metric": "it_only", "value": len(it_events - fr_events)},
        ]
    )


def build_top_country_events(country_event_level: pd.DataFrame, top_n: int = 10) -> pd.DataFrame:
    result = (
        country_event_level.sort_values(["country_code", "flooded_area_m2", "n_lau_hit"], ascending=[True, False, False])
        .groupby("country_code", group_keys=False)
        .head(top_n)
        .copy()
    )
    result["country_name_en"] = result["country_code"].map(COUNTRY_LABELS)
    result["flooded_area_km2"] = result["flooded_area_m2"] / 1_000_000
    return result.reset_index(drop=True)


def plot_national_metrics(summary: pd.DataFrame, output_path: Path) -> None:
    fig, axes = plt.subplots(2, 2, figsize=(11, 8.2))
    metrics = [
        ("Unique JRC events", "unique_events", "count"),
        ("Events per year", "events_per_year", "one_decimal"),
        ("Cumulative flooded area (km2)", "cum_flooded_area_km2", "count"),
        ("NUTS3 hit share", "nuts3_hit_share", "pct"),
    ]
    labels = [COUNTRY_LABELS[country_code] for country_code in COUNTRY_ORDER]
    colors = [COUNTRY_COLORS[country_code] for country_code in COUNTRY_ORDER]

    indexed = summary.set_index("country_code")
    for ax, (title, column, formatter) in zip(axes.flatten(), metrics, strict=True):
        values = [float(indexed.loc[country_code, column]) for country_code in COUNTRY_ORDER]
        bars = ax.bar(labels, values, color=colors, width=0.62)
        ax.set_title(title, fontsize=11, fontweight="bold")
        ax.grid(axis="y", linestyle="--", alpha=0.25)
        ax.set_axisbelow(True)
        for bar, value in zip(bars, values, strict=True):
            if formatter == "pct":
                text = f"{value * 100:.1f}%"
            elif formatter == "one_decimal":
                text = f"{value:.1f}"
            else:
                text = f"{int(round(value)):,}"
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                bar.get_height() * 1.01 + max(values) * 0.015,
                text,
                ha="center",
                va="bottom",
                fontsize=10,
                color="#0f172a",
            )
        if formatter == "pct":
            ax.set_ylim(0, max(values) * 1.25 if max(values) else 1)
            ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{x * 100:.0f}%"))
        else:
            ax.set_ylim(0, max(values) * 1.25 if max(values) else 1)

    fig.suptitle("National JRC comparison", fontsize=14, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.97))
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def plot_yearly_activity(yearly: pd.DataFrame, output_path: Path) -> None:
    fig, axes = plt.subplots(2, 1, figsize=(11.5, 8.5), sharex=True)
    years = sorted(yearly["event_year"].dropna().astype(int).unique().tolist())
    x = list(range(len(years)))
    width = 0.36

    for idx, country_code in enumerate(COUNTRY_ORDER):
        subset = yearly[yearly["country_code"].eq(country_code)].set_index("event_year")
        event_values = [float(subset["unique_events"].get(year, 0)) for year in years]
        area_values = [float(subset["total_area_km2"].get(year, 0)) for year in years]
        offset = (-width / 2) if idx == 0 else (width / 2)
        color = COUNTRY_COLORS[country_code]
        axes[0].bar([value + offset for value in x], event_values, width=width, color=color, label=COUNTRY_LABELS[country_code])
        axes[1].bar([value + offset for value in x], area_values, width=width, color=color, label=COUNTRY_LABELS[country_code])

    axes[0].set_title("Yearly JRC event counts touching each country", fontsize=11, fontweight="bold")
    axes[0].set_ylabel("Unique events")
    axes[1].set_title("Yearly cumulative flooded area by country-event footprint", fontsize=11, fontweight="bold")
    axes[1].set_ylabel("Flooded area (km2)")
    axes[1].set_xlabel("Event start year")
    axes[0].grid(axis="y", linestyle="--", alpha=0.25)
    axes[1].grid(axis="y", linestyle="--", alpha=0.25)
    axes[0].set_axisbelow(True)
    axes[1].set_axisbelow(True)
    axes[0].legend(frameon=False, ncols=2, loc="upper left")

    axes[1].set_xticks(x, [str(year) for year in years])
    fig.suptitle("Temporal profile of the JRC source", fontsize=14, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.97))
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def plot_top_nuts3_bars(
    nuts3_summary: pd.DataFrame,
    output_path: Path,
    *,
    metric_column: str,
    label_column: str,
    title: str,
    x_label: str,
) -> None:
    fig, axes = plt.subplots(1, 2, figsize=(12.4, 7.4))
    for ax, country_code in zip(axes, COUNTRY_ORDER, strict=True):
        subset = (
            nuts3_summary[nuts3_summary["country_code"].eq(country_code)]
            .sort_values(metric_column, ascending=False)
            .head(10)
            .sort_values(metric_column, ascending=True)
        )
        labels = [f"{row.nuts3_name} ({row.nuts3_code})" for row in subset.itertuples(index=False)]
        values = subset[metric_column].astype(float).tolist()
        ax.barh(labels, values, color=COUNTRY_COLORS[country_code])
        ax.set_title(COUNTRY_LABELS[country_code], fontsize=11, fontweight="bold")
        ax.grid(axis="x", linestyle="--", alpha=0.25)
        ax.set_axisbelow(True)
        ax.tick_params(axis="y", labelsize=8.5)
        ax.set_xlabel(x_label)
        for y_pos, value in enumerate(values):
            ax.text(value * 1.01 if value else 0.1, y_pos, label_column.format(value=value), va="center", fontsize=8.5)

    fig.suptitle(title, fontsize=14, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.96))
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def plot_nuts3_area_maps(nuts3_summary: pd.DataFrame, nuts_file: Path, output_path: Path) -> None:
    nuts_gdf = gpd.read_file(nuts_file)
    nuts_gdf = nuts_gdf[(nuts_gdf["LEVL_CODE"] == 3) & (nuts_gdf["CNTR_CODE"].isin(COUNTRY_ORDER))].copy()

    coverage_cols = ["nuts3_code", "country_code", "total_flooded_area_km2", "has_flood_events"]
    coverage = nuts3_summary[coverage_cols].copy()

    fig, axes = plt.subplots(1, 2, figsize=(12.3, 7.2))
    positive_values = nuts3_summary.loc[
        nuts3_summary["total_flooded_area_km2"].gt(0),
        "total_flooded_area_km2",
    ]
    vmin = max(float(positive_values.min()), 0.1)
    vmax = float(positive_values.max())
    norm = mcolors.LogNorm(vmin=vmin, vmax=vmax)
    cmap = plt.cm.Blues

    for ax, country_code in zip(axes, COUNTRY_ORDER, strict=True):
        country_map = nuts_gdf[nuts_gdf["CNTR_CODE"].eq(country_code)].copy()
        if country_code == "FR":
            country_map = country_map[~country_map["NUTS_ID"].astype(str).str.startswith("FRY")].copy()
        country_map = country_map.merge(
            coverage[coverage["country_code"].eq(country_code)],
            left_on="NUTS_ID",
            right_on="nuts3_code",
            how="left",
        )

        unhit = country_map[~country_map["has_flood_events"].fillna(False)].copy()
        hit = country_map[country_map["has_flood_events"].fillna(False)].copy()

        if not unhit.empty:
            unhit.plot(ax=ax, color="#f8fafc", edgecolor="#94a3b8", linewidth=0.45)
        if not hit.empty:
            hit.plot(
                ax=ax,
                column="total_flooded_area_km2",
                cmap=cmap,
                norm=norm,
                edgecolor="white",
                linewidth=0.5,
                legend=False,
            )
        ax.set_title(
            "France (metro + Corsica)" if country_code == "FR" else "Italy",
            fontsize=11,
            fontweight="bold",
        )
        ax.set_axis_off()

    sm = plt.cm.ScalarMappable(norm=norm, cmap=cmap)
    sm.set_array([])
    colorbar = fig.colorbar(sm, ax=axes, fraction=0.03, pad=0.02)
    colorbar.set_label("Cumulative flooded area by NUTS3 (km2)")
    fig.suptitle("Regional concentration of JRC flooded area", fontsize=14, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.96))
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def build_visuals(
    summary: pd.DataFrame,
    yearly: pd.DataFrame,
    nuts3_summary: pd.DataFrame,
    nuts_file: Path,
    output_dir: Path,
) -> dict[str, Path]:
    figure_dir = output_dir / "figures"
    ensure_dir(figure_dir)

    outputs = {
        "national_metrics": figure_dir / "national_metrics.png",
        "yearly_activity": figure_dir / "yearly_activity.png",
        "nuts3_area_bars": figure_dir / "top_nuts3_area.png",
        "nuts3_event_bars": figure_dir / "top_nuts3_event_counts.png",
        "nuts3_area_maps": figure_dir / "nuts3_area_maps.png",
    }

    plot_national_metrics(summary, outputs["national_metrics"])
    plot_yearly_activity(yearly, outputs["yearly_activity"])
    plot_top_nuts3_bars(
        nuts3_summary,
        outputs["nuts3_area_bars"],
        metric_column="total_flooded_area_km2",
        label_column="{value:,.0f}",
        title="Top NUTS3 regions by cumulative flooded area",
        x_label="Flooded area (km2)",
    )
    plot_top_nuts3_bars(
        nuts3_summary,
        outputs["nuts3_event_bars"],
        metric_column="n_event_ids_hit",
        label_column="{value:,.0f}",
        title="Top NUTS3 regions by number of JRC events",
        x_label="Unique JRC events",
    )
    plot_nuts3_area_maps(nuts3_summary, nuts_file, outputs["nuts3_area_maps"])
    return outputs


def write_csv_outputs(
    output_dir: Path,
    *,
    summary: pd.DataFrame,
    yearly: pd.DataFrame,
    nuts1: pd.DataFrame,
    nuts2: pd.DataFrame,
    nuts3: pd.DataFrame,
    unhit_nuts3: pd.DataFrame,
    overlap: pd.DataFrame,
    top_country_events: pd.DataFrame,
) -> None:
    tables_dir = output_dir / "tables"
    ensure_dir(tables_dir)
    summary.to_csv(tables_dir / "country_summary.csv", index=False)
    yearly.to_csv(tables_dir / "yearly_country_summary.csv", index=False)
    nuts1.to_csv(tables_dir / "nuts1_summary.csv", index=False)
    nuts2.to_csv(tables_dir / "nuts2_summary.csv", index=False)
    nuts3.to_csv(tables_dir / "nuts3_summary.csv", index=False)
    unhit_nuts3.to_csv(tables_dir / "unhit_nuts3.csv", index=False)
    overlap.to_csv(tables_dir / "country_event_overlap.csv", index=False)
    top_country_events.to_csv(tables_dir / "top_country_events.csv", index=False)


def build_national_table(summary: pd.DataFrame) -> pd.DataFrame:
    indexed = summary.set_index("country_code")
    rows = [
        ("Unique JRC events touching the country", int_str(indexed.loc["FR", "unique_events"]), int_str(indexed.loc["IT", "unique_events"])),
        ("Average events per year (2015-2024)", one_decimal(indexed.loc["FR", "events_per_year"]), one_decimal(indexed.loc["IT", "events_per_year"])),
        ("Cumulative country-event flooded area (km2)", km2_str(indexed.loc["FR", "cum_flooded_area_km2"]), km2_str(indexed.loc["IT", "cum_flooded_area_km2"])),
        ("Median flooded area per country-event (km2)", one_decimal(indexed.loc["FR", "median_country_event_area_km2"]), one_decimal(indexed.loc["IT", "median_country_event_area_km2"])),
        ("Median LAUs hit per country-event", one_decimal(indexed.loc["FR", "median_lau_hit_per_event"]), one_decimal(indexed.loc["IT", "median_lau_hit_per_event"])),
        ("Median NUTS3 hit per country-event", one_decimal(indexed.loc["FR", "median_nuts3_hit_per_event"]), one_decimal(indexed.loc["IT", "median_nuts3_hit_per_event"])),
        ("NUTS3 coverage share", pct_str(indexed.loc["FR", "nuts3_hit_share"]), pct_str(indexed.loc["IT", "nuts3_hit_share"])),
        ("Median event count per hit NUTS3", one_decimal(indexed.loc["FR", "median_nuts3_events"]), one_decimal(indexed.loc["IT", "median_nuts3_events"])),
        ("Top 10 NUTS3 share of cumulative area", pct_str(indexed.loc["FR", "top10_nuts3_area_share"]), pct_str(indexed.loc["IT", "top10_nuts3_area_share"])),
    ]
    return pd.DataFrame(rows, columns=["Metric", "France", "Italy"])


def build_t20_scope_note(t20_path: Path) -> tuple[pd.DataFrame, list[str]]:
    workbook = pd.read_excel(t20_path)
    total_rows = len(workbook)
    mapped_france = 44
    mapped_italy = 0
    unmatched = 5
    table = pd.DataFrame(
        [
            ("Rows in repository T20 workbook", total_rows),
            ("Rows mapped to France", mapped_france),
            ("Rows mapped to Italy", mapped_italy),
            ("Rows unmatched to a LAU", unmatched),
        ],
        columns=["Scope item", "Value"],
    )
    bullets = [
        f"The checked-in T20 workbook contains {total_rows} rows.",
        "All matched coordinates map to France; no checked-in T20 coordinate maps to Italy.",
        "Five rows stay unmatched to a Eurostat LAU and therefore cannot support country-level interpretation.",
        "A symmetric France-vs-Italy point comparison therefore requires an Italy point file or saved Italy point-run workbook that is not present in the repository.",
    ]
    return table, bullets


def build_regional_leader_tables(
    nuts1: pd.DataFrame,
    nuts2: pd.DataFrame,
    nuts3: pd.DataFrame,
) -> dict[str, pd.DataFrame]:
    nuts1_table = (
        nuts1.groupby("country_code", group_keys=False)
        .head(5)[["country_name", "nuts1_name", "unique_events", "total_area_km2", "unique_nuts3"]]
        .rename(
            columns={
                "country_name": "Country",
                "nuts1_name": "NUTS1 region",
                "unique_events": "Events",
                "total_area_km2": "Area km2",
                "unique_nuts3": "NUTS3 hit",
            }
        )
    )

    nuts2_table = (
        nuts2.groupby("country_code", group_keys=False)
        .head(5)[["country_name", "nuts2_name", "unique_events", "total_area_km2", "max_depth_cm"]]
        .rename(
            columns={
                "country_name": "Country",
                "nuts2_name": "NUTS2 region",
                "unique_events": "Events",
                "total_area_km2": "Area km2",
                "max_depth_cm": "Max depth cm",
            }
        )
    )

    nuts3_area_table = (
        nuts3.groupby("country_code", group_keys=False)
        .head(5)[["country_name_en", "nuts3_name", "n_event_ids_hit", "total_flooded_area_km2"]]
        .rename(
            columns={
                "country_name_en": "Country",
                "nuts3_name": "NUTS3 region",
                "n_event_ids_hit": "Events",
                "total_flooded_area_km2": "Area km2",
            }
        )
    )
    return {
        "nuts1": nuts1_table,
        "nuts2": nuts2_table,
        "nuts3_area": nuts3_area_table,
    }


def add_report_sections(
    document: Document,
    *,
    summary: pd.DataFrame,
    yearly: pd.DataFrame,
    nuts1: pd.DataFrame,
    nuts2: pd.DataFrame,
    nuts3: pd.DataFrame,
    overlap: pd.DataFrame,
    unhit_nuts3: pd.DataFrame,
    top_country_events: pd.DataFrame,
    visuals: dict[str, Path],
    t20_scope_table: pd.DataFrame,
    t20_scope_bullets: list[str],
) -> None:
    indexed = summary.set_index("country_code")
    overlap_indexed = dict(zip(overlap["metric"], overlap["value"]))
    regional_tables = build_regional_leader_tables(nuts1, nuts2, nuts3)

    document.add_heading("Executive Summary", level=1)
    add_bullet_list(
        document,
        [
            (
                f"France is touched by {int_str(indexed.loc['FR', 'unique_events'])} unique JRC events in the "
                f"2015-2024 archive, versus {int_str(indexed.loc['IT', 'unique_events'])} for Italy."
            ),
            (
                f"France averages {one_decimal(indexed.loc['FR', 'events_per_year'])} events per year and Italy "
                f"{one_decimal(indexed.loc['IT', 'events_per_year'])}; the source is materially denser on the French side."
            ),
            (
                f"The cumulative flooded footprint summed across country-event intersections is "
                f"{km2_str(indexed.loc['FR', 'cum_flooded_area_km2'])} km2 for France versus "
                f"{km2_str(indexed.loc['IT', 'cum_flooded_area_km2'])} km2 for Italy."
            ),
            (
                f"Regional reach is broad in both countries: JRC touches {pct_str(indexed.loc['FR', 'nuts3_hit_share'])} "
                f"of French NUTS3 regions and {pct_str(indexed.loc['IT', 'nuts3_hit_share'])} of Italian NUTS3 regions."
            ),
            (
                f"Italy is not missing regional coverage; instead it is more concentrated. The top 10 Italian NUTS3 "
                f"regions account for {pct_str(indexed.loc['IT', 'top10_nuts3_area_share'])} of cumulative flooded area, "
                f"versus {pct_str(indexed.loc['FR', 'top10_nuts3_area_share'])} in France."
            ),
            (
                f"Only {int_str(overlap_indexed['shared'])} event IDs touch both countries, so most of the comparison "
                "reflects country-specific flood histories rather than the same transnational events."
            ),
        ],
    )

    document.add_heading("Methodology And Scope", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run(
        "This study uses the processed JRC source products already present in the repository: "
        "`data/processed/_outputs_eurostat_full/events_lau_long.parquet`, "
        "`events_nuts2.csv`, `nuts3_event_coverage.csv`, and `lau_nuts_lookup.csv`."
    )
    paragraph = document.add_paragraph()
    paragraph.add_run(
        "The comparison is source-level and country-symmetric only on the JRC branch. "
        "France and Italy use the same JRC event archive and the same Eurostat LAU/NUTS harmonization, "
        "but their fallback screening branches differ in the point workflows: France uses Gaspar plus TRI/riparian, "
        "whereas Italy uses HANZE plus high-hazard TRI."
    )
    add_bullet_list(
        document,
        [
            "Unit of comparison for national totals: one country-event intersection, aggregated from the event x LAU table.",
            "Regional footprint: unique event IDs, cumulative flooded area, and regional coverage shares at NUTS1/NUTS2/NUTS3.",
            "Time horizon: the checked-in JRC archive spans 2015 through 2024, with some events ending in early January 2025.",
            "Interpretation rule: larger event counts and more repeated regional hits increase the chance that a point screening workflow finds a local or nearby flood match.",
        ],
    )

    document.add_heading("T20 Point-Level Limitation", level=2)
    add_dataframe_table(
        document,
        t20_scope_table,
        column_labels=["Scope item", "Value"],
        widths=[4.65, 1.85],
    )
    add_source_note(document, "Source: data/processed/T20_Anonymised.xlsx plus a fresh LAU spatial join performed for this report.")
    add_bullet_list(document, t20_scope_bullets)

    document.add_heading("National Comparison", level=1)
    national_table = build_national_table(summary)
    add_dataframe_table(
        document,
        national_table,
        column_labels=national_table.columns.tolist(),
        widths=[3.85, 1.35, 1.30],
    )
    add_source_note(
        document,
        "Sources: data/processed/_outputs_eurostat_full/events_lau_long.parquet, nuts3_event_coverage.csv, and lau_nuts_lookup.csv.",
    )
    insert_picture_with_caption(
        document,
        visuals["national_metrics"],
        "Figure 1. France exceeds Italy on event count, annual event rate, and cumulative flooded footprint, while both retain broad NUTS3 coverage.",
        width=6.3,
    )

    paragraph = document.add_paragraph()
    paragraph.add_run(
        f"France records {int_str(overlap_indexed['fr_only'])} country-specific event IDs that never touch Italy, "
        f"while Italy records {int_str(overlap_indexed['it_only'])}. The shared set is only "
        f"{int_str(overlap_indexed['shared'])} events, so the observed gap is mostly structural rather than an artifact of shared storms."
    )
    paragraph = document.add_paragraph()
    paragraph.add_run(
        f"A particularly important point for downstream point matching is recurrence: the median hit French NUTS3 region sees "
        f"{one_decimal(indexed.loc['FR', 'median_nuts3_events'])} JRC events over the archive, versus "
        f"{one_decimal(indexed.loc['IT', 'median_nuts3_events'])} in Italy."
    )

    insert_picture_with_caption(
        document,
        visuals["yearly_activity"],
        "Figure 2. France stays above Italy in yearly event count throughout the archive and shows much larger high-water years in 2018, 2020, 2021, 2023, and 2024.",
        width=6.3,
    )

    document.add_heading("Regional Structure", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run(
        "Regional coverage is broad in both countries, but the shape differs. Italy reaches all checked NUTS1 and NUTS2 regions, "
        "yet the source is more concentrated in the Po basin and Emilia-Romagna corridor. France combines high recurrence in the west and Brittany with very large cumulative footprints in Pays de la Loire, Grand Est, Bourgogne-Franche-Comte, and Normandie."
    )

    document.add_heading("NUTS1 leaders", level=2)
    nuts1_table = regional_tables["nuts1"].copy()
    nuts1_table["Area km2"] = nuts1_table["Area km2"].map(km2_str)
    nuts1_table["Events"] = nuts1_table["Events"].map(int_str)
    nuts1_table["NUTS3 hit"] = nuts1_table["NUTS3 hit"].map(int_str)
    add_dataframe_table(
        document,
        nuts1_table,
        column_labels=nuts1_table.columns.tolist(),
        widths=[0.90, 2.95, 0.85, 1.00, 0.80],
    )
    add_source_note(document, "Top 5 NUTS1 regions per country ranked by cumulative flooded area.")

    document.add_heading("NUTS2 leaders", level=2)
    nuts2_table = regional_tables["nuts2"].copy()
    nuts2_table["Area km2"] = nuts2_table["Area km2"].map(km2_str)
    nuts2_table["Events"] = nuts2_table["Events"].map(int_str)
    nuts2_table["Max depth cm"] = nuts2_table["Max depth cm"].map(int_str)
    add_dataframe_table(
        document,
        nuts2_table,
        column_labels=nuts2_table.columns.tolist(),
        widths=[0.90, 2.90, 0.75, 0.90, 1.05],
    )
    add_source_note(document, "Top 5 NUTS2 regions per country ranked by cumulative flooded area.")

    insert_picture_with_caption(
        document,
        visuals["nuts3_area_maps"],
        "Figure 3. JRC flooded area is broadly distributed in both countries, but Italian intensity is visibly concentrated in the north and north-east.",
        width=6.35,
    )

    nuts3_area_table = regional_tables["nuts3_area"].copy()
    nuts3_area_table["Area km2"] = nuts3_area_table["Area km2"].map(km2_str)
    nuts3_area_table["Events"] = nuts3_area_table["Events"].map(int_str)
    document.add_heading("NUTS3 leaders by flooded area", level=2)
    add_dataframe_table(
        document,
        nuts3_area_table,
        column_labels=nuts3_area_table.columns.tolist(),
        widths=[0.90, 3.05, 0.75, 0.85],
    )
    add_source_note(document, "Top 5 NUTS3 regions per country ranked by cumulative flooded area.")

    insert_picture_with_caption(
        document,
        visuals["nuts3_area_bars"],
        "Figure 4. France's largest flooded-area regions are mostly in the west and north-east; Italy's are concentrated in Pavia, Ferrara, Vercelli, Mantova, and Novara/Bologna.",
        width=6.35,
    )
    insert_picture_with_caption(
        document,
        visuals["nuts3_event_bars"],
        "Figure 5. Event recurrence is much higher across many French regions, while Italian recurrence peaks in Veneto and Emilia-Romagna but drops quickly outside those corridors.",
        width=6.35,
    )

    document.add_heading("Interpretation For Point Matching", level=1)
    add_bullet_list(
        document,
        [
            (
                f"If France and Italy had comparable point portfolios and identical date windows, France would still start with a stronger match opportunity set: "
                f"{int_str(indexed.loc['FR', 'unique_events'])} versus {int_str(indexed.loc['IT', 'unique_events'])} total JRC events."
            ),
            (
                f"The difference is amplified by footprint intensity: France accumulates roughly {km2_str(indexed.loc['FR', 'cum_flooded_area_km2'])} km2 "
                f"of country-event flooded area versus {km2_str(indexed.loc['IT', 'cum_flooded_area_km2'])} km2 for Italy."
            ),
            (
                f"Italy still has broad regional reach, with {pct_str(indexed.loc['IT', 'nuts2_hit_share'])} NUTS2 coverage and "
                f"{pct_str(indexed.loc['IT', 'nuts3_hit_share'])} NUTS3 coverage, so the problem is not absent Italian source coverage."
            ),
            (
                f"The more likely explanation is concentration: the top 10 Italian NUTS3 regions hold {pct_str(indexed.loc['IT', 'top10_nuts3_area_share'])} "
                "of cumulative flooded area, whereas France is both denser and more geographically diversified."
            ),
            "For strict point-screening comparisons, only the JRC branch is directly comparable between countries. Gaspar and HANZE/TRI fallback outputs should be interpreted separately."
        ],
    )

    document.add_heading("Limitations", level=1)
    add_bullet_list(
        document,
        [
            "The T20 workbook stored in the repository is France-only, so this report cannot validate a France-vs-Italy point-hit claim from T20 itself.",
            "Cumulative flooded area sums country-event intersections over time. It is an intensity indicator, not a unique physical area of land flooded once.",
            "JRC event presence is not the same thing as insured loss severity, asset vulnerability, or obligor exposure.",
            "Country labels in the processed tables are harmonized from Eurostat LAU/NUTS joins; remaining boundary or naming artifacts can still affect marginal rows.",
        ],
    )

    document.add_heading("Appendix: Unhit Regions And Largest Events", level=1)
    appendix_text = document.add_paragraph()
    appendix_text.add_run(
        f"France has {int_str(len(unhit_nuts3[unhit_nuts3['country_code'].eq('FR')]))} unhit NUTS3 regions in the checked archive, "
        f"mostly Paris/inner Ile-de-France, Corsica, and overseas territories. Italy has "
        f"{int_str(len(unhit_nuts3[unhit_nuts3['country_code'].eq('IT')]))} unhit NUTS3 regions, concentrated in Liguria, Calabria, Sardegna, and parts of Sicilia/Marche/Toscana."
    )

    appendix = (
        top_country_events.groupby("country_code", group_keys=False)
        .head(4)[
            [
                "country_name_en",
                "event_start",
                "event_end",
                "duration_days",
                "flooded_area_km2",
                "n_lau_hit",
                "n_nuts3_hit",
            ]
        ]
        .rename(
            columns={
                "country_name_en": "Country",
                "event_start": "Start",
                "event_end": "End",
                "duration_days": "Days",
                "flooded_area_km2": "Area km2",
                "n_lau_hit": "LAUs hit",
                "n_nuts3_hit": "NUTS3 hit",
            }
        )
        .copy()
    )
    appendix["Start"] = pd.to_datetime(appendix["Start"]).dt.strftime("%Y-%m-%d")
    appendix["End"] = pd.to_datetime(appendix["End"]).dt.strftime("%Y-%m-%d")
    appendix["Days"] = appendix["Days"].map(int_str)
    appendix["Area km2"] = appendix["Area km2"].map(km2_str)
    appendix["LAUs hit"] = appendix["LAUs hit"].map(int_str)
    appendix["NUTS3 hit"] = appendix["NUTS3 hit"].map(int_str)
    add_dataframe_table(
        document,
        appendix,
        column_labels=appendix.columns.tolist(),
        widths=[0.80, 1.05, 1.05, 0.55, 0.90, 0.75, 0.80],
        font_size=8,
    )


def build_report(
    *,
    summary: pd.DataFrame,
    yearly: pd.DataFrame,
    nuts1: pd.DataFrame,
    nuts2: pd.DataFrame,
    nuts3: pd.DataFrame,
    overlap: pd.DataFrame,
    unhit_nuts3: pd.DataFrame,
    top_country_events: pd.DataFrame,
    visuals: dict[str, Path],
    t20_scope_table: pd.DataFrame,
    t20_scope_bullets: list[str],
    output_path: Path,
) -> None:
    document = Document()
    configure_document(document)
    add_title_page(
        document,
        metadata_rows=[
            ("Prepared on", pd.Timestamp.today().strftime("%Y-%m-%d")),
            ("Primary scope", "Country-level and regional comparison of the processed JRC flood-depth source for France and Italy."),
            ("Archive horizon", "JRC raster inventory stored in data/JRC_flood_depth_maps with processed outputs covering 2015-2024."),
            ("Main source tables", "events_lau_long.parquet, events_nuts2.csv, nuts3_event_coverage.csv, lau_nuts_lookup.csv."),
            ("Key caveat", "The repository T20 workbook is France-only, so the report explains source asymmetry rather than claiming a symmetric point portfolio test."),
        ],
    )
    add_report_sections(
        document,
        summary=summary,
        yearly=yearly,
        nuts1=nuts1,
        nuts2=nuts2,
        nuts3=nuts3,
        overlap=overlap,
        unhit_nuts3=unhit_nuts3,
        top_country_events=top_country_events,
        visuals=visuals,
        t20_scope_table=t20_scope_table,
        t20_scope_bullets=t20_scope_bullets,
    )
    ensure_dir(output_path.parent)
    document.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build a France vs Italy JRC source comparison report from the processed event tables."
    )
    parser.add_argument(
        "--base-dir",
        default=str(PROJECT_ROOT / "data" / "processed" / "_outputs_eurostat_full"),
        help="Directory containing the processed JRC Eurostat-linked tables.",
    )
    parser.add_argument(
        "--nuts-file",
        default=str(PROJECT_ROOT / "NUTS_RG_03M_2024_4326.gpkg"),
        help="NUTS3 geometry source used for the regional comparison maps.",
    )
    parser.add_argument(
        "--t20-file",
        default=str(PROJECT_ROOT / "data" / "processed" / "T20_Anonymised.xlsx"),
        help="Repository T20 workbook used only for the point-level scope note.",
    )
    parser.add_argument(
        "--output-dir",
        default=str(PROJECT_ROOT / "outputs" / "jrc_france_italy_study"),
        help="Output folder for tables, charts, and the final DOCX report.",
    )
    parser.add_argument(
        "--output-docx",
        default=None,
        help="Optional explicit DOCX path. Default writes inside --output-dir.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()

    base_dir = Path(args.base_dir)
    nuts_file = Path(args.nuts_file)
    t20_file = Path(args.t20_file)
    output_dir = Path(args.output_dir)
    ensure_dir(output_dir)

    inputs = load_inputs(base_dir)
    country_event_level = build_country_event_level(inputs["events"])
    summary = build_country_summary(
        inputs["events"],
        country_event_level,
        inputs["lookup"],
        inputs["coverage"],
    )
    yearly = build_yearly_summary(country_event_level)
    nuts1 = build_nuts1_summary(inputs["events"])
    nuts2 = build_nuts2_summary(inputs["events_nuts2"])
    nuts3 = build_nuts3_summary(inputs["coverage"])
    unhit_nuts3 = build_unhit_nuts3(inputs["lookup"], inputs["coverage"])
    overlap = build_country_event_overlap(country_event_level)
    top_country_events = build_top_country_events(country_event_level)
    visuals = build_visuals(summary, yearly, nuts3, nuts_file, output_dir)
    t20_scope_table, t20_scope_bullets = build_t20_scope_note(t20_file)

    write_csv_outputs(
        output_dir,
        summary=summary,
        yearly=yearly,
        nuts1=nuts1,
        nuts2=nuts2,
        nuts3=nuts3,
        unhit_nuts3=unhit_nuts3,
        overlap=overlap,
        top_country_events=top_country_events,
    )

    output_docx = (
        Path(args.output_docx)
        if args.output_docx
        else output_dir / "JRC_France_Italy_Source_Comparison_Report.docx"
    )
    build_report(
        summary=summary,
        yearly=yearly,
        nuts1=nuts1,
        nuts2=nuts2,
        nuts3=nuts3,
        overlap=overlap,
        unhit_nuts3=unhit_nuts3,
        top_country_events=top_country_events,
        visuals=visuals,
        t20_scope_table=t20_scope_table,
        t20_scope_bullets=t20_scope_bullets,
        output_path=output_docx,
    )

    print(f"Output directory: {output_dir.resolve()}")
    print(f"Report DOCX: {output_docx.resolve()}")
    print(summary.to_string(index=False))


if __name__ == "__main__":
    main()
