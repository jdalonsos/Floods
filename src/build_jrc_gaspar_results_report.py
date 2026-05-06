from __future__ import annotations

import argparse
import math
import textwrap
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPORT_TITLE = "France JRC vs Gaspar Flood Comparison"
REPORT_SUBTITLE = "JRC flood-depth maps, Gaspar CATNAT recognitions, and match results"


@dataclass
class PaperMetadata:
    title: str | None
    doi: str | None
    journal: str | None
    year: str | None
    authors: list[str]
    abstract: str | None


def read_metric_csv(path: Path) -> dict[str, float | int | str]:
    df = pd.read_csv(path)
    return dict(zip(df["metric"], df["value"]))


def read_coverage_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)


def numeric(value: object) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return float("nan")


def extract_paper_metadata(xml_path: Path) -> PaperMetadata:
    if not xml_path.exists():
        return PaperMetadata(None, None, None, None, [], None)

    root = ET.parse(xml_path).getroot()
    title_node = root.find(".//article-title")
    abstract_nodes = root.findall(".//abstract//p")
    authors: list[str] = []
    for contrib in root.findall(".//contrib[@contrib-type='author']"):
        surname = contrib.findtext(".//surname") or ""
        given = contrib.findtext(".//given-names") or ""
        full = f"{given} {surname}".strip()
        if full:
            authors.append(full)

    abstract = " ".join(" ".join(p.itertext()).strip() for p in abstract_nodes).strip()
    if abstract:
        abstract = " ".join(abstract.split())

    return PaperMetadata(
        title="".join(title_node.itertext()).strip() if title_node is not None else None,
        doi=root.findtext(".//article-id[@pub-id-type='doi']"),
        journal=root.findtext(".//journal-title"),
        year=root.findtext(".//pub-date/year"),
        authors=authors,
        abstract=abstract,
    )


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path(r"C:\Windows\Fonts\segoeuib.ttf") if bold else Path(r"C:\Windows\Fonts\segoeui.ttf"),
        Path(r"C:\Windows\Fonts\arialbd.ttf") if bold else Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        proposal = word if not current else f"{current} {word}"
        if draw.textlength(proposal, font=font) <= width:
            current = proposal
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def create_year_bar_chart(year_counts: dict[int, int], output_path: Path) -> None:
    width, height = 1400, 720
    margin_left, margin_right, margin_top, margin_bottom = 110, 60, 100, 110
    background = (248, 250, 252)
    ink = (31, 41, 55)
    grid = (203, 213, 225)
    accent = (15, 118, 110)

    image = Image.new("RGB", (width, height), background)
    draw = ImageDraw.Draw(image)
    title_font = get_font(34, bold=True)
    label_font = get_font(22)
    tick_font = get_font(18)

    draw.text((margin_left, 28), "JRC France events by start year", fill=ink, font=title_font)
    draw.text(
        (margin_left, 66),
        "Unique JRC event IDs after France LAU to INSEE harmonization",
        fill=(71, 85, 105),
        font=label_font,
    )

    years = sorted(year_counts)
    max_value = max(year_counts.values())
    chart_x0 = margin_left
    chart_y0 = margin_top
    chart_x1 = width - margin_right
    chart_y1 = height - margin_bottom
    chart_w = chart_x1 - chart_x0
    chart_h = chart_y1 - chart_y0

    for step in range(0, max_value + 5, 5):
        y = chart_y1 - (step / max_value) * chart_h
        draw.line((chart_x0, y, chart_x1, y), fill=grid, width=1)
        draw.text((25, y - 10), str(step), fill=(100, 116, 139), font=tick_font)

    bar_gap = 18
    bar_width = (chart_w - bar_gap * (len(years) - 1)) / len(years)
    for idx, year in enumerate(years):
        value = year_counts[year]
        x0 = chart_x0 + idx * (bar_width + bar_gap)
        x1 = x0 + bar_width
        y0 = chart_y1 - (value / max_value) * chart_h
        draw.rounded_rectangle((x0, y0, x1, chart_y1), radius=8, fill=accent)
        year_text = str(year)
        text_w = draw.textlength(year_text, font=tick_font)
        draw.text((x0 + (bar_width - text_w) / 2, chart_y1 + 12), year_text, fill=ink, font=tick_font)
        value_text = str(value)
        v_w = draw.textlength(value_text, font=tick_font)
        draw.text((x0 + (bar_width - v_w) / 2, y0 - 28), value_text, fill=ink, font=tick_font)

    draw.line((chart_x0, chart_y1, chart_x1, chart_y1), fill=ink, width=2)
    image.save(output_path)


def create_sensitivity_chart(rows: list[tuple[str, float, float]], output_path: Path) -> None:
    width, height = 1500, 760
    background = (248, 250, 252)
    ink = (31, 41, 55)
    grid = (226, 232, 240)
    color_7d = (37, 99, 235)
    color_30d = (15, 118, 110)

    image = Image.new("RGB", (width, height), background)
    draw = ImageDraw.Draw(image)
    title_font = get_font(34, bold=True)
    label_font = get_font(22)
    tick_font = get_font(18)
    draw.text((90, 28), "Strict 7-day vs flexible 30-day event match shares", fill=ink, font=title_font)
    draw.text(
        (90, 66),
        "Event-level match rates: stricter dates recover fewer pairs, especially at department level",
        fill=(71, 85, 105),
        font=label_font,
    )

    chart_x0, chart_y0, chart_x1, chart_y1 = 120, 130, 1440, 650
    chart_h = chart_y1 - chart_y0
    row_gap = 32
    pair_gap = 24
    bar_h = 24
    rows_count = len(rows)
    band_h = (chart_h - row_gap * (rows_count - 1)) / rows_count

    for pct in range(0, 101, 10):
        x = chart_x0 + (pct / 100) * (chart_x1 - chart_x0)
        draw.line((x, chart_y0, x, chart_y1), fill=grid, width=1)
        label = f"{pct}%"
        tw = draw.textlength(label, font=tick_font)
        draw.text((x - tw / 2, chart_y1 + 12), label, fill=(100, 116, 139), font=tick_font)

    for idx, (label, share_7d, share_30d) in enumerate(rows):
        band_top = chart_y0 + idx * (band_h + row_gap)
        draw.text((20, band_top + 8), label, fill=ink, font=label_font)
        bars_top = band_top + 44

        for offset, share, color, tag in [
            (0, share_7d, color_7d, "7d"),
            (bar_h + pair_gap, share_30d, color_30d, "30d"),
        ]:
            y0 = bars_top + offset
            y1 = y0 + bar_h
            x1 = chart_x0 + share * (chart_x1 - chart_x0)
            draw.rounded_rectangle((chart_x0, y0, x1, y1), radius=7, fill=color)
            label_text = f"{tag}: {share:.1%}"
            draw.text((x1 + 12, y0 - 2), label_text, fill=ink, font=tick_font)

    draw.rounded_rectangle((1180, 92, 1415, 146), radius=10, fill=(255, 255, 255), outline=(203, 213, 225))
    draw.rounded_rectangle((1200, 106, 1230, 130), radius=4, fill=color_7d)
    draw.text((1242, 104), "7-day rule", fill=ink, font=tick_font)
    draw.rounded_rectangle((1200, 134, 1230, 158), radius=4, fill=color_30d)
    draw.text((1242, 132), "30-day rule", fill=ink, font=tick_font)

    image.save(output_path)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_default_font(document: Document, font_name: str = "Aptos", size_pt: int = 10) -> None:
    normal = document.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(size_pt)
    normal._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)

    for style_name, size, color in [
        ("Title", 22, RGBColor(15, 23, 42)),
        ("Heading 1", 16, RGBColor(15, 23, 42)),
        ("Heading 2", 13, RGBColor(15, 118, 110)),
        ("Heading 3", 11, RGBColor(30, 41, 59)),
    ]:
        style = document.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)


def add_title_page(document: Document, output_root: Path) -> None:
    p = document.add_paragraph()
    p.style = "Title"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(REPORT_TITLE)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(REPORT_SUBTITLE)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(71, 85, 105)

    box = document.add_table(rows=3, cols=2)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.style = "Table Grid"
    pairs = [
        ("Report scope", "France, JRC flood-depth maps vs Gaspar flood recognitions"),
        ("Main result folders", "data/processed/jrc_gaspar_comparison_flexible_7d and ..._30d"),
        ("Prepared from", "Current processed outputs and JRC source documentation"),
    ]
    for row, (left, right) in zip(box.rows, pairs):
        row.cells[0].text = left
        row.cells[1].text = right
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.8)
        set_cell_shading(row.cells[0], "E2E8F0")
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document.add_paragraph("")
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "This document explains what the comparison outputs mean, how the matching logic works, "
        "and how to interpret the strict 7-day and flexible 30-day variants."
    )
    run.italic = True
    run.font.color.rgb = RGBColor(71, 85, 105)

    document.add_page_break()


def add_bullet_list(document: Document, items: Iterable[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_executive_summary(document: Document, bullets: list[str]) -> None:
    document.add_heading("Executive summary", level=1)
    table = document.add_table(rows=len(bullets), cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, bullet in enumerate(bullets):
        cell = table.rows[idx].cells[0]
        set_cell_shading(cell, "F8FAFC" if idx % 2 == 0 else "EEF2FF")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("• ")
        run.bold = True
        p.add_run(bullet)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_dataset_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = table.rows[0].cells
    headers[0].text = "Field"
    headers[1].text = "Value"
    for cell in headers:
        set_cell_shading(cell, "DBEAFE")
        cell.paragraphs[0].runs[0].bold = True
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
        row[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(row[0], "F8FAFC")


def add_key_field_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = table.rows[0].cells
    headers[0].text = "Field"
    headers[1].text = "Meaning"
    for cell in headers:
        set_cell_shading(cell, "CCFBF1")
        cell.paragraphs[0].runs[0].bold = True
    for field, meaning in rows:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = meaning
        row[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(row[0], "F8FAFC")


def add_coverage_table(document: Document, cover7: pd.DataFrame, cover30: pd.DataFrame) -> None:
    document.add_heading("Match coverage: strict 7-day vs flexible 30-day", level=1)
    p = document.add_paragraph()
    p.add_run("How to read this table: ").bold = True
    p.add_run(
        "“Unique events” counts one event ID per source. “Canonical rows” counts one comparison row per commune-event "
        "or department-event record, so row totals are much larger than event totals."
    )

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Level / measure", "JRC 7d", "JRC 30d", "Gaspar 7d", "Gaspar 30d"]
    for cell, label in zip(table.rows[0].cells, headers):
        cell.text = label
        set_cell_shading(cell, "DBEAFE")
        cell.paragraphs[0].runs[0].bold = True

    merged = cover7.merge(cover30, on=["level", "measurement"], suffixes=("_7d", "_30d"))
    order = [
        ("commune", "unique_events"),
        ("commune", "canonical_rows"),
        ("department", "unique_events"),
        ("department", "canonical_rows"),
    ]
    for level, measurement in order:
        row_data = merged[(merged["level"] == level) & (merged["measurement"] == measurement)].iloc[0]
        row = table.add_row().cells
        row[0].text = f"{level.title()} / {measurement.replace('_', ' ')}"
        row[1].text = f"{int(row_data['jrc_matched_7d'])}/{int(row_data['jrc_total_7d'])} ({numeric(row_data['jrc_match_share_7d']):.1%})"
        row[2].text = f"{int(row_data['jrc_matched_30d'])}/{int(row_data['jrc_total_30d'])} ({numeric(row_data['jrc_match_share_30d']):.1%})"
        row[3].text = f"{int(row_data['gaspar_matched_7d'])}/{int(row_data['gaspar_total_7d'])} ({numeric(row_data['gaspar_match_share_7d']):.1%})"
        row[4].text = f"{int(row_data['gaspar_matched_30d'])}/{int(row_data['gaspar_total_30d'])} ({numeric(row_data['gaspar_match_share_30d']):.1%})"
        set_cell_shading(row[0], "F8FAFC")


def add_year_table(document: Document, jrc_year_counts: dict[int, int], gaspar_year_counts: dict[int, int]) -> None:
    document.add_heading("Annual profile of the two sources", level=1)
    p = document.add_paragraph()
    p.add_run("Important caveat: ").bold = True
    p.add_run(
        "JRC counts are unique physical map events. Gaspar counts below are recognition-period proxies "
        "built as cod_nat_catnat + start date + end date, so they are not directly equivalent to physical disaster episodes."
    )

    years = sorted(set(jrc_year_counts) | set(gaspar_year_counts))
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Start year", "JRC unique events", "Gaspar recognition-period groups"]
    for cell, label in zip(table.rows[0].cells, headers):
        cell.text = label
        set_cell_shading(cell, "DBEAFE")
        cell.paragraphs[0].runs[0].bold = True
    for year in years:
        row = table.add_row().cells
        row[0].text = str(year)
        row[1].text = str(jrc_year_counts.get(year, 0))
        row[2].text = str(gaspar_year_counts.get(year, 0))
        if year % 2 == 0:
            for cell in row:
                set_cell_shading(cell, "F8FAFC")


def add_interpretation_section(document: Document) -> None:
    document.add_heading("Interpretation of the 30-day flexible results", level=1)
    add_bullet_list(
        document,
        [
            "Department-level matching is substantially stronger than commune-level matching. This is expected because JRC event footprints and Gaspar recognitions rarely line up commune by commune.",
            "The flexible 30-day rule is materially more permissive than the strict 7-day rule. It recovers additional matches when the same flood episode is fragmented differently across the two datasets.",
            "Most department best-match pairs are 'Gaspar full, JRC partial'. In practice, this means a Gaspar recognition-period group is often fully contained inside a broader JRC event footprint or duration, rather than the other way around.",
            "Exact same start and end dates are rare. Spatial overlap plus a tolerant time window is doing most of the matching work.",
        ],
    )


def add_limitations_section(document: Document, comp30: dict[str, object]) -> None:
    document.add_heading("Limitations and recommended reading of the outputs", level=1)
    bullets = [
        "Gaspar is an administrative recognition database, not a satellite event catalogue. In the current comparison code, gaspar_event_uid is a recognition-period proxy built as cod_nat_catnat + start date + end date.",
        "This means Gaspar 'unique event' counts are best read as recognition-period groups, not as perfectly reconstructed physical flood events.",
        f"The 30-day diagnostics show {int(comp30['gaspar_unique_decrees'])} unique decree IDs but {int(comp30['gaspar_unique_event_uids'])} unique Gaspar recognition-period groups, which illustrates how a single decree can branch into many dated groups.",
        "Commune-level row match shares are low by construction because both sources partition the same episode differently across space and time. Department-level metrics are more stable for management reporting.",
        "JRC flood maps cover land only, use 20 m pixels, and encode permanent or seasonal water as dummy value 9999. They are therefore very strong for flood extent and relative timing, but not a direct mirror of compensation or administrative recognition processes.",
        "Recommended baseline: use the 30-day department-level outputs for the main narrative, keep the 7-day outputs as a stricter sensitivity test, and use commune tables only for local audit or deep dives.",
    ]
    add_bullet_list(document, bullets)


def add_file_guide(document: Document) -> None:
    document.add_heading("Which output files to open first", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "File"
    headers[1].text = "What it is for"
    for cell in headers:
        set_cell_shading(cell, "DBEAFE")
        cell.paragraphs[0].runs[0].bold = True

    rows = [
        ("comparison_summary.xlsx", "Headline counts for each run."),
        ("coverage_overview.xlsx", "The key distinction between unique events and canonical rows."),
        ("best_match_overview_commune.xlsx", "One suggested best Gaspar partner per JRC event at commune level."),
        ("best_match_overview_department.xlsx", "Same idea at department level; usually the best first analytical table."),
        ("jrc_gaspar_comparison_flexible.xlsx", "Combined workbook for people who prefer one Excel file."),
        ("details/", "Raw canonical tables, unmatched rows, and audit tables for debugging."),
    ]
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
        row[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(row[0], "F8FAFC")


def add_references(document: Document, paper: PaperMetadata) -> None:
    document.add_heading("References", level=1)
    refs = [
        "JRC dataset page: https://data.jrc.ec.europa.eu/dataset/0bc96690-b89c-4909-9166-c2c322a20130",
        "JRC dataset DOI: 10.2905/0bc96690-b89c-4909-9166-c2c322a20130",
        "JRC storage: https://jeodpp.jrc.ec.europa.eu/ftp/jrc-opendata/CEMS-EFAS/European_Satellite-Derived_Flood_Depth_Maps/",
        "GASPAR glossary page: https://www.georisques.gouv.fr/gaspar",
        "GASPAR administrative-procedures download page: https://www.georisques.gouv.fr/donnees/bases-de-donnees/procedures-administratives-relatives-aux-risques",
        "GASPAR official data dictionary PDF: https://www.georisques.gouv.fr/sites/default/files/Georisques_DictionnaireDonneesGaspar_1.0.pdf",
        "GASPAR CatNat open-data listing: https://www.data.gouv.fr/datasets/risques-arretes-catastrophes-naturelles",
        "Géorisques explanation of the CatNat indemnification mechanism: https://georisques.gouv.fr/le-dispositif-dindemnisation-des-catastrophes-naturelles",
    ]
    if paper.title or paper.doi:
        author_text = ", ".join(paper.authors[:2]) if paper.authors else "Andrea Betterle and Peter Salamon"
        refs.append(
            f"{author_text} ({paper.year or '2024'}), {paper.title or 'Water depth estimate and flood extent enhancement for satellite-based inundation maps'}, "
            f"{paper.journal or 'Natural Hazards and Earth System Sciences'}, DOI: {paper.doi or '10.5194/nhess-24-2817-2024'}"
        )
    refs.append(
        "Processed Gaspar source used by the comparison scripts: data/processed/Gaspar_2015_2024.xlsx "
        "(sheet Gaspar20152024FloodsClean)."
    )
    refs.append(
        "France JRC comparison input used by the scripts: data/processed/france_lau_insee_documentation/events_fr_insee_long.csv."
    )
    add_bullet_list(document, refs)


def insert_picture_with_caption(document: Document, image_path: Path, caption: str, width_in: float) -> None:
    document.add_picture(str(image_path), width=Inches(width_in))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(71, 85, 105)


def build_report(args: argparse.Namespace) -> Path:
    root = Path(args.project_root).resolve()
    out_path = Path(args.output).resolve()
    ensure_dir(out_path.parent)

    assets_dir = out_path.parent / "_report_assets"
    ensure_dir(assets_dir)

    comp7 = read_metric_csv(root / "data/processed/jrc_gaspar_comparison_flexible_7d/comparison_summary.csv")
    comp30 = read_metric_csv(root / "data/processed/jrc_gaspar_comparison_flexible_30d/comparison_summary.csv")
    cover7 = read_coverage_csv(root / "data/processed/jrc_gaspar_comparison_flexible_7d/coverage_overview.csv")
    cover30 = read_coverage_csv(root / "data/processed/jrc_gaspar_comparison_flexible_30d/coverage_overview.csv")

    jrc = pd.read_csv(
        root / "data/processed/france_lau_insee_documentation/events_fr_insee_long.csv",
        parse_dates=["start_date", "end_date"],
    )
    event_col = "jrc_event_id" if "jrc_event_id" in jrc.columns else "event_id"
    jrc_events = jrc[[event_col, "start_date", "end_date"]].drop_duplicates()
    jrc_events["year"] = jrc_events["start_date"].dt.year
    jrc_year_counts = {int(k): int(v) for k, v in jrc_events.groupby("year").size().to_dict().items()}

    gaspar = pd.read_csv(
        root / "data/processed/jrc_gaspar_comparison_flexible_30d/details/gaspar_commune_events_canonical.csv",
        parse_dates=["gaspar_start_date", "gaspar_end_date"],
    )
    gaspar_workbook = pd.read_excel(
        root / "data/processed/Gaspar_2015_2024.xlsx",
        sheet_name="Gaspar20152024FloodsClean",
    )
    gaspar_year_counts = {
        int(k): int(v)
        for k, v in gaspar[["gaspar_event_uid", "gaspar_start_date"]]
        .drop_duplicates()
        .assign(year=lambda df: df["gaspar_start_date"].dt.year)
        .groupby("year")
        .size()
        .to_dict()
        .items()
    }
    gaspar_risk_counts = {
        str(k): int(v) for k, v in gaspar_workbook["lib_risque_jo"].value_counts().to_dict().items()
    }

    paper = extract_paper_metadata(Path(args.paper_xml))

    year_chart = assets_dir / "jrc_events_by_year.png"
    create_year_bar_chart(jrc_year_counts, year_chart)

    sensitivity_chart = assets_dir / "strict_vs_flexible_event_match_shares.png"
    sensitivity_rows = [
        (
            "JRC commune unique events",
            numeric(cover7[(cover7["level"] == "commune") & (cover7["measurement"] == "unique_events")]["jrc_match_share"].iloc[0]),
            numeric(cover30[(cover30["level"] == "commune") & (cover30["measurement"] == "unique_events")]["jrc_match_share"].iloc[0]),
        ),
        (
            "JRC department unique events",
            numeric(cover7[(cover7["level"] == "department") & (cover7["measurement"] == "unique_events")]["jrc_match_share"].iloc[0]),
            numeric(cover30[(cover30["level"] == "department") & (cover30["measurement"] == "unique_events")]["jrc_match_share"].iloc[0]),
        ),
        (
            "Gaspar commune unique event groups",
            numeric(cover7[(cover7["level"] == "commune") & (cover7["measurement"] == "unique_events")]["gaspar_match_share"].iloc[0]),
            numeric(cover30[(cover30["level"] == "commune") & (cover30["measurement"] == "unique_events")]["gaspar_match_share"].iloc[0]),
        ),
        (
            "Gaspar department unique event groups",
            numeric(cover7[(cover7["level"] == "department") & (cover7["measurement"] == "unique_events")]["gaspar_match_share"].iloc[0]),
            numeric(cover30[(cover30["level"] == "department") & (cover30["measurement"] == "unique_events")]["gaspar_match_share"].iloc[0]),
        ),
    ]
    create_sensitivity_chart(sensitivity_rows, sensitivity_chart)

    document = Document()
    set_default_font(document)
    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    add_title_page(document, out_path.parent)

    executive_bullets = [
        f"The France comparison uses {int(comp30['jrc_unique_events'])} JRC flood events and {int(comp30['gaspar_unique_event_uids'])} Gaspar recognition-period groups derived from {int(comp30['gaspar_unique_decrees'])} CATNAT decree IDs.",
        f"At the department level, the strict 7-day rule matches {int(comp7['matched_jrc_events_department_level'])} of {int(comp7['jrc_unique_events'])} JRC events ({numeric(cover7[(cover7['level'] == 'department') & (cover7['measurement'] == 'unique_events')]['jrc_match_share'].iloc[0]):.1%}), while the flexible 30-day rule matches {int(comp30['matched_jrc_events_department_level'])} ({numeric(cover30[(cover30['level'] == 'department') & (cover30['measurement'] == 'unique_events')]['jrc_match_share'].iloc[0]):.1%}).",
        f"At the commune level, match rates remain much lower: {int(comp7['matched_jrc_events_commune_level'])} JRC events under the 7-day rule and {int(comp30['matched_jrc_events_commune_level'])} under the 30-day rule.",
        "The department-level comparison is easier to interpret than the commune-level comparison because the two sources break the same flood episode into space and time differently.",
        "The flexible 30-day outputs are the best operational baseline, but event-level Gaspar counts should be read as administrative recognition-period proxies rather than true physical disaster-event IDs.",
    ]
    add_executive_summary(document, executive_bullets)

    document.add_heading("Data sources used in the comparison", level=1)
    document.add_paragraph(
        "This report documents the processed France comparison between the European Satellite-Derived Flood Depth Maps "
        "published by the Joint Research Centre (JRC) and the French Gaspar CATNAT recognition database."
    )

    document.add_heading("JRC flood-depth maps", level=2)
    add_dataset_table(
        document,
        [
            ("Dataset name", "European Satellite-Derived Flood Depth Maps"),
            ("Maintainer", "Copernicus Emergency Management Service / JRC catalogue"),
            ("DOI", "10.2905/0bc96690-b89c-4909-9166-c2c322a20130"),
            ("Coverage", "Europe and surroundings, land only"),
            ("Time span", "2015-01-01 to 2024-12-31 in the official release"),
            ("Resolution", "20 m"),
            ("Projection", "Azimuthal equidistant"),
            ("File format", "GeoTIFF, water depth in cm, uint16"),
            ("Important dummy value", "9999 indicates permanent or seasonal water bodies"),
            ("Method summary", "Flood extents from Copernicus GFM aggregated with terrain topography using FLEXTH to estimate depth and enhance extent"),
        ],
    )
    if paper.title:
        p = document.add_paragraph()
        p.add_run("Method reference: ").bold = True
        p.add_run(
            f"{', '.join(paper.authors[:2])} ({paper.year}), {paper.title}, "
            f"{paper.journal}, DOI {paper.doi}."
        )
    if paper.abstract:
        p = document.add_paragraph()
        p.add_run("Paper takeaway: ").bold = True
        p.add_run(
            "The NHESS paper explains that FLEXTH expands flood masks into no-data areas using morphology and estimates "
            "water depths from terrain elevation along the flood boundary, allowing large-scale unsupervised processing."
        )

    document.add_heading("Gaspar source used here", level=2)
    add_dataset_table(
        document,
        [
            ("System name", "GASPAR: Gestion Assistee des Procedures Administratives relatives aux Risques"),
            ("Official role", "Backbone of the DGPR natural-risk information system, according to Georisques"),
            ("What the broader database contains", "CatNat recognitions plus preventive and regulatory procedures such as PPR, DICRIM, TIM, AZI and others"),
            ("What we use here", "Only the CatNat recognition table, not the full administrative database"),
            ("Official update note", "Georisques says CatNat procedures are updated within 30 days after publication in the Journal Officiel"),
            ("Open-data access", "National, regional, departmental and annual extracts are offered on Georisques; an open-data listing is also published on data.gouv"),
            ("Licence on data.gouv", "Licence Ouverte / Open Licence version 2.0"),
            ("Observed metadata date", "The data.gouv CatNat page accessed for this report showed a last update of 4 March 2026"),
        ],
    )
    document.add_paragraph(
        "Official Géorisques documentation describes GASPAR as an administrative risk database rather than a physical flood-event catalogue. "
        "This distinction is central for interpreting any comparison with the JRC satellite event set."
    )

    document.add_heading("What one Gaspar CatNat row means", level=3)
    document.add_paragraph(
        "The official GASPAR data dictionary states that one line in catnat_gaspar.csv corresponds to one recognition concerning one commune, "
        "one event period and one risk. Several recognitions can therefore exist for the same commune and the same broad event when different "
        "risks or administrative treatments are involved."
    )
    add_key_field_table(
        document,
        [
            ("cod_nat_catnat", "Unique recognition code generated by Gaspar"),
            ("cod_commune / lib_commune", "Affected commune code and commune name"),
            ("num_risque_jo / lib_risque_jo", "Risk code and official risk wording in the Journal Officiel"),
            ("dat_deb / dat_fin", "Recognized start and end dates of the event period"),
            ("dat_pub_arrete", "Date of the interministerial order"),
            ("dat_pub_jo", "Date of publication in the Journal Officiel"),
            ("dat_maj", "Date the Gaspar record was last updated"),
        ],
    )
    document.add_paragraph(
        "This structure explains why Gaspar does not naturally offer a physical event identifier equivalent to JRC jrc_event_id. "
        "In the current comparison scripts we therefore use a practical proxy, gaspar_event_uid = cod_nat_catnat + start date + end date. "
        "That is useful for matching, but it is still closer to an administrative recognition-period group than to a fully reconstructed flood episode."
    )

    document.add_heading("How the CatNat mechanism shapes the data", level=3)
    add_bullet_list(
        document,
        [
            "Géorisques explains that the CatNat indemnification regime is activated after an interministerial order is published in the Journal Officiel.",
            "The mayor requests recognition for the commune, and the published order specifies the recognized communes, the period of the event and the nature of the damage.",
            "Géorisques also notes that insured parties have 30 days after publication in the Journal Officiel to declare the loss to their insurer.",
            "Because the database follows this recognition workflow, Gaspar dates and event boundaries may legitimately differ from JRC satellite event boundaries.",
        ],
    )

    document.add_heading("Gaspar flood subset used in this project", level=3)
    add_dataset_table(
        document,
        [
            ("Processed workbook", "data/processed/Gaspar_2015_2024.xlsx, sheet Gaspar20152024FloodsClean"),
            ("Rows in the cleaned flood sheet", str(len(gaspar_workbook))),
            ("Canonical commune rows after downstream deduplication", str(int(comp30["gaspar_canonical_commune_rows"]))),
            ("Unique Gaspar recognition-period groups in comparison", str(int(comp30["gaspar_unique_event_uids"]))),
            ("Unique decree IDs in comparison", str(int(comp30["gaspar_unique_decrees"]))),
        ],
    )
    wave_label = "Chocs Mécaniques liés à l'action des Vagues"
    add_bullet_list(
        document,
        [
            f"Inondations et/ou Coulees de Boue: {gaspar_risk_counts.get('Inondations et/ou Coulées de Boue', 0)} rows in the cleaned flood sheet.",
            f"Inondations Remontee Nappe: {gaspar_risk_counts.get('Inondations Remontée Nappe', 0)} rows.",
            f"Chocs Mecaniques lies a l'action des Vagues: {gaspar_risk_counts.get(wave_label, 0)} rows.",
            "Comparison geography: commune level first, then department level derived from commune INSEE codes.",
        ],
    )

    document.add_heading("Comparison workflow", level=1)
    add_bullet_list(
        document,
        [
            "1. JRC rasters are tabularized to LAU, then harmonized to France INSEE communes in data/processed/france_lau_insee_documentation/events_fr_insee_long.csv.",
            "2. Gaspar rows are cleaned, deduplicated, and filtered to flood-related hazards in the preprocessing notebook.",
            "3. The strict run matches the same commune or department when both start and end dates are within 7 days.",
            "4. The flexible run uses a 30-day window and also accepts cross start/end proximity and expanded interval overlap inside the same 30-day tolerance.",
            "5. Department tables are derived from commune INSEE codes; best-match overviews then suggest one strongest Gaspar partner per JRC event.",
            "6. unmatched_jrc_events_* tables count unique JRC events with no match at the chosen level, whereas unmatched_jrc_*_events tables count every unmatched commune or department row.",
        ],
    )

    document.add_heading("JRC France subset used for this report", level=1)
    add_dataset_table(
        document,
        [
            ("Unique JRC events", str(int(comp30["jrc_unique_events"]))),
            ("Canonical commune rows", str(int(comp30["jrc_canonical_commune_rows"]))),
            ("Unique communes touched", str(int(comp30["jrc_unique_communes"]))),
            ("Unique departments touched", str(int(comp30["jrc_unique_departments"]))),
            ("First start date in France subset", str(jrc_events["start_date"].min().date())),
            ("Last end date in France subset", str(jrc_events["end_date"].max().date())),
        ],
    )
    insert_picture_with_caption(
        document,
        year_chart,
        "Figure 1. Unique JRC event IDs by start year after France harmonization. The highest count in the current France subset is 2024 (35 starts).",
        6.8,
    )
    add_year_table(document, jrc_year_counts, gaspar_year_counts)

    add_coverage_table(document, cover7, cover30)
    insert_picture_with_caption(
        document,
        sensitivity_chart,
        "Figure 2. Event-level sensitivity of match rates to the stricter 7-day and more permissive 30-day rules.",
        6.9,
    )
    insert_picture_with_caption(
        document,
        root / "data/processed/jrc_gaspar_comparison_flexible_30d/plots/comparison_overview.png",
        "Figure 3. Coverage overview from the 30-day comparison. Unique-event shares and row-level shares tell different stories and should not be mixed.",
        6.9,
    )

    add_interpretation_section(document)
    insert_picture_with_caption(
        document,
        root / "data/processed/jrc_gaspar_comparison_flexible_30d/plots/department_event_summary.png",
        "Figure 4. Department-level best event-pair summary for the 30-day run. Most best pairs are Gaspar-full and JRC-partial rather than full overlaps on both sides.",
        6.9,
    )

    document.add_heading("Why event-level counts are asymmetric", level=1)
    add_bullet_list(
        document,
        [
            "JRC uses physical flood map events. Gaspar uses administrative recognition rows and the current comparison groups them with a proxy key: cod_nat_catnat + start date + end date.",
            "One JRC event can overlap many Gaspar recognition-period groups, and one Gaspar group can cover only part of a broader JRC event footprint.",
            "That is why department-level event matching is much stronger than commune-level event matching, and why matched row totals differ from matched unique event totals.",
            "In the 30-day run, only one best-match pair has exact same start and end dates at commune level, and only one at department level. Exact dates are the exception, not the rule.",
        ],
    )

    add_file_guide(document)
    add_limitations_section(document, comp30)
    add_references(document, paper)

    document.save(out_path)
    return out_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a Word report for the France JRC vs Gaspar comparison.")
    parser.add_argument(
        "--project-root",
        default=".",
        help="Project root containing data/processed outputs.",
    )
    parser.add_argument(
        "--paper-xml",
        default=r"D:\Users\Juan David Alonso\Downloads\nhess-24-2817-2024.xml",
        help="Path to the NHESS XML paper used for the methodology summary.",
    )
    parser.add_argument(
        "--output",
        default=r"data\processed\JRC_Gaspar_France_Comparison_Report.docx",
        help="Output DOCX path.",
    )
    args = parser.parse_args()
    out_path = build_report(args)
    print(out_path)


if __name__ == "__main__":
    main()
