from __future__ import annotations

import argparse
import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from build_jrc_gaspar_results_report import (
    add_bullet_list,
    add_coverage_table,
    add_dataset_table,
    add_executive_summary,
    add_file_guide,
    add_limitations_section,
    create_sensitivity_chart,
    create_year_bar_chart,
    ensure_dir,
    insert_picture_with_caption,
    numeric,
    read_coverage_csv,
    read_metric_csv,
    set_cell_shading,
    set_default_font,
)


REPORT_TITLE = "France JRC vs Gaspar Flood Comparison"
REPORT_SUBTITLE = (
    "Full analysis report: commune and department results, 2015-2024 horizon audit, "
    "July 2021 case study, and app interpretation"
)

SOURCE_PACK = [
    (
        "data/processed/jrc_gaspar_comparison_flexible_7d",
        "Strict-window comparison outputs used for the narrower sensitivity check.",
    ),
    (
        "data/processed/jrc_gaspar_comparison_flexible_30d",
        "Flexible-window comparison outputs used for the main operational reading.",
    ),
    (
        "docs/gaspar_jrc_match_audit_en.md",
        "National match statistics plus mapped manual checks for selected periods.",
    ),
    (
        "docs/gaspar_jrc_horizon_audit_en.md",
        "2015-2024 region-quarter mismatch ranking and extended manual review.",
    ),
    (
        "docs/july_2021_gaspar_jrc_mismatch_evidence_report.md",
        "External-evidence note used to interpret the July 2021 mismatch.",
    ),
    (
        "docs/france_commune_activity_app_deep_guide.md",
        "Functional explanation of the Streamlit commune-activity comparison app.",
    ),
]

APP_FEATURES = [
    "Switch between Gaspar, JRC, and direct comparison modes on the same France commune map.",
    "Filter by exact date, month, year, quarter, or custom interval.",
    "Inspect active communes, filtered source rows, matched rows, and unresolved Gaspar rows.",
    "Download the current commune table as CSV and the current rendered map as HTML.",
    "Overlay department boundaries while keeping the fill layer at current-commune resolution.",
]

APP_READING = [
    "The app is a temporal commune-activity viewer, not a raw raster viewer.",
    "Gaspar rows are resolved to current INSEE communes before they contribute to the map.",
    "JRC rows are already close to map-ready once the France LAU to INSEE harmonization has been run.",
    "Comparison mode is the fastest way to see both-active, Gaspar-only, and JRC-only communes for one period.",
]

JULY_2021_MONTHLY_SNAPSHOT = {
    "period": "July 2021 (app filter: Year 2021 -> Month 07)",
    "both": 133,
    "gaspar_only": 363,
    "jrc_only": 1175,
}

JULY_2021_EVIDENCE_ROWS = [
    (
        "Jura (39): Arbois / Lons-le-Saunier / Bletterans",
        "Gaspar only",
        "Strong",
        "The earlier note found local reporting, an official basin report, and a CatNat decree that all support real flood impacts.",
    ),
    (
        "Meuse (55): Bar-le-Duc / Behonne",
        "Gaspar only",
        "Strong",
        "The earlier note found local reporting, prefecture material, and CatNat recognition for 13-15 July 2021 floods.",
    ),
    (
        "Seine-et-Marne (77): Claye-Souilly",
        "Gaspar only",
        "Moderate to strong",
        "The earlier note found clear local media evidence of flooded communes and emergency interventions.",
    ),
    (
        "Haute-Saone (70): Autet and nearby communes",
        "JRC only",
        "Moderate",
        "The earlier note found local flood cleanup reporting and department-level vigilance evidence.",
    ),
    (
        "Ardennes (08): Asfeld / Sedan sector",
        "JRC only",
        "Moderate to strong",
        "The earlier note found local reporting on pumping operations, continuing flood conditions, and vigilance context.",
    ),
    (
        "Aisne / Ardennes floodplain: Savigny-sur-Aisne / Vieux-les-Asfeld",
        "JRC only",
        "Strong",
        "The earlier note found site-level flood-mark evidence on the official reperes de crues platform.",
    ),
    (
        "Saone-et-Loire (71): Louhans / Seille corridor",
        "JRC only",
        "Moderate to strong",
        "The earlier note found both government return-experience material and local reporting on flood impacts.",
    ),
]

SELECTED_EXTERNAL_REFERENCES = [
    "JRC dataset DOI: 10.2905/0bc96690-b89c-4909-9166-c2c322a20130",
    "JRC dataset page: https://data.jrc.ec.europa.eu/dataset/0bc96690-b89c-4909-9166-c2c322a20130",
    "GASPAR glossary page: https://www.georisques.gouv.fr/gaspar",
    "GASPAR official data dictionary PDF: https://www.georisques.gouv.fr/sites/default/files/Georisques_DictionnaireDonneesGaspar_1.0.pdf",
    "GASPAR CatNat open-data listing: https://www.data.gouv.fr/datasets/risques-arretes-catastrophes-naturelles",
    "July 2021 official CatNat decree: https://www.legifrance.gouv.fr/jorf/id/JORFSCTA000043879099",
    "Centre-Val de Loire prefecture floods page: https://www.prefectures-regions.gouv.fr/centre-val-de-loire/Actualites/Principales/Inondations-en-region-Centre-Val-de-Loire",
    "Government return-experience report on the 2021 floods: https://www.economie.gouv.fr/cge/retour-inondations2021",
]


def pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def pct_number(value: float) -> str:
    return f"{value:.1f}%"


def pp_delta(high: float, low: float) -> str:
    return f"{(high - low) * 100:.1f} pp"


def int_str(value: int | float) -> str:
    return f"{int(round(float(value))):,}"


def add_title_page(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(REPORT_TITLE)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(REPORT_SUBTITLE)
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(71, 85, 105)

    document.add_paragraph("")

    meta_table = document.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    meta_rows = [
        ("Prepared on", "2026-06-11"),
        ("Scope", "France, processed JRC-Gaspar comparison outputs covering 2015-2024."),
        ("Main purpose", "Consolidate the national comparison, the commune-versus-department reading, the horizon audit, the July 2021 mismatch note, and the app interpretation in one Word report."),
        ("Recommended baseline", "Use the 30-day department-level outputs for the main narrative, and use commune-level outputs plus the app for diagnosis and manual review."),
    ]
    for row_cells, (label, value) in zip(meta_table.rows, meta_rows, strict=True):
        row_cells.cells[0].text = label
        row_cells.cells[1].text = value
        set_cell_shading(row_cells.cells[0], "E2E8F0")
        for cell in row_cells.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    document.add_paragraph("")

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "This report synthesizes the existing local analyses already present in the repository. "
        "It does not introduce a new matching algorithm; it explains and consolidates the results we already produced."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(51, 65, 85)
    document.add_page_break()


def set_table_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = Inches(width)


def format_table_text(cell, *, font_size: int = 9, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.bold = bold


def add_matrix_table(
    document: Document,
    *,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    header_fill: str = "E2E8F0",
    font_size: int = 9,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_widths(table, widths)

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], header_fill)
        format_table_text(header_cells[idx], font_size=font_size, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            format_table_text(cells[idx], font_size=font_size)
    set_table_widths(table, widths)


def find_share(coverage: pd.DataFrame, level: str, measurement: str, source: str) -> float:
    mask = (coverage["level"] == level) & (coverage["measurement"] == measurement)
    return float(coverage.loc[mask, f"{source}_match_share"].iloc[0])


def build_commune_vs_department_rows(cover7: pd.DataFrame, cover30: pd.DataFrame) -> list[list[str]]:
    rows: list[list[str]] = []
    for label, cover in [("7 days", cover7), ("30 days", cover30)]:
        jrc_commune = find_share(cover, "commune", "unique_events", "jrc")
        jrc_department = find_share(cover, "department", "unique_events", "jrc")
        gaspar_commune = find_share(cover, "commune", "unique_events", "gaspar")
        gaspar_department = find_share(cover, "department", "unique_events", "gaspar")
        rows.append(
            [
                label,
                pct(jrc_commune),
                pct(jrc_department),
                pp_delta(jrc_department, jrc_commune),
                pct(gaspar_commune),
                pct(gaspar_department),
                pp_delta(gaspar_department, gaspar_commune),
            ]
        )
    return rows


def add_source_pack_section(document: Document) -> None:
    document.add_heading("Report scope and source pack", level=1)
    document.add_paragraph(
        "This document is a consolidation report. It brings together the earlier national comparison outputs, "
        "the commune-activity app logic, the 2015-2024 horizon audit, and the July 2021 external-evidence note."
    )
    add_matrix_table(
        document,
        headers=["Repository artifact", "Role in this report"],
        rows=[[path, role] for path, role in SOURCE_PACK],
        widths=[2.8, 3.7],
        font_size=9,
    )
    document.add_paragraph("")
    add_bullet_list(
        document,
        [
            "The national comparison sections use the processed 7-day and 30-day output folders directly.",
            "The horizon section uses the saved region-quarter ranking and its mapped examples.",
            "The July 2021 section uses the earlier external-evidence note as an interpretation layer on top of the local comparison.",
            "The app section explains how the same commune-activity logic can be used interactively for validation and manual QA.",
        ],
    )


def add_data_and_method_section(
    document: Document,
    *,
    comp30: dict[str, float | int | str],
    gaspar_workbook: pd.DataFrame,
    gaspar_risk_counts: dict[str, int],
) -> None:
    document.add_heading("Data and matching logic", level=1)
    document.add_paragraph(
        "The comparison is between two very different source families: JRC flood-depth events derived from satellite-based products, "
        "and Gaspar catastrophe-naturelle recognition rows that follow an administrative process. "
        "That source asymmetry is the central reason the outputs require careful interpretation."
    )

    document.add_heading("JRC flood-depth source", level=2)
    add_dataset_table(
        document,
        [
            ("Dataset name", "European Satellite-Derived Flood Depth Maps"),
            ("Institutional source", "Copernicus Emergency Management Service / Joint Research Centre"),
            ("Coverage", "Europe and surroundings, land only"),
            ("Time span used in the local France subset", "2015-01-01 to 2024-12-31"),
            ("Raster meaning", "Water depth in centimetres on 20 m pixels"),
            ("Important dummy code", "9999 marks permanent or seasonal water bodies"),
            ("Analytical role here", "Physical flood-footprint event source"),
        ],
    )

    document.add_heading("Gaspar source used in this project", level=2)
    add_dataset_table(
        document,
        [
            ("System", "GASPAR: administrative risk-procedure information system"),
            ("Subset used here", "Processed CatNat flood-related rows only"),
            ("Processed workbook", "data/processed/Gaspar_2015_2024.xlsx (sheet Gaspar20152024FloodsClean)"),
            ("Rows in cleaned flood sheet", int_str(len(gaspar_workbook))),
            ("Canonical commune rows used downstream", int_str(comp30["gaspar_canonical_commune_rows"])),
            ("Unique decree IDs in comparison", int_str(comp30["gaspar_unique_decrees"])),
            ("Unique Gaspar event groups in comparison", int_str(comp30["gaspar_unique_event_uids"])),
            ("Analytical role here", "Administrative recognition-period event proxy"),
        ],
    )
    document.add_paragraph(
        "One Gaspar row is not the same thing as one physical flood event. "
        "For the comparison scripts, the practical event key is gaspar_event_uid = cod_nat_catnat + start date + end date. "
        "That makes matching possible, but it still represents an administrative recognition-period grouping rather than a natural disaster episode reconstructed from first principles."
    )

    wave_label = "Chocs Mécaniques liés à l'action des Vagues"
    add_bullet_list(
        document,
        [
            f"Inondations et/ou Coulées de Boue rows in the cleaned sheet: {int_str(gaspar_risk_counts.get('Inondations et/ou Coulées de Boue', 0))}.",
            f"Inondations Remontée Nappe rows in the cleaned sheet: {int_str(gaspar_risk_counts.get('Inondations Remontée Nappe', 0))}.",
            f"Chocs Mécaniques liés à l'action des Vagues rows in the cleaned sheet: {int_str(gaspar_risk_counts.get(wave_label, 0))}.",
            "The comparison geography starts at commune level and also produces department-level rollups derived from commune codes.",
        ],
    )

    document.add_heading("How the matching works", level=2)
    add_bullet_list(
        document,
        [
            "JRC rasters are first tabularized to LAU and then harmonized to current French INSEE communes.",
            "Gaspar rows are cleaned, deduplicated, filtered to flood-related risks, and resolved to current communes where possible.",
            "The flexible comparison is then run at commune level and at department level.",
            "The direct rule keeps pairs whose start dates are within the chosen window and whose end dates are also within that window.",
            "The broader rules also keep pairs when opposite endpoints are close or when the two periods overlap once the same date window is applied to both sides.",
            "The repository already contains both a 7-day run and a 30-day run, which lets us compare a stricter and a more permissive interpretation.",
        ],
    )


def add_national_results_section(
    document: Document,
    *,
    root: Path,
    comp7: dict[str, float | int | str],
    comp30: dict[str, float | int | str],
    cover7: pd.DataFrame,
    cover30: pd.DataFrame,
    diagnostics30: dict[str, object],
    year_chart: Path,
    sensitivity_chart: Path,
) -> None:
    document.add_heading("National comparison results", level=1)
    document.add_paragraph(
        "The headline result is stable across the project outputs: commune-level event matching is low, but department-level matching is materially higher. "
        "That pattern strongly suggests that the two sources often describe related flood episodes with different commune footprints and timing slices rather than having no overlap at all."
    )

    insert_picture_with_caption(
        document,
        year_chart,
        "Figure 1. Unique JRC event starts by year in the France harmonized subset.",
        6.7,
    )

    add_coverage_table(document, cover7, cover30)
    document.add_paragraph("")
    add_matrix_table(
        document,
        headers=[
            "Window",
            "JRC commune",
            "JRC department",
            "JRC uplift",
            "Gaspar commune",
            "Gaspar department",
            "Gaspar uplift",
        ],
        rows=build_commune_vs_department_rows(cover7, cover30),
        widths=[0.9, 0.9, 0.95, 0.85, 0.95, 0.95, 0.95],
        font_size=9,
    )

    insert_picture_with_caption(
        document,
        sensitivity_chart,
        "Figure 2. Event-match shares under the stricter 7-day rule and the more permissive 30-day rule.",
        6.8,
    )
    insert_picture_with_caption(
        document,
        root / "data/processed/jrc_gaspar_comparison_flexible_30d/plots/comparison_overview.png",
        "Figure 3. 30-day comparison overview. The event-level and row-level views tell different stories and should be read separately.",
        6.8,
    )
    insert_picture_with_caption(
        document,
        root / "data/processed/jrc_gaspar_comparison_flexible_30d/plots/department_event_summary.png",
        "Figure 4. 30-day department-level best event-pair summary. Gaspar-full / JRC-partial pairs dominate the best-match picture.",
        6.8,
    )

    avg_groups_per_decree = float(comp30["gaspar_unique_event_uids"]) / float(comp30["gaspar_unique_decrees"])
    max_date_pairs = int(diagnostics30["gaspar_diagnostics"]["max_unique_date_pairs_within_one_decree"])
    add_bullet_list(
        document,
        [
            f"At commune level, the 30-day run matches {int_str(comp30['matched_jrc_events_commune_level'])} of {int_str(comp30['jrc_unique_events'])} JRC events ({pct(find_share(cover30, 'commune', 'unique_events', 'jrc'))}) and {int_str(comp30['matched_gaspar_event_uids_commune_level'])} of {int_str(comp30['gaspar_unique_event_uids'])} Gaspar event groups ({pct(find_share(cover30, 'commune', 'unique_events', 'gaspar'))}).",
            f"At department level, the same 30-day run rises to {int_str(comp30['matched_jrc_events_department_level'])} matched JRC events ({pct(find_share(cover30, 'department', 'unique_events', 'jrc'))}) and {int_str(comp30['matched_gaspar_event_uids_department_level'])} matched Gaspar event groups ({pct(find_share(cover30, 'department', 'unique_events', 'gaspar'))}).",
            f"The uplift from commune to department is +{pp_delta(find_share(cover30, 'department', 'unique_events', 'jrc'), find_share(cover30, 'commune', 'unique_events', 'jrc'))} for JRC and +{pp_delta(find_share(cover30, 'department', 'unique_events', 'gaspar'), find_share(cover30, 'commune', 'unique_events', 'gaspar'))} for Gaspar in the 30-day run.",
            f"Gaspar expands from {int_str(comp30['gaspar_unique_decrees'])} decree IDs to {int_str(comp30['gaspar_unique_event_uids'])} recognition-period groups in the comparison logic, or about {avg_groups_per_decree:.1f} groups per decree on average. The saved diagnostics also show one decree branching into as many as {max_date_pairs} distinct date pairs.",
            "This asymmetry is why department-level matching is a better narrative baseline, while commune-level mismatches are better treated as signals that deserve diagnosis rather than as automatic source errors.",
        ],
    )


def add_app_section(document: Document) -> None:
    document.add_heading("What the Streamlit app adds", level=1)
    document.add_paragraph(
        "The repository now includes a Streamlit commune-activity app that turns the comparison into a practical manual-QA workflow. "
        "The app does not compare rasters live on the map. Instead, it shows which current communes are active in Gaspar, JRC, or both for a chosen period."
    )
    add_bullet_list(document, APP_FEATURES)
    document.add_paragraph("")
    add_bullet_list(document, APP_READING)
    document.add_paragraph(
        "This matters for interpretation. The app is where the July 2021 and horizon-audit examples become visually understandable: "
        "the user can move from national counts to one date range, one region, and one cluster of communes without changing the underlying reconciliation logic."
    )


def add_horizon_section(document: Document, ranking: pd.DataFrame, root: Path) -> None:
    document.add_page_break()
    document.add_heading("2015-2024 horizon audit", level=1)
    document.add_paragraph(
        "The horizon audit extends the analysis beyond the original national summary by ranking every region-quarter from 2015-Q1 to 2024-Q4 with the same commune-activity logic used in the app. "
        "The ranking therefore measures period-overlap mismatch, not event-pair matching in the flexible comparison script."
    )

    top10 = ranking.sort_values(["mismatch_communes", "overlap_share"], ascending=[False, True]).head(10)
    horizon_rows = []
    for row in top10.itertuples():
        horizon_rows.append(
            [
                row.period_label,
                row.region_name,
                int_str(row.active_communes),
                int_str(row.both),
                int_str(row.gaspar_only),
                int_str(row.jrc_only),
                int_str(row.mismatch_communes),
                pct(float(row.overlap_share)),
            ]
        )
    add_matrix_table(
        document,
        headers=[
            "Period",
            "Region",
            "Active",
            "Both",
            "Gaspar only",
            "JRC only",
            "Mismatch",
            "Overlap",
        ],
        rows=horizon_rows,
        widths=[0.75, 1.55, 0.7, 0.55, 0.8, 0.7, 0.8, 0.65],
        font_size=8,
    )
    document.add_paragraph("")
    add_bullet_list(
        document,
        [
            "The largest mismatch quarters are not confined to one year. They span 2016, 2018, 2021, and 2024.",
            "The strongest high-activity Gaspar-dominant quarter is Centre-Val de Loire, 2016-Q2.",
            "The strongest JRC-dominant quarters cluster around Grand Est and Bourgogne-Franche-Comte in early 2018, with additional very large recent quarters in 2024.",
            "This means the low match rate is not a single-event anomaly. It recurs across different hydrologic settings and different parts of the source horizon.",
        ],
    )

    for image_path, caption in [
        (
            root / "docs/assets/gaspar_jrc_horizon_audit/centre_val_de_loire_2016_q2.png",
            "Figure 5. Centre-Val de Loire, Q2 2016. This is the strongest high-activity Gaspar-dominant quarter in the saved horizon audit.",
        ),
        (
            root / "docs/assets/gaspar_jrc_horizon_audit/grand_est_2018_q1.png",
            "Figure 6. Grand Est, Q1 2018. This is the largest saved region-quarter mismatch and is strongly JRC-dominant.",
        ),
        (
            root / "docs/assets/gaspar_jrc_horizon_audit/grand_est_2024_q2.png",
            "Figure 7. Grand Est, Q2 2024. This recent quarter shows large disagreement in both directions near the end of the source horizon.",
        ),
    ]:
        insert_picture_with_caption(document, image_path, caption, 6.8)


def add_july_case_study(document: Document, ranking: pd.DataFrame, root: Path) -> None:
    document.add_page_break()
    document.add_heading("July 2021 case study", level=1)
    snapshot = JULY_2021_MONTHLY_SNAPSHOT
    document.add_paragraph(
        "The July 2021 mismatch was already audited in a separate external-evidence note. "
        "That note matters because it tests a key question: are the largest Gaspar-only and JRC-only clusters just noise, or do they correspond to real flood evidence on the ground?"
    )
    add_matrix_table(
        document,
        headers=["Monthly app snapshot", "Both active", "Gaspar only", "JRC only"],
        rows=[
            [
                snapshot["period"],
                int_str(snapshot["both"]),
                int_str(snapshot["gaspar_only"]),
                int_str(snapshot["jrc_only"]),
            ]
        ],
        widths=[3.4, 0.9, 1.0, 1.0],
        font_size=9,
    )
    document.add_paragraph("")
    add_bullet_list(
        document,
        [
            "The mismatch is material in both directions: there are many Gaspar-only communes and many more JRC-only communes in the month-level app snapshot.",
            "The earlier evidence note found supporting public evidence on both sides, which means the mismatch should not be read as an automatic false-positive problem for one source only.",
            "The most useful interpretation is that Gaspar and JRC capture related but non-identical aspects of the same flood family: administrative recognition versus mapped inundation footprint.",
        ],
    )

    july_cases = [
        (
            "Grand Est",
            "2021-Q3",
            root / "docs/assets/gaspar_jrc_match_audit/grand_est_2021_q3.png",
            "Figure 8. Grand Est, Q3 2021. The overlap exists, but the two source-specific clusters are much larger than the shared zone.",
        ),
        (
            "Bourgogne-Franche-Comte",
            "2021-Q3",
            root / "docs/assets/gaspar_jrc_match_audit/bourgogne_franche_comte_2021_q3.png",
            "Figure 9. Bourgogne-Franche-Comte, Q3 2021. JRC spreads across a much wider corridor, while Gaspar still preserves real supported pockets.",
        ),
    ]
    for region_name, period_label, image_path, caption in july_cases:
        row = ranking[(ranking["region_name"] == region_name) & (ranking["period_label"] == period_label)].iloc[0]
        document.add_heading(f"{region_name}, {period_label}", level=2)
        document.add_paragraph(
            f"Saved quarter-level counts for this case: active communes {int_str(row.active_communes)}, both {int_str(row.both)}, "
            f"Gaspar only {int_str(row.gaspar_only)}, JRC only {int_str(row.jrc_only)}, overlap share {pct(float(row.overlap_share))}."
        )
        insert_picture_with_caption(document, image_path, caption, 6.8)

    document.add_heading("Evidence-supported July 2021 clusters", level=2)
    add_matrix_table(
        document,
        headers=["Cluster", "Side", "Evidence strength", "Reading"],
        rows=[list(row) for row in JULY_2021_EVIDENCE_ROWS],
        widths=[2.15, 0.8, 1.0, 2.55],
        font_size=8,
    )
    document.add_paragraph("")
    add_bullet_list(
        document,
        [
            "Gaspar-only does not automatically mean false positive. The earlier note found strong support in Jura and Meuse, and additional support in Seine-et-Marne.",
            "JRC-only does not automatically mean false positive either. The earlier note found support in Haute-Saone, Ardennes, the Aisne-Ardennes floodplain sector, and Saone-et-Loire.",
            "The July 2021 case therefore behaves like a real source disagreement rather than a simple cleaning failure.",
        ],
    )


def add_conclusion_section(document: Document) -> None:
    document.add_heading("Overall interpretation", level=1)
    add_bullet_list(
        document,
        [
            "The commune-level match rate is genuinely low in the saved outputs, even after moving from a 7-day to a 30-day window.",
            "The department-level uplift is too large to ignore. It is the clearest sign that commune allocation and timing fragmentation drive much of the mismatch.",
            "Gaspar event counts must be read as administrative recognition-period groups, not as clean physical event IDs.",
            "The horizon audit shows that the mismatch is structural across several major flood families, not confined to one famous month or one region.",
            "The July 2021 case study shows that both Gaspar-only and JRC-only clusters can still correspond to real flood evidence.",
            "Operationally, the strongest workflow is to use the 30-day department-level outputs for the main narrative, and then use commune-level diagnostics plus the app when a specific event family needs manual review.",
        ],
    )

    document.add_heading("Local reading path after this report", level=2)
    add_bullet_list(
        document,
        [
            "Read the 30-day comparison directory first when you want the main national story.",
            "Use the 7-day directory as a stricter sensitivity check rather than as the default benchmark.",
            "Use the app when you need to understand where a mismatch sits geographically and which communes drive it.",
            "Use the horizon and July notes when you need manual, human-readable interpretation rather than raw tables.",
        ],
    )


def add_references_section(document: Document) -> None:
    document.add_heading("References and project artifacts", level=1)
    document.add_paragraph(
        "Selected external references used across the earlier notes and the comparison documentation:"
    )
    add_bullet_list(document, SELECTED_EXTERNAL_REFERENCES)


def build_report(args: argparse.Namespace) -> Path:
    root = Path(args.project_root).resolve()
    out_path = Path(args.output).resolve()
    ensure_dir(out_path.parent)

    assets_dir = out_path.parent / "_report_assets_full"
    ensure_dir(assets_dir)

    comp7 = read_metric_csv(root / "data/processed/jrc_gaspar_comparison_flexible_7d/comparison_summary.csv")
    comp30 = read_metric_csv(root / "data/processed/jrc_gaspar_comparison_flexible_30d/comparison_summary.csv")
    cover7 = read_coverage_csv(root / "data/processed/jrc_gaspar_comparison_flexible_7d/coverage_overview.csv")
    cover30 = read_coverage_csv(root / "data/processed/jrc_gaspar_comparison_flexible_30d/coverage_overview.csv")
    diagnostics30 = json.loads(
        (root / "data/processed/jrc_gaspar_comparison_flexible_30d/details/comparison_diagnostics.json").read_text(
            encoding="utf-8"
        )
    )

    jrc = pd.read_csv(
        root / "data/processed/france_lau_insee_documentation/events_fr_insee_long.csv",
        parse_dates=["start_date", "end_date"],
    )
    event_col = "jrc_event_id" if "jrc_event_id" in jrc.columns else "event_id"
    jrc_events = jrc[[event_col, "start_date", "end_date"]].drop_duplicates()
    jrc_events["year"] = jrc_events["start_date"].dt.year
    jrc_year_counts = {int(k): int(v) for k, v in jrc_events.groupby("year").size().to_dict().items()}

    gaspar = pd.read_csv(
        root / "data/processed/jrc_gaspar_comparison_flexible_30d/details/gaspar_commune_events_canonical.csv",
        parse_dates=["gaspar_start_date", "gaspar_end_date"],
    )
    gaspar["year"] = gaspar["gaspar_start_date"].dt.year
    gaspar_year_counts = {
        int(k): int(v)
        for k, v in gaspar[["gaspar_event_uid", "year"]].drop_duplicates().groupby("year").size().to_dict().items()
    }
    gaspar_workbook = pd.read_excel(
        root / "data/processed/Gaspar_2015_2024.xlsx",
        sheet_name="Gaspar20152024FloodsClean",
    )
    gaspar_risk_counts = {
        str(k): int(v) for k, v in gaspar_workbook["lib_risque_jo"].value_counts().to_dict().items()
    }

    year_chart = assets_dir / "jrc_events_by_year.png"
    create_year_bar_chart(jrc_year_counts, year_chart)

    sensitivity_chart = assets_dir / "strict_vs_flexible_event_match_shares.png"
    sensitivity_rows = [
        (
            "JRC commune unique events",
            numeric(cover7[(cover7["level"] == "commune") & (cover7["measurement"] == "unique_events")]["jrc_match_share"].iloc[0]),
            numeric(cover30[(cover30["level"] == "commune") & (cover30["measurement"] == "unique_events")]["jrc_match_share"].iloc[0]),
        ),
        (
            "JRC department unique events",
            numeric(cover7[(cover7["level"] == "department") & (cover7["measurement"] == "unique_events")]["jrc_match_share"].iloc[0]),
            numeric(cover30[(cover30["level"] == "department") & (cover30["measurement"] == "unique_events")]["jrc_match_share"].iloc[0]),
        ),
        (
            "Gaspar commune unique event groups",
            numeric(cover7[(cover7["level"] == "commune") & (cover7["measurement"] == "unique_events")]["gaspar_match_share"].iloc[0]),
            numeric(cover30[(cover30["level"] == "commune") & (cover30["measurement"] == "unique_events")]["gaspar_match_share"].iloc[0]),
        ),
        (
            "Gaspar department unique event groups",
            numeric(cover7[(cover7["level"] == "department") & (cover7["measurement"] == "unique_events")]["gaspar_match_share"].iloc[0]),
            numeric(cover30[(cover30["level"] == "department") & (cover30["measurement"] == "unique_events")]["gaspar_match_share"].iloc[0]),
        ),
    ]
    create_sensitivity_chart(sensitivity_rows, sensitivity_chart)

    ranking = pd.read_csv(root / "docs/assets/gaspar_jrc_horizon_audit/region_quarter_mismatch_2015_2024.csv")

    document = Document()
    set_default_font(document)
    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    add_title_page(document)
    add_executive_summary(
        document,
        [
            f"The France comparison uses {int_str(comp30['jrc_unique_events'])} unique JRC events and {int_str(comp30['gaspar_unique_event_uids'])} Gaspar recognition-period groups derived from {int_str(comp30['gaspar_unique_decrees'])} decree IDs.",
            f"Commune-level event matching remains low even under the 30-day rule: {pct(find_share(cover30, 'commune', 'unique_events', 'jrc'))} for JRC and {pct(find_share(cover30, 'commune', 'unique_events', 'gaspar'))} for Gaspar.",
            f"Department-level event matching is much higher: {pct(find_share(cover30, 'department', 'unique_events', 'jrc'))} for JRC and {pct(find_share(cover30, 'department', 'unique_events', 'gaspar'))} for Gaspar in the 30-day run.",
            "The horizon audit shows that this mismatch is structural across several flood families from 2016 to 2024 rather than an isolated one-off case.",
            "The July 2021 evidence note found support for both Gaspar-only clusters and JRC-only clusters, which means the disagreement is not just noise on one side.",
            "The Streamlit app turns those statistics into a map-based validation workflow that can be used for manual interpretation and QA.",
        ],
    )

    add_source_pack_section(document)
    add_data_and_method_section(
        document,
        comp30=comp30,
        gaspar_workbook=gaspar_workbook,
        gaspar_risk_counts=gaspar_risk_counts,
    )
    add_national_results_section(
        document,
        root=root,
        comp7=comp7,
        comp30=comp30,
        cover7=cover7,
        cover30=cover30,
        diagnostics30=diagnostics30,
        year_chart=year_chart,
        sensitivity_chart=sensitivity_chart,
    )
    add_app_section(document)
    add_horizon_section(document, ranking, root)
    add_july_case_study(document, ranking, root)
    add_conclusion_section(document)
    add_file_guide(document)
    add_limitations_section(document, comp30)
    add_references_section(document)

    document.save(out_path)
    return out_path


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Build a full Word report for the France JRC vs Gaspar comparison."
    )
    parser.add_argument(
        "--project-root",
        default=".",
        help="Project root containing the processed comparison outputs.",
    )
    parser.add_argument(
        "--output",
        default=r"data\processed\JRC_Gaspar_France_Full_Analysis_Report.docx",
        help="Output DOCX path.",
    )
    args = parser.parse_args()
    out_path = build_report(args)
    print(out_path)


if __name__ == "__main__":
    main()
